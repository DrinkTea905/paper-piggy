# -*- coding: utf-8 -*-
"""
本地知识库 web 服务（检索 + RAG对话 + 三档索引，一个进程全包）。
数据源：直接读 zotero.sqlite（绕过 Better BibTeX）。
两种用法：① 内置 chat（用户自备 LLM key）② 检索 API（任何 agent 调 /search，不需 key）。
错误日志：后端异常 + 前端上报都写 logs/errors.log，GET /errors 可导出（方便反馈问题）。
"""
import sys, os, json, time, threading, subprocess, traceback, re
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
import config as C
import retriever as R
import llm as L
import wiki_store as W
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional, List
import uvicorn

app = FastAPI(title="本地知识库")
#  proc/cancelled：整库深索(scope=all)跑在 subprocess 里，此前不留句柄 → 一旦开始无法停，
#  可能空跑数小时并烧掉 API 额度。存下 Popen 供 POST /build/cancel 终止。
BUILD = {"running": False, "stage": None, "log": [], "started": None, "rc": None,
         "proc": None, "cancelled": False,
         "bulk": False}   # BF35：当前 deep 构建是否整库深索(scope=all)，队列批次为 False
# B1：build 守卫锁——把「判 running + 置 True」做成原子并在调用线程内同步置位，
# 杜绝多路触发并发起两个 build_all 子进程写坏同一 LanceDB/bm25/papers.jsonl。
_BUILD_LOCK = threading.Lock()

@app.middleware("http")
async def _no_cache(request, call_next):
    """前端频繁迭代阶段禁用浏览器缓存，避免 HTML/JS 版本不匹配。"""
    resp = await call_next(request)
    p = request.url.path
    if p == "/" or p.startswith("/static"):
        resp.headers["Cache-Control"] = "no-store, must-revalidate"
        # 任务六：给文本类响应补 charset=utf-8，避免 WebView2 按系统 GBK 解码整页出现乱码
        # （StaticFiles 在 Windows 上常返回不带 charset 的 text/html、text/css）。
        ct = resp.headers.get("content-type", "")
        if ct and "charset" not in ct.lower() and (
                ct.startswith("text/") or "javascript" in ct or "json" in ct):
            resp.headers["content-type"] = ct + "; charset=utf-8"
    return resp
ERRLOG = C.LOGS / "errors.log"

def _last4(s):
    """K3：返回 key 的末4位（用于掩码显示「已填」状态）；不足4位或空则返回 ""。绝不回完整明文。"""
    s = (s or "").strip()
    return s[-4:] if len(s) >= 4 else ""

# ── 错误日志 ──────────────────────────────────────────────
def log_error(where, err, tb=""):
    try:
        ERRLOG.parent.mkdir(parents=True, exist_ok=True)
        with open(ERRLOG, "a", encoding="utf-8") as f:
            f.write(f"\n===== {time.strftime('%Y-%m-%d %H:%M:%S')} · {where} =====\n{err}\n")
            if tb:
                f.write(tb + "\n")
    except Exception:
        pass

@app.exception_handler(Exception)
async def _all_exc(request: Request, exc: Exception):
    log_error(f"{request.method} {request.url.path}", repr(exc), traceback.format_exc())
    return JSONResponse({"error": "服务器内部错误（已记入错误日志，可在设置里导出）", "detail": str(exc)}, status_code=500)

class LogQ(BaseModel):
    level: str = "error"
    msg: str = ""
    ctx: str = ""

@app.post("/log")
def client_log(q: LogQ):
    log_error(f"前端[{q.level}] {q.ctx}", q.msg)
    return {"ok": True}

@app.get("/errors")
def get_errors(n: int = 200):
    if not ERRLOG.exists():
        return {"errors": "（暂无错误记录）", "lines": 0}
    lines = ERRLOG.read_text(encoding="utf-8", errors="replace").splitlines()
    return {"errors": "\n".join(lines[-n:]), "lines": len(lines)}

@app.post("/errors/clear")
def clear_errors():
    try:
        if ERRLOG.exists():
            ERRLOG.unlink()
    except Exception:
        pass
    return {"ok": True}

# ── 自动更新：源变化(Zotero 新条目/文件夹新 PDF)时后台定时增量更新 ──
AUTO = {"sig": None, "last": 0}
def _source_signature():
    """返回当前数据源的"内容指纹"，变化即代表有新增/改动。"""
    try:
        import settings as S
        if S.source() == "folder":
            import folder_source as FS
            pdfs = FS.scan(S.folder_dir())
            mt = max((Path(p).stat().st_mtime for p in pdfs), default=0)
            return f"folder:{len(pdfs)}:{int(mt)}"
        import zotero_source as Z
        dd = Z.detect_data_dir()
        sq = (Path(dd) / "zotero.sqlite") if dd else None
        if not (sq and sq.exists()):
            return "zotero:none"
        # R7：WAL 模式下新增条目常只落 -wal、主库 mtime 暂不动 → 自动更新检测不到。
        # 把 zotero.sqlite-wal 的 mtime/size 一并纳入指纹，新增即可被捕获。
        sig = f"zotero:{int(sq.stat().st_mtime)}"
        wal = sq.with_name("zotero.sqlite-wal")
        if wal.exists():
            try:
                ws = wal.stat()
                sig += f":wal{int(ws.st_mtime)}:{ws.st_size}"
            except Exception:
                pass
        return sig
    except Exception:
        return None

def _auto_update_loop():
    import settings as S
    time.sleep(20)                         # 等首次加载完
    AUTO["sig"] = _source_signature()      # 建基线：启动时不误触发
    while True:
        try:
            conf = S.load().get("auto_update", {}) or {}
            # BF24：兜底值统一取 settings.DEFAULT（此前 30/15/60 三处各说各话，改设置后行为不可预期）
            interval = max(5, int(conf.get("interval_min", S.DEFAULT["auto_update"]["interval_min"]))) * 60
            time.sleep(60)
            if not conf.get("enabled", True):
                continue
            if time.time() - AUTO["last"] < interval:
                continue
            AUTO["last"] = time.time()
            sig = _source_signature()
            if sig and sig != AUTO["sig"] and not BUILD["running"]:
                stage = "folder" if S.source() == "folder" else "all"   # 只轻量层+语义，深索永远手动
                # A4：只有 build 真正成功(rc==0)才推进 sig，失败留待下轮重试（与 A3 returncode 贯通）。
                def _auto_done(rc, _sig=sig):
                    if rc == 0:
                        AUTO["sig"] = _sig
                        BUILD["log"].insert(0, "[auto] 检测到新增文献，已自动增量更新。")
                    else:
                        BUILD["log"].insert(0, f"[auto] 自动增量更新未成功(rc={rc})，下轮将重试。")
                    try:
                        _drain_deep_queue()   # 沿用旧默认行为：build 后推进自动深索队列
                    except Exception as e:
                        log_error("auto build drain", repr(e))
                _run_build(stage, on_done=_auto_done)
        except Exception as e:
            log_error("auto_update loop", repr(e))
            time.sleep(60)

class AutoUpdateQ(BaseModel):
    enabled: Optional[bool] = None
    interval_min: Optional[int] = None

@app.get("/setup/auto_update")
def get_auto_update():
    import settings as S
    c = S.load().get("auto_update", {}) or {}
    # BF24：兜底统一走 settings.DEFAULT（60 分钟）
    return {"enabled": bool(c.get("enabled", True)),
            "interval_min": int(c.get("interval_min", S.DEFAULT["auto_update"]["interval_min"]))}

@app.post("/setup/auto_update")
def set_auto_update(q: AutoUpdateQ):
    import settings as S
    patch = {}
    if q.enabled is not None: patch["enabled"] = bool(q.enabled)
    if q.interval_min is not None: patch["interval_min"] = max(5, int(q.interval_min))
    S.save({"auto_update": patch})
    c = S.load().get("auto_update", {})
    # BF24：兜底统一走 settings.DEFAULT（60 分钟）
    return {"ok": True, "enabled": bool(c.get("enabled", True)),
            "interval_min": int(c.get("interval_min", S.DEFAULT["auto_update"]["interval_min"]))}

# ── 启动加载 ──────────────────────────────────────────────
@app.on_event("startup")
def _startup():
    try:
        # 建 wiki 目录 + 把 WIKI.md 升到当前 schema 版本。它会被 MCP 整篇下发给 agent，
        # 过期的话 agent 就照着旧规约干活（不知道 entity 页、不知道该 mark_stale 而非覆盖）。
        W.ensure_scaffold()
    except Exception as e:
        log_error("startup ensure_scaffold", repr(e))
    try:
        # BF34：深索标记文件（追加写）在重试/并发下会积累重复行，计数与判断随之虚高——启动自愈去重
        _dedup_deep_marks()
    except Exception as e:
        log_error("startup dedup deep marks", repr(e))
    threading.Thread(target=_safe_load, daemon=True).start()
    threading.Thread(target=_auto_update_loop, daemon=True).start()

def _safe_load():
    try:
        R.load_all()
    except Exception as e:
        log_error("startup load_all", repr(e), traceback.format_exc())
        print("[server] 索引加载失败：", e, flush=True)
    # F38-B：启动后台预热当前学科的期刊分级缓存（首次冷算 20+s，异步不阻塞；之后落盘永久快）
    try:
        import grading_svc as GS
        GS.warm_async(_load_papers())
    except Exception as e:
        log_error("startup grading warm", repr(e))
    # F10：回灌上次崩溃残留的自动深索队列，续跑
    try:
        _q_boot()
    except Exception as e:
        log_error("startup deep queue boot", repr(e))

# ── 状态 ──────────────────────────────────────────────────
@app.get("/health")
def health():
    mode = R.STATE.get("mode")
    papers = len(R.M.get("papers", {}))          # 去重篇数（题录数，L/F 档都在内存）
    blocks = len(R.M.get("records", {})) if mode == "full" else 0  # 正文块数（段）
    n = blocks if mode == "full" else papers     # 兼容旧字段
    return {"ready": R.STATE.get("ready", False), "mode": mode,
            "n": n, "papers": papers, "blocks": blocks, "building": BUILD["running"],
            "deep": len(_deep_keys())}          # F10：「全部文献」显示 已深索/总数

# ── Agent / MCP 接入信息（给应用内 Agent 页，吐出本机真实可用的接入命令）──
@app.get("/agent/mcp-config")
def agent_mcp_config():
    py = sys.executable                        # 正在跑本服务的解释器＝mcp_server 拉起 server 用的同一个
    # 兜底：若 launcher 用非 python.exe 宿主启动（sys.executable 指向 launcher），
    # 改拼分发版内置的 LocalKB/python/python.exe（已确认该目录存在）。
    try:
        cand = C.APP.parent / "python" / "python.exe"
        if (not py or not Path(py).exists() or Path(py).name.lower() not in ("python.exe", "pythonw.exe", "python", "python3")) and cand.exists():
            py = str(cand)
    except Exception:
        pass
    mcp = str(C.APP / "mcp_server.py")
    def q(s): return '"' + s + '"'             # 路径带空格/中文，命令行统一加引号
    add_core = f'claude mcp add localkb -- {q(py)} {q(mcp)}'
    mcp_json = json.dumps(
        {"mcpServers": {"localkb": {"command": py, "args": [mcp]}}},
        ensure_ascii=False, indent=2)
    codex_toml = (f'[mcp_servers.localkb]\n'
                  f'command = {json.dumps(py)}\n'
                  f'args = [{json.dumps(mcp)}]')
    try:                                       # 真实 MCP 工具数（前端"看到 N 个工具"用它，别写死）
        import mcp_server as MCP
        tool_count = len(MCP.TOOLS)
    except Exception:
        tool_count = 0
    return {
        "python": py, "mcp_server": mcp,
        "daemon_url": C.DAEMON_URL, "server_running": True,
        "wiki_schema_md": str(C.WIKI_SCHEMA_MD), "tool_count": tool_count,
        "claude_cmd": add_core,
        "claude_cmd_user": add_core.replace("claude mcp add localkb ",
                                            "claude mcp add localkb --scope user "),
        "mcp_json": mcp_json, "codex_toml": codex_toml,
    }

@app.get("/setup/detect")
def setup_detect():
    zdir = None
    try:
        import zotero_source as Z
        d = Z.detect_data_dir()
        zdir = str(d) if d else None
    except Exception as e:
        log_error("setup/detect zotero", repr(e))
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    import settings as S
    st = S.load()
    backend = st.get("backend", "local")
    api_key_set = bool((st.get("api") or {}).get("key"))
    models_local = (C.MODELS / "bge-m3-onnx" / "model_quantized.onnx").exists()
    reranker_local = (C.MODELS / "bge-reranker-v2-m3-onnx" / "model_quantized.onnx").exists()
    # 数据源：优先 settings.source（用户/向导已选），否则据 manifest / 探到 zotero 推断
    src = st.get("source")
    if src not in ("zotero", "folder"):
        man_src = str(manifest.get("source") or "")
        src = "folder" if man_src.startswith("folder") else ("zotero" if zdir else None)
    try:
        import folder_meta as FM
        meta_ready = FM.available()
    except Exception:
        meta_ready = api_key_set or bool((st.get("sac") or {}).get("key"))
    return {
        "zotero_dir": zdir,
        "source": src,                             # zotero | folder | None
        "folder_dir": st.get("folder_dir", ""),
        "import_only_pdf": bool(st.get("import_only_pdf")),
        "zotero_detected": bool(zdir),             # 与 source 解耦：探到 zotero≠一定用它
        "meta_ready": bool(meta_ready),            # 抽题录 LLM key 是否就绪（folder 模式用）
        "folder_meta_ready": bool(meta_ready),
        "backend": backend,                        # local | api
        "api_key_set": api_key_set,
        # K3：各 key 末4位掩码（只回末4位、绝不回明文），供前端展示「已填 ••••1234」
        "api_key_last4": _last4((st.get("api") or {}).get("key")),
        "sac_key_last4": _last4((st.get("sac") or {}).get("key")),
        # 引擎就绪：本地模式看模型文件；API 模式看 key 是否已填
        "models_ready": models_local if backend == "local" else api_key_set,
        "reranker_ready": reranker_local if backend == "local" else api_key_set,
        "models_local": models_local, "reranker_local": reranker_local,
        "indexed": C.INDEX_MANIFEST.exists(),
        "mode": R.STATE.get("mode"),
    }

class ConnectQ(BaseModel):
    zotero_dir: Optional[str] = None
    source: Optional[str] = None                  # "folder" 时走文件夹分支
    folder_dir: Optional[str] = None

@app.post("/setup/connect")
def setup_connect(q: ConnectQ):
    """校验数据源并返回条目数。source=folder → 选/建文件夹并计 PDF 数；否则读 zotero.sqlite。"""
    import settings as S
    # BF17b：建库子进程正在读源（zotero 临时副本/受管文件夹），此时切换数据源会与之互踩
    if BUILD["running"]:
        return JSONResponse({"ok": False, "msg": "正在建库，稍后再试"}, status_code=400)
    if q.source == "folder":
        try:
            import folder_source as FS
            p = Path(q.folder_dir or "")
            if not str(p).strip():
                return JSONResponse({"ok": False, "msg": "请指定一个文件夹路径"}, status_code=400)
            p.mkdir(parents=True, exist_ok=True)      # 支持"新建空文件夹"
            n = len(FS.scan(str(p)))
            S.save({"source": "folder", "folder_dir": str(p.resolve())})
            try:
                import folder_meta as FM
                mr = FM.available()
            except Exception:
                mr = False
            return {"ok": True, "source": "folder", "entries": n,
                    "dir": str(p.resolve()), "folder_meta_ready": mr}
        except Exception as e:
            log_error("setup/connect folder", repr(e))
            return JSONResponse({"ok": False, "msg": f"无法创建/访问该文件夹：{e}"}, status_code=400)
    try:
        import zotero_source as Z
        if Z.available(q.zotero_dir):
            papers = Z.load_papers(q.zotero_dir)
            n = len(papers)
            with_pdf = sum(1 for p in papers if p.get("has_pdf"))   # F1：向导计数区分「全库 / 将入库」
            # BF2：用户手选的 zotero 数据目录此前校验完即丢，重启后 zotero_source 又退回自动探测；
            # 空串=没手选（沿用自动探测），键名 zotero_dir 与 settings.DEFAULT / zotero_source 约定一致。
            S.save({"source": "zotero", "zotero_dir": q.zotero_dir or ""})
            return {"ok": True, "source": "zotero.sqlite", "entries": n,
                    "with_pdf": with_pdf, "no_pdf": n - with_pdf,
                    "dir": q.zotero_dir or str(Z.detect_data_dir())}
    except Exception as e:
        log_error("setup/connect zotero", repr(e))
        return JSONResponse({"ok": False, "msg": f"读取 zotero.sqlite 失败：{e}"}, status_code=400)
    return JSONResponse({"ok": False, "msg": "未探测到 zotero.sqlite（请确认已安装 Zotero 且库中有文献）"}, status_code=400)

class FolderQ(BaseModel):
    folder_dir: str

@app.post("/setup/folder")
def setup_folder(q: FolderQ):
    """保存文件夹模式选择（选/建文件夹）。"""
    import settings as S
    try:
        import folder_source as FS
        p = Path(q.folder_dir)
        p.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return JSONResponse({"ok": False, "msg": f"无法创建/访问该文件夹：{e}"}, status_code=400)
    n = len(FS.scan(str(p)))
    S.save({"source": "folder", "folder_dir": str(p.resolve())})
    try:
        import folder_meta as FM
        mr = FM.available()
    except Exception:
        mr = False
    return {"ok": True, "folder_dir": str(p.resolve()), "pdf_count": n, "folder_meta_ready": mr}

def _default_papers_dir():
    """文件夹模式的建议默认目录：应用自己的数据目录旁 <HOME>/papers（DATA=HOME/data）。"""
    try:
        return str((C.DATA.parent / "papers"))
    except Exception:
        return ""

@app.get("/setup/folder_default")
def setup_folder_default():
    return {"default_dir": _default_papers_dir()}

class ImportOptQ(BaseModel):
    only_pdf: bool = False

@app.post("/setup/import_only_pdf")
def setup_import_only_pdf(q: ImportOptQ):
    """Zotero 模式：是否只导入有 PDF 的条目（切换后需重建即时索引才生效）。"""
    import settings as S
    S.save({"import_only_pdf": bool(q.only_pdf)})
    return {"ok": True, "only_pdf": bool(q.only_pdf)}

@app.post("/setup/open_folder")
def setup_open_folder():
    """在系统文件管理器里打开受管文件夹（副本#7：跳转让用户直接放 PDF）。"""
    import settings as S
    d = S.folder_dir() or _default_papers_dir()
    try:
        Path(d).mkdir(parents=True, exist_ok=True)
        if sys.platform == "win32":
            os.startfile(d)  # noqa
        elif sys.platform == "darwin":
            subprocess.Popen(["open", d])
        else:
            subprocess.Popen(["xdg-open", d])
        return {"ok": True, "dir": d}
    except Exception as e:
        return JSONResponse({"ok": False, "msg": str(e), "dir": d}, status_code=400)

@app.post("/setup/pick_folder")
def setup_pick_folder():
    """pywebview 原生文件夹选择对话框。无 pywebview（浏览器回退）时返回 ok=False，前端隐藏「浏览…」。"""
    try:
        import webview
        w = webview.windows[0] if getattr(webview, "windows", None) else None
        if not w:
            return {"ok": False, "msg": "无原生窗口"}
        res = w.create_file_dialog(webview.FOLDER_DIALOG)
        d = (res[0] if isinstance(res, (list, tuple)) and res else res) or ""
        return {"ok": bool(d), "dir": str(d)}
    except Exception as e:
        return {"ok": False, "msg": str(e)}

class BackendQ(BaseModel):
    backend: str = "local"                       # local | api
    base: Optional[str] = None
    key: Optional[str] = None
    embed_model: Optional[str] = None
    rerank_model: Optional[str] = None

@app.post("/setup/backend")
def setup_backend(q: BackendQ):
    """保存检索引擎后端选择（本地/API）。API 模式存 SiliconFlow 等的 key。"""
    import settings as S
    patch = {"backend": "api" if q.backend == "api" else "local"}
    api = {}
    if q.base: api["base"] = q.base
    if q.key is not None: api["key"] = q.key
    if q.embed_model: api["embed_model"] = q.embed_model
    if q.rerank_model: api["rerank_model"] = q.rerank_model
    if api: patch["api"] = api
    st = S.save(patch)
    # 一致性提醒：若已建库且原引擎与新选不同，语义/深索向量会不匹配，需重建
    warn = None
    if C.INDEX_MANIFEST.exists():
        try:
            man = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8"))
            built = man.get("backend")
            if built and built != st.get("backend") and (man.get("meta_done") or R.STATE.get("mode") == "full"):
                warn = "检索引擎已切换，但现有语义/深索索引是用另一引擎建的，向量不兼容——建议重建索引。"
        except Exception:
            pass
    return {"ok": True, "backend": st.get("backend"),
            "api_key_set": bool((st.get("api") or {}).get("key")), "warn": warn}

@app.post("/setup/test_api")
def setup_test_api(q: BackendQ):
    """用一小段文本测试 API：嵌入 + 重排是否可用，返回向量维度/延迟。"""
    import time as _t
    from embedder import APIEmbedder
    from reranker import APIReranker
    import settings as S
    a = S.api_conf()
    base = q.base or a.get("base"); key = q.key if q.key is not None else a.get("key", "")
    em = q.embed_model or a.get("embed_model", "BAAI/bge-m3")
    rm = q.rerank_model or a.get("rerank_model", "BAAI/bge-reranker-v2-m3")
    if not key:
        # 诊断接口：测试“失败”不是 HTTP 错误，返回 200 + ok:false，让前端 jpost 能读到 msg
        # （否则前端只会显示无意义的“/setup/test_api 400”，真实原因被吞）。
        return JSONResponse({"ok": False, "msg": "未填 API key"}, status_code=200)
    try:
        t0 = _t.time()
        v = APIEmbedder(base, key, em).encode(["测试：涉罪未成年人分流转处"])
        sc = APIReranker(base, key, rm).scores("未成年人分流", ["涉罪未成年人分流转处", "无关文本"])
        return {"ok": True, "dim": int(v.shape[1]), "rerank_ok": len(sc) == 2,
                "latency_ms": int((_t.time() - t0) * 1000)}
    except Exception as e:
        return JSONResponse({"ok": False, "msg": f"API 测试失败：{e}"}, status_code=200)

class SacQ(BaseModel):
    enabled: Optional[bool] = None
    generator: Optional[str] = None       # K2：server（服务端用API Key）| agent（交给Agent）| off（不生成）
    base: Optional[str] = None
    key: Optional[str] = None
    model: Optional[str] = None

@app.get("/setup/sac")
def get_sac():
    import settings as S, sac as SAC
    sc = S.sac_conf()
    return {"enabled": bool(sc.get("enabled")), "generator": sc.get("generator"),
            "base": sc.get("base"), "model": sc.get("model"),
            "key_set": bool(sc.get("key")), "key_last4": _last4(sc.get("key")),   # K3 掩码
            "effective_ready": SAC.enabled()}

@app.post("/setup/sac")
def setup_sac(q: SacQ):
    """配置深索摘要生成方（K2 generator 三选一）。key 空时会自动复用 API 后端的 key。
       generator=server→服务端自动生成；agent→交给 Agent（服务端不生成）；off→不生成。"""
    import settings as S, sac as SAC
    patch = {}
    if q.generator in ("server", "agent", "off"):
        patch["generator"] = q.generator
        patch["enabled"] = (q.generator == "server")     # 与旧 enabled 同步（向后兼容读 enabled 的代码）
    elif q.enabled is not None:                           # 老前端只传 enabled 时的兼容路径
        patch["enabled"] = bool(q.enabled)
        patch["generator"] = "server" if q.enabled else "off"
    # BF契约3：判空一律用 is not None——空字符串=用户清空该项要落盘；
    # 旧的 if q.base 会把"清空 base/model"静默吞掉（前端 base/model 每次都发，key 仅 dirty 才发）。
    if q.base is not None: patch["base"] = q.base
    if q.key is not None: patch["key"] = q.key
    if q.model is not None: patch["model"] = q.model
    S.save({"sac": patch})
    sc = S.sac_conf()
    return {"ok": True, "generator": sc.get("generator"), "effective_ready": SAC.enabled()}

# ── 期刊分级学科（整库锁定单学科；journal_grading 期刊权重引擎用）──
class DiscQ(BaseModel):
    discipline: Optional[str] = None

@app.get("/setup/discipline")
def get_discipline():
    """当前锁定学科 + 可选学科清单（供设置面板下拉）。personal=True 的是个人档（如 law_personal）。"""
    import settings as S
    cur = S.discipline()
    items = []
    try:
        import journal_grading as JG
        for did, meta in JG.load_data().disciplines().items():
            items.append({"id": did, "name": meta.get("name", did),
                          "personal": did.endswith("_personal")})
    except Exception:
        items = [{"id": cur, "name": cur, "personal": cur.endswith("_personal")}]
    return {"current": cur, "disciplines": items}

@app.post("/setup/discipline")
def setup_discipline(q: DiscQ):
    """切换整库锁定学科。检索期动态读设置——保存后下次检索即生效，无需重建索引。
       并后台预热新学科的分级分布缓存，让库总览/浏览稍后刷新即见新口径（不阻塞本请求）。"""
    import settings as S
    if q.discipline:
        S.save({"journal_discipline": q.discipline})
        try:
            import grading_svc as GS
            GS.warm_async(_load_papers())
        except Exception as e:
            log_error("discipline warm", repr(e))
    return {"ok": True, "current": S.discipline()}

@app.post("/setup/reset")
def setup_reset():
    """设置页「恢复默认」：settings.json 覆盖为默认（清 API/SAC key、学科回标准法学、后端回本地）。
       浏览器里的对话 LLM key 存 localStorage，由前端另清。"""
    import settings as S
    st = S.reset()
    return {"ok": True, "backend": st.get("backend"), "discipline": st.get("journal_discipline")}

@app.get("/setup/models_status")
def setup_models_status():
    """本地模型是否齐全（供本地模式的下载步骤判断）。"""
    try:
        import models_bootstrap as MB
        miss = MB.missing_models()
        return {"present": not miss, "missing": miss}
    except Exception as e:
        log_error("setup/models_status", repr(e))
        return {"present": False, "missing": ["bge-m3-onnx", "bge-reranker-v2-m3-onnx"], "err": str(e)}

@app.post("/setup/download_models")
def setup_download_models():
    """本地模式：从云端下载缺失模型，SSE 汇报进度。"""
    import models_bootstrap as MB
    def gen():
        q = []
        def cb(name, done, total, phase):
            q.append({"name": name, "done": done, "total": total, "phase": phase})
        import threading
        state = {"done": False, "ok": False, "msg": ""}
        def work():
            try:
                ok, msg = MB.ensure_models(progress_cb=cb, log=lambda *a: None)
                state["ok"], state["msg"] = ok, msg
            except Exception as e:
                state["ok"], state["msg"] = False, str(e)
                log_error("download_models", repr(e))
            state["done"] = True
        th = threading.Thread(target=work, daemon=True); th.start()
        while not state["done"] or q:
            if q:
                yield "data: " + json.dumps(q.pop(0), ensure_ascii=False) + "\n\n"
            else:
                time.sleep(0.15)
                yield ": keepalive\n\n"
        yield "data: " + json.dumps({"final": True, "ok": state["ok"], "msg": state["msg"]}, ensure_ascii=False) + "\n\n"
        yield "data: [DONE]\n\n"
    return StreamingResponse(gen(), media_type="text/event-stream")

@app.get("/index/status")
def index_status():
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    # BF33：去重计数——标记文件出现重复行时 len(split()) 虚高，与 /health 的 _deep_keys()（set）口径不一
    ek = C.STATE / "embedded_keys.txt"
    deep = len(set(ek.read_text(encoding="utf-8").split())) if ek.exists() else 0
    meta_done = len(C.META_EMBEDDED.read_text(encoding="utf-8").split()) if C.META_EMBEDDED.exists() else 0
    # C1/A2：扫描件/无文本篇数（记在 deep_no_text.txt，不算已深索）——供前端与深索汇总提示。
    nt = C.STATE / "deep_no_text.txt"
    deep_no_text = len(nt.read_text(encoding="utf-8").split()) if nt.exists() else 0
    with _Q_LOCK:
        q_pending = len(QUEUE["pending"]); q_inflight = len(QUEUE["in_flight"])
    return {
        "mode": R.STATE.get("mode"), "ready": R.STATE.get("ready", False),
        "light_done": manifest.get("light_done", False), "source": manifest.get("source"),
        "papers": manifest.get("papers", 0), "with_pdf": manifest.get("with_pdf", 0),
        "meta_done": meta_done, "deep_done": deep, "deep_no_text": deep_no_text,
        "building": BUILD["running"], "stage": BUILD["stage"], "log": BUILD["log"][-40:],
        "queue_pending": q_pending, "queue_in_flight": q_inflight,
        # BF14：cancelled=本次构建是否被 /build/cancel 取消；rc=上次构建子进程退出码（未结束为 null）。
        # 前端据此把「取消后仍显示构建中/误报完成」纠正为真实终态。
        "cancelled": bool(BUILD.get("cancelled")), "rc": BUILD.get("rc"),
    }

@app.get("/stats")
def stats_ep():
    if not C.STATS_CACHE.exists():
        return JSONResponse({"error": "尚未建立索引"}, status_code=404)
    s = json.loads(C.STATS_CACHE.read_text(encoding="utf-8"))
    # 深索数以实时 embedded_keys.txt 为准：stats_cache 在建 L 档时算，深索后不重算→会偏旧
    ek = C.STATE / "embedded_keys.txt"
    if ek.exists() and isinstance(s.get("coverage"), dict):
        # BF33：去重计数——重复行会让「已深索 N 篇」虚高甚至超过总篇数
        s["coverage"]["deep_indexed"] = len(_deep_keys())
    # 最近入库补 has_pdf/deep（供前端三态深索按钮，F45/副本#13）
    try:
        pap = _load_papers(); deepk = _deep_keys()
        for r in s.get("recent", []):
            p = pap.get(r.get("key"))
            r["has_pdf"] = bool(p.get("has_pdf")) if p else False
            r["deep"] = is_deep(r.get("key"), deepk)
    except Exception as e:
        # BF18：静默 pass 会把 papers.jsonl 损坏等真故障吞成「最近入库按钮不对」，至少记一笔
        log_error("stats recent enrich", repr(e))
    # F38-B：期刊分级分布按当前锁定学科现算（命中缓存则替换 by_tier/by_journal；
    # 未命中则后台预热、本次先返回建库时的旧分布并标 grading_pending，前端稍后刷新）。
    try:
        import grading_svc as GS
        import settings as _S
        s["grading_discipline"] = _S.discipline()
        try:                                          # F3：概览显示中文学科名
            import journal_grading as JG
            s["grading_discipline_name"] = JG.load_data().disciplines().get(
                _S.discipline(), {}).get("name", _S.discipline())
        except Exception:
            s["grading_discipline_name"] = _S.discipline()
        dist = GS.weight_dist(_load_papers())
        if dist:
            s["by_tier"], s["by_journal"] = dist
            s["grading_pending"] = False
        else:
            s["grading_pending"] = True
    except Exception as e:
        log_error("stats grading", repr(e))
    return s

# ── 浏览：收藏夹 + 推荐「值得读」──────────────────────────
_PC = {"data": None, "mtime": 0}
def _load_papers():
    if not C.PAPERS_JSONL.exists():
        return {}
    mt = C.PAPERS_JSONL.stat().st_mtime
    if _PC["data"] is None or _PC["mtime"] != mt:
        d = {}
        for line in open(C.PAPERS_JSONL, encoding="utf-8"):
            if line.strip():
                p = json.loads(line); d[p["key"]] = p
        _PC["data"] = d; _PC["mtime"] = mt
    return _PC["data"]

def _load_cats():
    f = C.CATEGORIES_DIR / "zotero_collections.json"
    return json.loads(f.read_text(encoding="utf-8")) if f.exists() else {"tree": [], "by_collection": {}, "by_key": {}}

def _deep_keys():
    ek = C.STATE / "embedded_keys.txt"
    return set(ek.read_text(encoding="utf-8").split()) if ek.exists() else set()

def _deep_no_text_keys():
    """C1/A2：扫描件/无可抽文本的 safe_name(stem) 集合（格式同 embedded_keys.txt）。
       这些篇不算已深索、也无法深索——前端据此标「🚫 扫描件·需OCR」而非「未深索」。"""
    nt = C.STATE / "deep_no_text.txt"
    return set(nt.read_text(encoding="utf-8").split()) if nt.exists() else set()

def _dedup_deep_marks():
    """BF34：深索标记文件是追加写，批次重试/进程中断会留下重复行——启动时保序去重、
       tmp+os.replace 原子重写自愈（与 BF33 的读侧去重配套，治本在此）。"""
    for f in (C.STATE / "embedded_keys.txt", C.STATE / "deep_no_text.txt"):
        try:
            if not f.exists():
                continue
            lines = f.read_text(encoding="utf-8").split()
            uniq = list(dict.fromkeys(lines))          # 保序去重
            if len(uniq) < len(lines):
                tmp = f.with_name(f.name + ".tmp")
                tmp.write_text("\n".join(uniq) + "\n", encoding="utf-8")
                os.replace(tmp, f)
                print(f"[server] 已清理 {len(lines) - len(uniq)} 条重复深索标记（{f.name}）", flush=True)
        except Exception as e:
            log_error("dedup deep marks", repr(e))

def _rec_score(p, g=None):
    """值得读打分：期刊权重为主 + 新近度 + 有 PDF（可深读）。
       学科感知权重(g)优先（0–1→0–10），缺则回退旧离散 tier_rank（峰值同为 ~12）。"""
    if g and g.get("weight") is not None:
        s = g["weight"] * 10.0
    else:
        s = (6 - p.get("tier_rank", 6)) * 2.0
    try:
        yr = int(p.get("year") or 0)
    except Exception:
        yr = 0
    if yr >= 2015:
        s += min(3.0, (yr - 2015) * 0.3)
    if p.get("has_pdf"):
        s += 1.0
    return round(s, 2)

def _load_ai_topics():
    f = C.CATEGORIES_DIR / "ai_topics.json"
    return json.loads(f.read_text(encoding="utf-8")) if f.exists() else {"topics": [], "by_key": {}}

# ══════════════════════════════════════════════════════════════════
#  F10/F11：自建「知识库分类」 + 分类→keys 解析 + 自动深索串行队列
# ══════════════════════════════════════════════════════════════════
import uuid
import textutil as T

def _atomic_json_write(path, obj, indent=None):
    """原子写 JSON，对 OneDrive/杀软临时占用导致的 os.replace WinError 5 做重试 + 兜底直写。
       生产环境 LOCALKB_DATA 指向 %LOCALAPPDATA%（非 OneDrive），此重试主要保开发/便携场景稳。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    data = json.dumps(obj, ensure_ascii=False, indent=indent)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(data, encoding="utf-8")
    last = None
    for i in range(6):
        try:
            os.replace(tmp, path)
            return
        except PermissionError as e:
            last = e
            time.sleep(0.15 * (i + 1))
    # 重试仍失败：退化为直写目标（放弃原子性，但保证不丢数据）
    try:
        path.write_text(data, encoding="utf-8")
        try:
            if tmp.exists():
                tmp.unlink()
        except Exception:
            pass
    except Exception:
        raise last

def is_deep(key, deepk=None):
    """唯一的深索判定入口：分类文件存原始 key，embedded_keys.txt 存 safe_name(stem)。"""
    if deepk is None:
        deepk = _deep_keys()
    return T.safe_name(key) in deepk

_KBC_FILE = C.CATEGORIES_DIR / "kb_categories.json"
_KBC_LOCK = threading.Lock()
_KBC = {"data": None, "mtime": 0}

def _kbc_load():
    if not _KBC_FILE.exists():
        return {"version": 1, "categories": [], "updated_at": ""}
    mt = _KBC_FILE.stat().st_mtime
    if _KBC["data"] is None or _KBC["mtime"] != mt:
        _KBC["data"] = json.loads(_KBC_FILE.read_text(encoding="utf-8"))
        _KBC["mtime"] = mt
    return _KBC["data"]

def _kbc_save(doc):
    """原子写（对 OneDrive 占用重试）。调用方须持 _KBC_LOCK。"""
    doc["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    _atomic_json_write(_KBC_FILE, doc, indent=1)
    _KBC["data"] = doc; _KBC["mtime"] = _KBC_FILE.stat().st_mtime

def _kbc_find(doc, cid):
    for c in doc["categories"]:
        if c["id"] == cid:
            return c
    return None

# ── 自动深索串行队列 ──────────────────────────────────────
_DEEP_QUEUE_FILE = C.STATE / "deep_queue.json"
_Q_LOCK = threading.Lock()
# K1：paused=深索暂停标志（不再起新批）；spp=近批「秒/篇」用于 ETA 外推；batch_start=当前批起跑时刻。
QUEUE = {"pending": [], "in_flight": [], "timer": None, "fails": 0,
         "paused": False, "spp": None, "batch_start": None}
_DEEP_BATCH = 50
_DEEP_DEBOUNCE = 4.0
_DEEP_MAX_RETRY = 3        # A3：同一批连续软失败超过此次数则暂停自动重试（避免坏 PDF/余额0 无限重试烧钱）

def _q_persist():
    """调用方须持 _Q_LOCK。"""
    try:
        _atomic_json_write(_DEEP_QUEUE_FILE, {
            "pending": QUEUE["pending"], "in_flight": QUEUE["in_flight"],
            "paused": bool(QUEUE.get("paused")),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S")})
    except Exception as e:
        log_error("deep_queue persist", repr(e))

def _q_boot():
    """startup 调用：回灌上次崩溃残留（in_flight 可能未跑完 → 回 pending，extract 幂等跳过已完成）。"""
    if _DEEP_QUEUE_FILE.exists():
        try:
            d = json.loads(_DEEP_QUEUE_FILE.read_text(encoding="utf-8"))
            QUEUE["pending"] = list(dict.fromkeys((d.get("in_flight") or []) + (d.get("pending") or [])))
            QUEUE["in_flight"] = []
            QUEUE["paused"] = bool(d.get("paused"))    # K1：恢复暂停标志（暂停态跨重启保留）
        except Exception:
            pass
    if QUEUE["pending"] and not QUEUE.get("paused"):
        _drain_deep_queue()

def _queue_keyset():
    with _Q_LOCK:
        return set(QUEUE["pending"]) | set(QUEUE["in_flight"])

def enqueue_deep(keys):
    """入队一批 raw key：剔除已深索/已在队/在跑。返回真正入队的。触发防抖排空。"""
    deepk = _deep_keys()
    with _Q_LOCK:
        have = set(QUEUE["pending"]) | set(QUEUE["in_flight"])
        newk = [k for k in keys if k not in have and not is_deep(k, deepk)]
        if newk:
            QUEUE["pending"].extend(newk); QUEUE["fails"] = 0; _q_persist()   # 新动作复位失败计数，重启自动重试
            if QUEUE["timer"]:
                QUEUE["timer"].cancel()
            QUEUE["timer"] = threading.Timer(_DEEP_DEBOUNCE, _drain_deep_queue)
            QUEUE["timer"].daemon = True; QUEUE["timer"].start()
    return newk

def _drain_deep_queue():
    """把一批 pending 转 in_flight 并起 deep build。撞锁则回滚，等 build 结束再排。"""
    with _Q_LOCK:
        QUEUE["timer"] = None
        # K1：暂停中不再起新批（正在跑的那批已在子进程里自然跑完，队列保留）。
        if BUILD["running"] or QUEUE.get("paused") or not QUEUE["pending"]:
            return
        batch = QUEUE["pending"][:_DEEP_BATCH]
        QUEUE["pending"] = QUEUE["pending"][_DEEP_BATCH:]
        QUEUE["in_flight"] = batch
        QUEUE["batch_start"] = time.time()   # K1：记本批起跑，用于 ETA「秒/篇」外推
        _q_persist()
    started = _run_build("deep", ["--scope", "keys:" + ",".join(batch)], on_done=_on_deep_done)
    if not started:
        with _Q_LOCK:
            QUEUE["pending"] = batch + QUEUE["pending"]
            QUEUE["in_flight"] = []; _q_persist()

def _on_deep_done(rc=0):
    """A3：深索子进程结束回调（收 returncode）。
       rc==0：本批已成功入库，清 in_flight、复位失败计数、继续排空后续。
       rc!=0：软失败——把 in_flight 退回 pending **队首**而非清空（否则该批 key 永久蒸发），
              在重试预算内自动重排；超预算则保留在 pending 但停自动重排，待人工/重启再试。"""
    should_drain = False
    with _Q_LOCK:
        batch = QUEUE["in_flight"]; QUEUE["in_flight"] = []
        # K1：成功批用「本批耗时/篇数」更新 seconds-per-paper 的 EMA，供 /index/queue 外推 ETA。
        if rc == 0 and batch and QUEUE.get("batch_start"):
            el = time.time() - QUEUE["batch_start"]
            spp = el / max(1, len(batch))
            prev = QUEUE.get("spp")
            QUEUE["spp"] = spp if not prev else (prev * 0.5 + spp * 0.5)
        QUEUE["batch_start"] = None
        if rc != 0 and batch:
            QUEUE["fails"] = QUEUE.get("fails", 0) + 1
            QUEUE["pending"] = batch + QUEUE["pending"]        # 退回队首，保住这批 key
            if QUEUE["fails"] <= _DEEP_MAX_RETRY:
                should_drain = True
                BUILD["log"].append(f"[deep] 本批深索失败(rc={rc})，已退回队列重试（第 {QUEUE['fails']} 次）。")
            else:
                BUILD["log"].append(f"[deep] 本批深索连续失败 {QUEUE['fails']} 次，暂停自动重试"
                                    f"（{len(batch)} 篇仍在队列，稍后可重新触发深索）。")
        else:
            QUEUE["fails"] = 0
            should_drain = bool(QUEUE["pending"])
        _q_persist()
    if should_drain:
        _drain_deep_queue()

@app.get("/index/queue")
def deep_queue_status():
    """K1：深索队列/详情/ETA。items=当前在深索或在队首的若干篇 key+标题；
       eta_seconds 按近批速率外推（取不到给 null）。"""
    papers = _load_papers()
    deep_done = len(_deep_keys())
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    with _Q_LOCK:
        pending = len(QUEUE["pending"]); in_flight = len(QUEUE["in_flight"])
        paused = bool(QUEUE.get("paused")); spp = QUEUE.get("spp")
        inflight_keys = list(QUEUE["in_flight"])
        shown = (inflight_keys + list(QUEUE["pending"]))[:8]
    items = [{"key": k, "title": (papers.get(k) or {}).get("title", "")} for k in shown]
    remaining = pending + in_flight
    eta_seconds = int(remaining * spp) if (spp and remaining) else None
    return {"pending": pending, "in_flight": in_flight, "paused": paused,
            "deep_done": deep_done, "with_pdf": manifest.get("with_pdf", 0),
            "eta_seconds": eta_seconds, "items": items,
            # 兼容旧前端字段
            "in_flight_keys": inflight_keys,
            # BF35：bulk=当前 deep 构建是整库深索(scope=all)；队列批次为 false。
            # 前端据此区分「整库深索中」与「队列批次在跑」两种进度文案。
            "bulk": bool(BUILD.get("bulk")),
            "building": BUILD["running"], "stage": BUILD["stage"]}

class DeepPauseQ(BaseModel):
    paused: bool = True

@app.post("/index/deep/pause")
def deep_pause(q: DeepPauseQ):
    """K1：置深索暂停标志。暂停后 _drain_deep_queue 不再起新批（正在跑的自然跑完），队列保留；
       恢复(paused=False)立即尝试继续排空。"""
    with _Q_LOCK:
        QUEUE["paused"] = bool(q.paused)
        _q_persist()
    if not q.paused:
        try:
            _drain_deep_queue()
        except Exception as e:
            log_error("deep resume drain", repr(e))
    return {"ok": True, "paused": bool(q.paused)}

# ── F11：分类 → keys 白名单解析 ──────────────────────────
def _resolve_category_keys(category):
    """把 category 值解析成原始 key 白名单 set；None/"" → None(不过滤)。
       前缀：kbc_（用户分类）/ topic:<id>（AI主题）/ zotero:<path>（收藏夹镜像）。未知前缀→空集(命中0，安全)。"""
    if not category:
        return None
    if category.startswith("kbc_"):
        with _KBC_LOCK:
            doc = _kbc_load(); c = _kbc_find(doc, category)
        return set(c["keys"]) if c else set()
    if category.startswith("topic:"):
        try:
            tid = int(category.split(":", 1)[1])
        except Exception:
            return set()
        ait = _load_ai_topics()
        return {k for k, v in ait.get("by_key", {}).items() if v == tid}
    if category.startswith("zotero:"):
        path = category.split(":", 1)[1]
        return set(_load_cats().get("by_collection", {}).get(path, []))
    return set()

# ── F10：知识库分类 CRUD 端点 ────────────────────────────
@app.get("/kb/categories")
def kb_categories_list():
    papers = _load_papers(); deepk = _deep_keys(); qset = _queue_keyset()
    with _KBC_LOCK:
        doc = _kbc_load()
        cats = [dict(c) for c in doc["categories"]]
    out = []
    for c in cats:
        live = [k for k in c.get("keys", []) if k in papers]
        deep = sum(1 for k in live if is_deep(k, deepk))
        nopdf = sum(1 for k in live if not papers[k].get("has_pdf"))
        pend = sum(1 for k in live if k in qset)
        out.append({"id": c["id"], "name": c["name"], "source": c.get("source", "user"),
                    "count": len(live), "deep_count": deep, "no_pdf": nopdf,
                    "pending": pend, "updated_at": c.get("updated_at", "")})
    return {"categories": out}

class KbCatNewQ(BaseModel):
    name: str
    source: str = "user"

@app.post("/kb/categories")
def kb_cat_create(q: KbCatNewQ):
    name = (q.name or "").strip()
    if not name:
        return JSONResponse({"ok": False, "detail": "分类名不能为空"}, status_code=400)
    cid = "kbc_" + uuid.uuid4().hex[:8]
    now = time.strftime("%Y-%m-%d %H:%M:%S")
    with _KBC_LOCK:
        doc = _kbc_load()
        doc["categories"].append({"id": cid, "name": name, "source": q.source,
            "keys": [], "note": "", "created_at": now, "updated_at": now})
        _kbc_save(doc)
    return {"ok": True, "id": cid, "name": name}

class KbCatRenameQ(BaseModel):
    name: str

@app.patch("/kb/categories/{cid}")
def kb_cat_rename(cid: str, q: KbCatRenameQ):
    name = (q.name or "").strip()
    if not name:
        return JSONResponse({"ok": False, "detail": "分类名不能为空"}, status_code=400)
    with _KBC_LOCK:
        doc = _kbc_load(); c = _kbc_find(doc, cid)
        if not c:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        c["name"] = name; c["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _kbc_save(doc)
    return {"ok": True, "id": cid, "name": name}

@app.delete("/kb/categories/{cid}")
def kb_cat_delete(cid: str):
    with _KBC_LOCK:
        doc = _kbc_load()
        n = len(doc["categories"])
        doc["categories"] = [c for c in doc["categories"] if c["id"] != cid]
        if len(doc["categories"]) == n:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        _kbc_save(doc)
    return {"ok": True, "id": cid}

class KbMembersQ(BaseModel):
    keys: List[str] = []

@app.post("/kb/categories/{cid}/members")
def kb_cat_add_members(cid: str, q: KbMembersQ):
    papers = _load_papers(); deepk = _deep_keys()
    req = [k for k in dict.fromkeys(q.keys or []) if k]
    added, already = [], []
    with _KBC_LOCK:
        doc = _kbc_load(); c = _kbc_find(doc, cid)
        if not c:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        have = set(c["keys"])
        for k in req:
            (already if k in have else added).append(k)
        c["keys"].extend(added)
        c["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _kbc_save(doc)
    will_deep, no_pdf, already_deep = [], [], []
    for k in added:
        p = papers.get(k)
        if not p:
            continue
        if is_deep(k, deepk):
            already_deep.append(k)
        elif p.get("has_pdf"):
            will_deep.append(k)
        else:
            no_pdf.append(k)
    queued = enqueue_deep(will_deep) if will_deep else []
    return {"ok": True, "added": added, "already": already,
            "will_deep": queued, "no_pdf": no_pdf, "already_deep": already_deep,
            "queued": len(queued)}

@app.delete("/kb/categories/{cid}/members")
def kb_cat_del_members(cid: str, q: KbMembersQ):
    rm = set(q.keys or [])
    with _KBC_LOCK:
        doc = _kbc_load(); c = _kbc_find(doc, cid)
        if not c:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        before = len(c["keys"])
        c["keys"] = [k for k in c["keys"] if k not in rm]
        removed = before - len(c["keys"])
        c["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _kbc_save(doc)
    return {"ok": True, "removed": removed}

def _enrich_cat_deep(cats):
    """给收藏夹树每个节点补 count_indexed（该夹+子孙里已深索的篇数，实时算）。
       深索是后置手动动作，建库时算会偏旧——故检索期用 embedded_keys.txt 现算，
       让浏览左树显示「已深索/总数」双数字，而非满屏 0（旧版只显示总数、无深索维度）。"""
    import textutil as T
    deepk = _deep_keys()
    by_col = cats.get("by_collection", {})
    def walk(n):
        direct = by_col.get(n.get("path", ""), [])
        di = sum(1 for k in direct if T.safe_name(k) in deepk)
        for c in n.get("children", []):
            di += walk(c)
        n["count_indexed"] = di
        return di
    for r in cats.get("tree", []):
        walk(r)
    return cats

@app.get("/categories")
def categories():
    return _enrich_cat_deep(_load_cats())

@app.get("/topics")
def topics():
    ait = _load_ai_topics()
    deepk = _deep_keys()   # F10：每主题补已深索数，前端显示 已深索/总篇数
    out = []
    for t in ait.get("topics", []):
        keys = t.get("keys") or []
        deep = sum(1 for k in keys if is_deep(k, deepk))
        out.append({"id": t["id"], "name": t["name"], "size": t["size"], "deep": deep})
    return {"topics": out}

# ── AI 主题聚类：用向量 KMeans 把已索引文献自动归类（无 LLM，簇数自适应）──
TOPICS_BUILD = {"running": False, "msg": ""}
@app.post("/topics/rebuild")
def topics_rebuild():
    """后台重跑 AI 主题聚类（build_ai_topics.py，需已建语义/深索索引）。"""
    if TOPICS_BUILD["running"]:
        return {"ok": False, "msg": "AI 主题正在归类中，请稍候"}
    if BUILD["running"]:
        return {"ok": False, "msg": "正在建库/深索中，请等结束后再归类"}
    if R.STATE.get("mode") != "full":
        return JSONResponse({"ok": False, "msg": "需先深索或建语义层（AI 主题按向量聚类，纯题录库还没有向量）", "need_index": True}, status_code=200)
    def run():
        TOPICS_BUILD["running"] = True; TOPICS_BUILD["msg"] = "归类中…"
        try:
            env = dict(os.environ); env.pop("PYTHONUTF8", None)
            p = subprocess.run([sys.executable, str(C.APP / "build_ai_topics.py")],
                               stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env, timeout=900)
            TOPICS_BUILD["msg"] = "完成" if p.returncode == 0 else f"失败(code={p.returncode})"
        except Exception as e:
            log_error("topics/rebuild", repr(e)); TOPICS_BUILD["msg"] = f"异常：{e}"
        finally:
            TOPICS_BUILD["running"] = False
    threading.Thread(target=run, daemon=True).start()
    return {"ok": True}

@app.get("/topics/status")
def topics_status():
    return {"running": TOPICS_BUILD["running"], "msg": TOPICS_BUILD["msg"]}

# ── AI 抽词「找相似」：从标题抽实词做检索词（默认 AI，缺 key 退本地）──
class SimilarQ(BaseModel):
    title: str = ""
    provider: str = "siliconflow"
    base_url: str = ""
    api_key: str = ""
    model: str = ""

@app.post("/similar/keywords")
def similar_keywords(q: SimilarQ):
    """把一篇标题抽成 2-4 个检索关键词。优先 LLM（准），无 key 退本地实词。"""
    title = (q.title or "").strip()
    if not title:
        return {"ok": True, "keywords": ""}
    base, model = L.resolve(q.provider, q.base_url, q.model)
    key = q.api_key
    # BF8：key 外发守卫（仿 /chat）——检索引擎的 SiliconFlow key 只准回填给 SiliconFlow 的 base，
    # 否则用户任填一个第三方 base 就会把 key 发出去。无 key 时返回空 keywords，前端退本地分词。
    if not key and "siliconflow" in (base or "").lower():
        try:
            import settings as S
            key = (S.api_conf() or {}).get("key", "")
        except Exception:
            key = ""
    if key:
        try:
            msgs = [{"role": "system", "content": "你是学术检索助手。把用户给的论文标题提炼成 2-4 个最核心的中文检索关键词，"
                     "用空格分隔，只输出关键词本身，不要解释、不要标点。"},
                    {"role": "user", "content": title}]
            kw = L.chat_once(msgs, base, key, model or "Qwen/Qwen2.5-7B-Instruct", temperature=0.2, timeout=30)
            kw = " ".join(re.findall(r"[一-鿿A-Za-z0-9]+", kw))[:60]
            if kw:
                return {"ok": True, "keywords": kw, "by": "ai"}
        except Exception as e:
            log_error("similar/keywords", repr(e))
    return {"ok": True, "keywords": "", "by": "none"}   # 前端据空回退本地抽词

@app.get("/similar/{key}")
def similar_vector(key: str, topk: int = 8):
    """C4/F6：真正的向量「找相似」。full 模式用该 key 已存向量（或现场 encode 其标题）在
       LanceDB 做 cosine 近邻、排除自身，results 每条结构与 /search 一致（前端复用 resultCard）。
       light 模式或取不到向量 → {ok:false}，前端回退现有 /similar/keywords 抽词法。"""
    try:
        res = R.neighbors(key, topk)
    except Exception as e:
        log_error("similar vector", repr(e))
        res = None
    if res is None:
        return {"ok": False}
    return {"ok": True, "key": key, "results": res}

# ── 综合层 wiki（答案沉淀 / 按需综述）──────────────────────────
class WikiSaveQ(BaseModel):
    query: str = ""
    answer: str
    sources: List[dict] = []      # 前端传 /chat 的 sources；服务端只认 key、重解析页级引用
    model: str = ""               # generated_by（可信度审计）
    by_agent: bool = False        # True=agent 经 MCP 写回（默认采纳、标 🤖 待人复看/剔除）

@app.post("/wiki/answer")
def wiki_answer(q: WikiSaveQ):
    """Phase 0：把一次问答沉淀成 answer 综合页（存盘 + 入表可检索）。网页由"保存此答案"、agent 由 save_synthesis 触发。"""
    try:
        meta = W.save_answer(q.query, q.answer, q.sources, generated_by=q.model, by_agent=q.by_agent)
        return {"ok": True, "id": meta["id"], "title": meta["title"],
                "indexed": meta.get("indexed", False), "n_sources": len(meta.get("sources", []))}
    except W.WikiWriteDenied as e:
        # 正常的权限拒绝（agent 想覆盖人工核验页），不是故障，不进 errors.log
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except Exception as e:
        log_error("wiki/answer", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/list")
def wiki_list():
    return {"pages": W.list_pages()}

@app.get("/wiki/page/{page_id}")
def wiki_page(page_id: str):
    p = W.get_page(page_id)
    if not p:
        return JSONResponse({"error": "无此综合页", "id": page_id}, status_code=404)
    return p

@app.delete("/wiki/page/{page_id}")
def wiki_delete(page_id: str):
    """§6.4 一键"不保存"：删 md + index 条目 + LanceDB 表行。**仅人用（UI）——不做成 MCP 工具，agent 无删权。**"""
    try:
        r = W.delete_page(page_id)
        if not r.get("deleted"):
            return JSONResponse({"ok": False, "detail": "无此综合页", "id": page_id}, status_code=404)
        return {"ok": True, "id": page_id, **r}
    except Exception as e:
        log_error("wiki/delete", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ── stale 写侧 + by_source 反查出口（此前两者都只建不用）──
class WikiStaleQ(BaseModel):
    stale: bool = True
    reason: str = ""

@app.post("/wiki/stale/{page_id}")
def wiki_stale(page_id: str, q: WikiStaleQ):
    """标记/清除某综合页「已过时」。检索降权立即生效，无需重启。"""
    try:
        return {"ok": True, **W.set_stale(page_id, q.stale, q.reason)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/stale", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

class WikiVerifyQ(BaseModel):
    page_id: str

@app.post("/wiki/verify")
def wiki_verify(q: WikiVerifyQ):
    """W3：人工核验盖章——页面 frontmatter/index/内存三处写 verified_at。
       只给 UI 用，不做成 MCP 工具（核验是人的动作，agent 不得给自己的产出盖章）。"""
    try:
        r = W.set_verified(q.page_id)
        return {"ok": True, "verified_at": r["verified_at"]}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/verify", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/backlinks")
def wiki_backlinks(key: Optional[str] = None, page_id: Optional[str] = None):
    """key=论文 → 引用它的综合页；page_id=页 → 它的来源与互链（含 orphan 判定）。"""
    try:
        return {"ok": True, **W.backlinks(key=key, page_id=page_id)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)
    except Exception as e:
        log_error("wiki/backlinks", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ── 波次2：互链写侧 / 任意页读写 / 体检 / 一源触多页 / 图 / 版本历史 ──
class WikiLinksQ(BaseModel):
    links: List[str] = []
    mode: str = "replace"           # replace | add | remove

@app.post("/wiki/links/{page_id}")
def wiki_set_links(page_id: str, q: WikiLinksQ):
    """写 links —— 把一堆孤岛补成一张图。拒绝自链与断链，返回 skipped。"""
    try:
        return {"ok": True, **W.set_links(page_id, q.links, q.mode)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/links", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

class WikiUpdateQ(BaseModel):
    kind: Optional[str] = None
    title: Optional[str] = None
    content: str = ""
    sources: Optional[List[dict]] = None
    mode: str = "replace"           # replace | append
    links: Optional[List[str]] = None
    model: str = ""
    by_agent: bool = False

@app.post("/wiki/page/{page_id}")
def wiki_update_page(page_id: str, q: WikiUpdateQ):
    """建 / 覆盖 / 追加任意 kind 的 wiki 页（含 entity / overview）。沿用 agent 写权护栏。"""
    try:
        m = W.update_page(page_id, kind=q.kind, title=q.title, content=q.content,
                          sources=q.sources, mode=q.mode, links=q.links,
                          generated_by=q.model, by_agent=q.by_agent)
        return {"ok": True, "id": m["id"], "kind": m.get("kind"), "title": m.get("title"),
                "indexed": m.get("indexed", False), "degraded": m.get("degraded", False),
                "links": m.get("links", []), "n_sources": len(m.get("sources", []))}
    except W.WikiWriteDenied as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)
    except Exception as e:
        log_error("wiki/update", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/lint")
def wiki_lint(min_mentions: int = 2):
    """gist 的 Lint：孤儿页 / 过时页 / 断链 / 无来源 / 降级页 / 缺失概念页。纯读，零副作用。"""
    try:
        return {"ok": True, **W.lint(min_mentions)}
    except Exception as e:
        log_error("wiki/lint", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/graph")
def wiki_graph():
    """图视图数据：节点=页，边=links，孤儿页标出。给应用内的图视图用（不依赖 Obsidian）。"""
    try:
        return {"ok": True, **W.graph()}
    except Exception as e:
        log_error("wiki/graph", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/propose/{source_key}")
def wiki_propose(source_key: str, topk: int = 12):
    """gist 核心：一篇源影响了哪些 wiki 页、每页该怎么改。只建议，不动手。"""
    try:
        return {"ok": True, **W.propose_updates(source_key, topk)}
    except Exception as e:
        log_error("wiki/propose", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/history/{page_id}")
def wiki_history(page_id: str, limit: int = 30):
    try:
        return {"ok": True, **W.page_history(page_id, limit)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/history", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

class WikiRestoreQ(BaseModel):
    rev: str

@app.post("/wiki/restore/{page_id}")
def wiki_restore(page_id: str, q: WikiRestoreQ):
    """回滚某页到历史版本。回滚本身也记一版，可再滚回去。**仅人用**，不做成 MCP 工具。"""
    try:
        return {"ok": True, **W.restore_page(page_id, q.rev)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/restore", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/vcs")
def wiki_vcs_status():
    """版本历史后端：git（开发机）或 snapshot（多数 exe 用户机器上没装 git）。"""
    try:
        import wiki_vcs as V
        return {"ok": True, **V.status()}
    except Exception as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ── Phase 1：按需生成概念/主题综述页（命中缓存 0 成本；LLM 综合 + LLM 命名）──
class WikiSynthQ(BaseModel):
    concept: str = ""
    topic_id: Optional[int] = None
    force: bool = False           # 忽略缓存、强制重生
    provider: str = "siliconflow"
    base_url: str = ""
    api_key: str = ""
    model: str = ""
    topk: int = 8

def _synth_llm(q):
    return {"provider": q.provider, "base_url": q.base_url, "api_key": q.api_key,
            "model": q.model, "topk": q.topk}

def _synth_ret(m):
    return {"ok": True, "id": m["id"], "title": m["title"], "kind": m.get("kind"),
            "cached": m.get("cached", False), "indexed": m.get("indexed", False),
            "n_sources": len(m.get("sources", []))}

@app.post("/wiki/concept")
def wiki_concept(q: WikiSynthQ):
    try:
        return _synth_ret(W.synthesize_concept(q.concept, force=q.force, **_synth_llm(q)))
    except Exception as e:
        log_error("wiki/concept", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/wiki/topic")
def wiki_topic(q: WikiSynthQ):
    try:
        return _synth_ret(W.synthesize_topic(q.topic_id, force=q.force, **_synth_llm(q)))
    except Exception as e:
        log_error("wiki/topic", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/wiki/regenerate/{page_id}")
def wiki_regen(page_id: str, q: WikiSynthQ):
    try:
        return _synth_ret(W.regenerate(page_id, **_synth_llm(q)))
    except Exception as e:
        log_error("wiki/regenerate", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ══════════════════════════════════════════════════════════════════
#  Phase 5：半自动研究助手（选题/框架/带页级引注的资料汇编/建议补文献）
# ══════════════════════════════════════════════════════════════════
class ResearchQ(BaseModel):
    query: str = ""
    topic: str = ""
    topk: int = 14
    force: bool = False
    provider: str = "siliconflow"
    base_url: str = ""
    api_key: str = ""
    model: str = ""
    by_agent: bool = False

def _research_llm(q):
    return {"provider": q.provider, "base_url": q.base_url, "api_key": q.api_key, "model": q.model}

@app.get("/research/pagemap/{key}")
def research_pagemap(key: str):
    """调试：查某篇 PDF页→期刊印刷页码 的映射与 quality。"""
    try:
        import page_map as PM
        doc = PM.build(key)
        if not doc:
            return JSONResponse({"ok": False, "detail": "该篇无提取文本（未深索或扫描件），无法建映射"}, status_code=404)
        return {"ok": True, **doc}
    except Exception as e:
        log_error("research/pagemap", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/digest")
def research_digest(q: ResearchQ):
    """能力二：给一个子题 query → 带印刷页引注的综述 + 覆盖评级 + 缺口。存 kind=digest wiki 页。"""
    try:
        import research_assistant as RA
        m = RA.digest(q.query or q.topic, topk=q.topk, llm=_research_llm(q), force=q.force, by_agent=q.by_agent)
        return {"ok": True, "id": m["id"], "title": m["title"], "kind": m.get("kind"),
                "cached": m.get("cached", False), "indexed": m.get("indexed", False),
                "coverage": m.get("coverage"), "n_sources": len(m.get("sources", []))}
    except W.WikiWriteDenied as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except Exception as e:
        log_error("research/digest", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/scope")
def research_scope(q: ResearchQ):
    """能力一：主题 → 范围映射 + 选题拆解 + 标题候选 + 三级大纲(★/☆)。存 kind=outline wiki 页。"""
    try:
        import research_assistant as RA
        m = RA.scope(q.topic or q.query, topk=q.topk, llm=_research_llm(q), force=q.force, by_agent=q.by_agent)
        return {"ok": True, "id": m["id"], "title": m["title"], "kind": m.get("kind"),
                "cached": m.get("cached", False), "indexed": m.get("indexed", False),
                "n_sources": len(m.get("sources", []))}
    except W.WikiWriteDenied as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except Exception as e:
        log_error("research/scope", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/suggest_sources")
def research_suggest(q: ResearchQ):
    """能力三：覆盖评估 + 脚注引文挖掘缺失文献 + 库内错配（按期刊层级排）。只读，不写库。"""
    try:
        import research_assistant as RA
        return {"ok": True, **RA.suggest_sources(q.topic or q.query, topk=q.topk, llm=_research_llm(q))}
    except Exception as e:
        log_error("research/suggest_sources", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

def _export_docx(page_id: str):
    """把 digest/outline wiki 页导出。有 python-docx 则出 .docx，否则退化为 .md（诚实降级）。
       C3：末尾用 p['sources'] 追加一节「参考文献」（页级引用），产出可直接拿去写作。"""
    p = W.get_page(page_id)
    if not p:
        return JSONResponse({"ok": False, "detail": "无此页"}, status_code=404)
    sources = p.get("sources", []) or []
    def _cite(s):
        if isinstance(s, dict):
            return s.get("citation") or s.get("key") or ""
        return str(s)
    try:
        import docx  # python-docx
        from docx import Document
        doc = Document()
        doc.add_heading(p.get("title", "资料汇编"), level=0)
        for line in (p.get("markdown") or "").splitlines():
            s = line.strip()
            if s.startswith("### "):
                doc.add_heading(s[4:], level=2)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=1)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s and not s.startswith("---"):
                doc.add_paragraph(s)
        if sources:
            doc.add_heading("参考文献", level=1)
            for i, s in enumerate(sources, 1):
                doc.add_paragraph(f"{i}. {_cite(s)}")
        out = C.WIKI_DIGEST_DIR / f"{page_id}.docx"
        doc.save(str(out))
        return FileResponse(str(out), filename=f"{page_id}.docx",
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except ImportError:
        # 降级：返回 markdown 文件（本机未装 python-docx）
        md = p.get("markdown", "")
        if sources:
            md += "\n\n## 参考文献\n\n" + "\n".join(f"{i}. {_cite(s)}" for i, s in enumerate(sources, 1))
        out = C.WIKI_DIGEST_DIR / f"{page_id}.md"
        out.write_text(md, encoding="utf-8")
        return FileResponse(str(out), filename=f"{page_id}.md", media_type="text/markdown")
    except Exception as e:
        log_error("research/export_docx", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/export_docx/{page_id}")
def research_export_docx(page_id: str):
    return _export_docx(page_id)

@app.get("/research/export_docx/{page_id}")
def research_export_docx_get(page_id: str):
    """C3：GET 版——前端用 <a href download> 直接触发下载。"""
    return _export_docx(page_id)

@app.get("/papers")
def papers(collection: Optional[str] = None, topic: Optional[int] = None,
           category: Optional[str] = None, deep: Optional[str] = None,
           sort: str = "recommend", limit: int = 300, offset: int = 0):
    papers = _load_papers(); cats = _load_cats(); deepk = _deep_keys()
    notextk = _deep_no_text_keys()   # C1/A2：扫描件集合
    if category and category.startswith("kbc_"):
        ks = _resolve_category_keys(category) or set()
        items = [papers[k] for k in ks if k in papers]
    elif topic is not None:
        ait = _load_ai_topics()
        tk = [k for k, tid in ait.get("by_key", {}).items() if tid == topic]
        items = [papers[k] for k in tk if k in papers]
    elif collection and collection in cats.get("by_collection", {}):
        items = [papers[k] for k in cats["by_collection"][collection] if k in papers]
    else:
        items = list(papers.values())
    import grading_svc as GS
    out = []
    for p in items:
        _isdeep = is_deep(p["key"], deepk)
        if deep == "yes" and not _isdeep:
            continue
        if deep == "no" and _isdeep:
            continue
        # F38-B：按当前学科取分级（快路径 compute=False，只用已预热 memo；未预热则回退旧 journal_tier）
        # 手动改档/法源报告规则在 grade_paper 里优先命中（不走 memo，即改即显）。
        g = GS.grade_paper(p, compute=False)
        out.append({
            "key": p["key"], "title": p.get("title", ""), "author": p.get("author", ""),
            "year": p.get("year", ""), "journal": p.get("journal", ""),
            "journal_tier": p.get("journal_tier", ""), "tier_rank": p.get("tier_rank", 6),
            "weight_tier": (g["cn"] if g else ""),           # 学科感知中文档名（未预热时空→前端兜旧）
            "weight_rank": (g["rank"] if g else 6),
            "journal_weight": (g["weight"] if g else None),
            "weight_needs_review": (g["needs_review"] if g else False),
            "weight_src": (g.get("src") if g else None),     # manual=手动改档 / rule=法源报告规则
            "official_pages": p.get("official_pages", ""), "has_pdf": p.get("has_pdf", False),
            "collections": p.get("collections", []),
            "needs_review": bool(p.get("needs_review", False)),   # folder 模式：AI 抽的题录待核对
            "no_text": T.safe_name(p["key"]) in notextk,          # C1/A2：扫描件（有 PDF 但无可抽文本，需 OCR，不可深索）
            "ingested_at": p.get("ingested_at", ""),               # 供「最新入库」排序
            "score": _rec_score(p, g), "deep": _isdeep,
        })
    if sort == "recommend":
        out.sort(key=lambda x: -x["score"])
    elif sort == "year":
        def _y(x):
            try: return int(x["year"] or 0)
            except Exception: return 0
        out.sort(key=lambda x: -_y(x))
    elif sort == "ingested":
        out.sort(key=lambda x: x.get("ingested_at", ""), reverse=True)   # 最新入库优先
    # W1：分页出口——大库此前只能看到前 limit 篇，其余永远翻不到；total=过滤后总数供前端算页数
    off = max(0, int(offset or 0))
    return {"papers": out[off:off + limit], "total": len(out),
            "collection": collection, "topic": topic, "category": category, "deep": deep, "sort": sort}

# ── 单篇手动改档（法源权重改造 2026-07-12）──────────────────
class TierOverrideQ(BaseModel):
    key: str
    tier: Optional[str] = None      # "T1"/"T1b"/"T2"/"T3"/"T4"/"T5"；None/空 = 恢复自动

@app.post("/paper/tier")
def set_paper_tier(q: TierOverrideQ):
    """手动提高/降低某篇的档位（存 data/tier_overrides.json，优先级最高，即改即生效）。"""
    import source_rules as SR
    import grading_svc as GS
    tier = (q.tier or "").strip() or None
    if tier and tier not in SR.TIER_W:
        return JSONResponse({"detail": f"非法档位 {tier}（可选 T1/T1b/T2/T3/T4/T5 或留空恢复自动）"},
                            status_code=400)
    p = _load_papers().get(q.key)
    if not p:
        return JSONResponse({"detail": f"未找到文献 {q.key}"}, status_code=404)
    try:
        SR.set_override(q.key, tier)
    except Exception as e:
        log_error("paper tier override", repr(e))
        return JSONResponse({"detail": f"写入改档失败：{e}"}, status_code=500)
    g = GS.grade_paper(p)            # 改完立刻算出生效档（含回落到规则/期刊分级的情形）
    return {"ok": True, "key": q.key, "override": tier,
            "effective": g or {"tier": None, "cn": p.get("journal_tier", "未知"),
                               "weight": None, "rank": p.get("tier_rank", 6),
                               "needs_review": False}}

# ── 打开原文 PDF（C2/D4：页级引注可回到原文核对）─────────────
class OpenPdfQ(BaseModel):
    key: str

@app.post("/open_pdf")
def open_pdf(q: OpenPdfQ):
    """用系统默认阅读器打开某篇的原文 PDF（供深索结果卡「📄 打开原文」）。
       pdf_path 取自 papers.jsonl（zotero/folder 两种源建库时都已落该字段）。"""
    p = _load_papers().get(q.key)
    if not p:
        return JSONResponse({"ok": False, "msg": "无此文献"}, status_code=404)
    path = (p.get("pdf_path") or "").strip()
    if not path or not Path(path).exists():
        return JSONResponse({"ok": False, "msg": "未找到原文 PDF 文件（可能仅题录、或文件已移动/未随库）"}, status_code=200)
    try:
        if sys.platform == "win32":
            os.startfile(path)  # noqa
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
        return {"ok": True, "path": path}
    except Exception as e:
        log_error("open_pdf", repr(e))
        return JSONResponse({"ok": False, "msg": f"打开失败：{e}"}, status_code=200)

# ══════════════════════════════════════════════════════════════════
#  读取论文全文（agent 的 ingest 地基）
#  此前 agent 只能拿到 220 字检索片段；逐页正文一直躺在 config.EXTRACTED 里，
#  却只有 _extracted_excerpt 内部自用、没有任何出口。gist 的 ingest 主干
#  "the LLM reads it, extracts the key information" 因此走不通。
# ══════════════════════════════════════════════════════════════════
@app.get("/source/{key}")
def read_source(key: str, from_page: int = 1, to_page: int = 0, max_chars: int = 20000):
    """按 PDF 顺序页返回该篇提取正文（每页附印刷页码）。

    诚实降级：读不到时明确区分「无此篇 / 只有题录无 PDF / 尚未深索 / 扫描件无文字」，
    各给 reason + 可执行的 detail——绝不返回空串让 agent 以为这篇没内容。"""
    import textutil as T
    papers = _load_papers()
    p = papers.get(key)
    if not p:
        return JSONResponse({"ok": False, "reason": "not_found",
                             "detail": f"知识库中没有 key={key} 的文献。可先用 /papers 或 list_sources 枚举。"},
                            status_code=404)
    title = p.get("title", "")
    if not p.get("has_pdf"):
        return {"ok": False, "reason": "no_pdf", "key": key, "title": title,
                "detail": "该篇只有题录、没有 PDF 原文，读不到全文。"}

    stem = p.get("stem") or T.safe_name(key)
    f = C.EXTRACTED / f"{stem}.json"
    if not f.exists():
        return {"ok": False, "reason": "not_deep_indexed", "key": key, "title": title,
                "detail": "该篇尚未深索，没有可读全文。请在应用「浏览」页勾选它深索，"
                          "或调用 localkb_build(stage='deep')。"}
    try:
        rec = json.loads(f.read_text(encoding="utf-8"))
    except Exception as e:
        log_error("source/read", repr(e))
        return JSONResponse({"ok": False, "reason": "unreadable", "key": key,
                             "detail": f"提取文件损坏：{e}"}, status_code=500)

    pages = rec.get("pages") or []
    if not any((pg.get("text") or "").strip() for pg in pages):
        return {"ok": False, "reason": "scanned_no_text", "key": key, "title": title,
                "detail": "该篇是扫描件（图片型 PDF），抽不到文字，需要先 OCR。"}

    lo = max(1, int(from_page or 1))
    hi = int(to_page) if to_page else len(pages)
    try:
        import page_map as PM
    except Exception:
        PM = None

    out, used, truncated, next_page = [], 0, False, None
    budget = max(500, int(max_chars or 20000))
    for pg in pages:
        n = int(pg.get("page", 0) or 0)
        if n < lo or n > hi:
            continue
        txt = (pg.get("text") or "").strip()
        if not txt:
            continue
        if used + len(txt) > budget and out:      # 至少给一页，别因预算太小返回空
            truncated, next_page = True, n
            break
        used += len(txt)
        pr = PM.printed(key, n) if PM else {}
        out.append({"pdf_page": n, "printed_page": (pr or {}).get("display", ""), "text": txt})

    return {"ok": True, "key": key, "title": title,
            "author": p.get("author", ""), "year": p.get("year", ""), "journal": p.get("journal", ""),
            "n_pages_total": len(pages), "returned_pages": len(out),
            "chars": used, "truncated": truncated, "next_page": next_page,
            "pages": out}

# ── 三档索引触发 ──────────────────────────────────────────
@app.post("/index/light")
def index_light_ep():
    # BF9：即时索引此前完全绕开构建锁——与 build 子进程并发重写 papers.jsonl/bm25 会写坏索引。
    # 与 /index/deep(scope=all) 同款约定：忙时返回 {ok:false,busy:true}，前端已能处理 ok:false。
    with _BUILD_LOCK:
        if BUILD["running"]:
            return {"ok": False, "busy": True, "msg": "已有构建任务在跑，请稍后再试"}
        BUILD["running"] = True; BUILD["stage"] = "light"
        BUILD["started"] = time.time(); BUILD["rc"] = None; BUILD["cancelled"] = False
        BUILD["log"] = ["[light] 即时索引启动…"]   # 不留上次构建的日志尾巴
    try:
        import importlib, index_light as IL
        importlib.reload(IL)
        stats = IL.main()
        try:
            import build_categories as BC
            importlib.reload(BC); BC.main()          # 只读 zotero.sqlite，约 1s；失败不影响词法索引
        except Exception as e:
            log_error("index/light categories", repr(e), traceback.format_exc())
        # BF9：仿 _run_build 的重载窗口——先置未就绪，防 tbl 新旧错位时检索拿到错误的空命中
        R.STATE["ready"] = False
        R.load_all()
        return {"ok": True, **stats["coverage"]}
    except Exception as e:
        log_error("index/light", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "msg": str(e)}, status_code=500)
    finally:
        with _BUILD_LOCK:
            BUILD["running"] = False
            BUILD["stage"] = ""

def _child_env():
    """任务五：子进程统一 env——强制 UTF-8 输出（PYTHONIOENCODING）+ 开 UTF-8 模式（PYTHONUTF8=1），
       让 build 子进程稳定输出 UTF-8，server 端按 utf-8 解码不再乱码。"""
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONUTF8"] = "1"
    return env

def _run_build(stage, extra=None, on_done=None):
    # B1：原子「判 running + 置 True」并在调用线程内同步置位（不再等子线程 start() 后才置），
    # 多路触发时后来者立即拿到 running=True → return False，不会并发跑两个子进程。
    with _BUILD_LOCK:
        if BUILD["running"]:
            return False
        BUILD["running"] = True; BUILD["stage"] = stage; BUILD["started"] = time.time()
        BUILD["log"] = [f"[{stage}] 启动…"]; BUILD["rc"] = None
        BUILD["proc"] = None; BUILD["cancelled"] = False
        # BF35：整库深索(scope=all)标记——extra 形如 ["--scope","all"]；队列批次是 "keys:..."，不算 bulk
        BUILD["bulk"] = bool(stage == "deep" and "all" in (extra or []))
    def run():
        rc = None
        try:
            env = _child_env()   # 任务五：稳定 UTF-8 输出，避免 build 日志乱码
            cmd = [sys.executable, str(C.APP / "build_all.py"), "--stage", stage] + (extra or [])
            p = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env)
            BUILD["proc"] = p            # 留句柄给 /build/cancel
            for raw in p.stdout:
                BUILD["log"].append(raw.decode("utf-8", errors="replace").rstrip())
                BUILD["log"] = BUILD["log"][-300:]
            # A3：取子进程 returncode（此前 p.wait() 取到却从不检查）。非 0 = 软失败（余额0/PDF异常）。
            rc = p.wait()
            if BUILD.get("cancelled"):
                BUILD["log"].append("[build] 已被用户取消。已完成的部分已保存，可随时继续。")
            elif rc != 0:
                BUILD["log"].append(f"[build] 子进程以非 0 退出（returncode={rc}）——本批未成功完成。")
        except Exception as e:
            log_error(f"build {stage}", repr(e), traceback.format_exc())
            BUILD["log"].append("[build] 异常：" + str(e))
            rc = -1
        finally:
            BUILD["rc"] = rc
            BUILD["proc"] = None
            # B2：重载期非原子——大库重载数秒内 tbl 已换新表而 records 仍旧，search_full 会拿空结果。
            # 重载全程保持 running=True（锁未释放前不接受新 build）+ 短暂 ready=False（检索先返回未就绪，
            # 而非错误的空命中）；load_all 成功后自会把 ready 置回 True。
            try:
                R.STATE["ready"] = False
                R.load_all(); BUILD["log"].append("[build] 索引已重载。")
            except Exception as e:
                BUILD["log"].append("[build] 重载失败：" + str(e))
            with _BUILD_LOCK:
                BUILD["running"] = False
                BUILD["stage"] = ""          # F4：构建结束复位 stage，避免残留 "deep" 让顶栏误报「深索中」
                BUILD["bulk"] = False        # BF35：构建结束复位整库深索标记
            # A3/A4：把 returncode 贯通给回调（失败时上层决定退回队列/不推进 sig）。
            if on_done:
                try:
                    on_done(rc)
                except Exception as e:
                    log_error("deep queue on_done", repr(e))
            elif not BUILD.get("cancelled"):
                # 被取消时不再排空队列——否则子进程一死就立刻起新批，用户会以为取消没生效
                try:
                    _drain_deep_queue()
                except Exception as e:
                    log_error("deep queue drain", repr(e))
    threading.Thread(target=run, daemon=True).start()
    return True

# ── 更新知识库（增量 all 档：读 Zotero 新增题录 + 补语义层，不动深索）──
# 前端「更新知识库」按钮 POST /build（此前无此路由、404 被静默吞→按钮假死）。
# _run_build("all") 走 build_all.py 的增量管线（已入库跳过），深索仍由用户手动触发。
@app.post("/build")
def build_ep():
    return {"ok": _run_build("all")}

def _kill_tree(p):
    """BF7：Windows 下 p.terminate() 只杀 build_all 本体，它派生的孙进程（嵌入/抽取 worker）
       会变孤儿继续跑、继续烧 API 额度——taskkill /T /F 终止整棵进程树；失败兜底 terminate 并记日志。"""
    if sys.platform == "win32":
        try:
            r = subprocess.run(["taskkill", "/PID", str(p.pid), "/T", "/F"],
                               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=15)
            if r.returncode == 0:
                return
            log_error("build/cancel taskkill", f"taskkill rc={r.returncode}，兜底 terminate")
        except Exception as e:
            log_error("build/cancel taskkill", repr(e))
    p.terminate()

@app.post("/build/cancel")
def build_cancel():
    """取消正在跑的建库/深索子进程。整库深索(scope=all)此前一旦开始就停不下来，
       可能空跑数小时并烧掉 API 额度。深索是增量的：已完成的篇已落盘，随时可继续。
       同时暂停深索队列——否则子进程一死，_drain_deep_queue 立刻起新批。"""
    p = BUILD.get("proc")
    if not BUILD["running"] or p is None:
        return JSONResponse({"ok": False, "detail": "当前没有正在运行的任务"}, status_code=409)
    BUILD["cancelled"] = True
    try:
        with _Q_LOCK:
            QUEUE["paused"] = True
            _q_persist()
    except Exception as e:
        log_error("build/cancel pause-queue", repr(e))
    try:
        _kill_tree(p)   # BF7：杀整棵进程树，孙进程不再漏网
    except Exception as e:
        log_error("build/cancel", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": f"终止失败：{e}"}, status_code=500)
    BUILD["log"].append("[build] 收到取消请求，正在停止…")
    return {"ok": True, "stage": BUILD.get("stage"), "queue_paused": True}

@app.get("/build/status")
def build_status():
    return {"running": BUILD["running"], "stage": BUILD["stage"],
            "log": BUILD["log"][-300:], "started": BUILD["started"],
            # cancellable：有活着的子进程句柄才能取消（供前端决定是否显示「取消」按钮）
            "cancellable": bool(BUILD["running"] and BUILD.get("proc") is not None),
            "cancelled": bool(BUILD.get("cancelled"))}

# ── 文件夹模式：建库（后台，含 N 次 LLM 抽题录）+ 拖入入库 ──
@app.post("/index/folder_build")
def index_folder_build():
    import settings as S
    if S.source() != "folder":
        return JSONResponse({"ok": False, "msg": "当前非文件夹模式"}, status_code=400)
    try:
        import folder_meta as FM
        ready = FM.available()
    except Exception:
        ready = False
    if not ready:
        return JSONResponse({"ok": False, "need_key": True,
                             "msg": "未配置 API Key（用于自动抽取题录）；也可先建粗库（文件名题名），配好 Key 后再更新"},
                            status_code=200)  # 200 + need_key：前端可选择继续无 key 建库
    return {"ok": _run_build("folder")}

@app.post("/index/folder_build_nokey")
def index_folder_build_nokey():
    """无 key 也建库：folder_ingest 在无 key 时正常退出（退回文件名题名），仍能建词法层。"""
    import settings as S
    if S.source() != "folder":
        return JSONResponse({"ok": False, "msg": "当前非文件夹模式"}, status_code=400)
    return {"ok": _run_build("folder")}

def _dedupe_name(dst):
    """同名不同容→加序号后缀，避免覆盖。"""
    if not dst.exists():
        return dst
    stem, suf = dst.stem, dst.suffix
    i = 2
    while True:
        cand = dst.with_name(f"{stem} ({i}){suf}")
        if not cand.exists():
            return cand
        i += 1

class IngestFile(BaseModel):
    name: str
    content_b64: str                              # 文件内容 base64（避免依赖 python-multipart）

class IngestQ(BaseModel):
    files: List[IngestFile] = []

@app.post("/ingest/files")
def ingest_files(q: IngestQ):
    """拖入/选择的 PDF（base64）复制进受管文件夹 → 去重 → 后台 folder build（抽题录+索引）。
       用 JSON base64 而非 multipart，避免分发版缺 python-multipart。"""
    import settings as S, hashlib as _hl, base64 as _b64
    if S.source() != "folder":
        return JSONResponse({"ok": False, "msg": "仅文件夹模式支持拖入入库"}, status_code=400)
    folder = S.folder_dir()
    if not folder:
        return JSONResponse({"ok": False, "msg": "未配置受管文件夹"}, status_code=400)
    if BUILD["running"]:
        return JSONResponse({"ok": False, "msg": "正在建库/入库中，请稍后再拖入", "building": True}, status_code=409)
    fp = Path(folder)
    # R2：先按文件大小建索引，只有大小撞车的候选才按需算 sha1（并缓存），
    # 避免每拖一篇就把受管文件夹里每个 PDF 全量读盘算 sha1（大库会卡几十秒~几分钟）。
    size_index = {}          # size -> [paths]
    for p in fp.rglob("*.pdf"):
        try:
            size_index.setdefault(p.stat().st_size, []).append(p)
        except Exception:
            pass
    _sha1_cache = {}         # path -> sha1（仅对大小相同的候选按需算一次）
    def _is_dup(data):
        cands = size_index.get(len(data))
        if not cands:
            return False
        h = _hl.sha1(data).hexdigest()
        for p in cands:
            hp = _sha1_cache.get(p)
            if hp is None:
                try:
                    hp = _hl.sha1(p.read_bytes()).hexdigest()
                except Exception:
                    hp = ""
                _sha1_cache[p] = hp
            if hp and hp == h:
                return True
        return False
    added, skipped, failed = [], [], []
    for f in (q.files or []):
        name = f.name or "untitled.pdf"
        if not name.lower().endswith(".pdf"):
            failed.append({"name": name, "reason": "非 PDF"}); continue
        try:
            data = _b64.b64decode((f.content_b64 or "").split(",")[-1])
            if _is_dup(data):
                skipped.append(name); continue
            dst = _dedupe_name(fp / Path(name).name)
            dst.write_bytes(data)
            # 把新写入的文件计入索引/缓存，供同批后续文件去重（原逻辑靠 existing.add）
            size_index.setdefault(len(data), []).append(dst)
            _sha1_cache[dst] = _hl.sha1(data).hexdigest()
            added.append(dst.name)
        except Exception as e:
            failed.append({"name": name, "reason": str(e)})
    try:
        import folder_meta as FM
        ready = FM.available()
    except Exception:
        ready = False
    building = False
    if added:
        building = _run_build("folder")   # 后台增量：只处理新 key（幂等）；无 key 时 folder_ingest 退文件名
    return {"ok": True, "added": len(added), "added_names": added,
            "skipped": len(skipped), "failed": failed,
            "building": building, "need_key": (not ready)}

@app.post("/index/semantic")
def index_semantic_ep():
    return {"ok": _run_build("semantic")}

class DeepQ(BaseModel):
    scope: str = "all"

@app.post("/index/deep")
def index_deep_ep(q: DeepQ):
    # C7/A1：手动深索不再静默丢弃。
    #  - scope="keys:k1,k2..."（勾选若干篇深索）→ 走持久队列 enqueue_deep：撞锁自动排队、
    #    崩溃可回灌续跑，绝不因忙而蒸发。返回真正入队数 queued。
    #  - scope="all"（整库深索）→ 忙时返回 {ok:false,busy:true} 让前端提示「已有任务在跑」，
    #    而非假装成功。
    scope = (q.scope or "all").strip()
    if scope.startswith("keys:"):
        raw = scope[len("keys:"):]
        keys = [k for k in (raw.split(",")) if k.strip()]
        queued = enqueue_deep(keys)
        return {"ok": True, "queued": len(queued), "scope": scope}
    started = _run_build("deep", ["--scope", scope])
    if not started:
        return {"ok": False, "busy": True, "scope": scope}
    return {"ok": True, "queued": scope, "scope": scope}

# ── #7：Agent 驱动深索（切块→Agent写摘要→带摘要嵌入，一趟完成，阻塞式逐批）──
class DeepAgentSummary(BaseModel):
    key: str
    summary: str = ""

class DeepAgentQ(BaseModel):
    summaries: Optional[List[DeepAgentSummary]] = None   # 带上一批 Agent 写的摘要则先写盘+嵌入
    batch: int = 15

def _run_stage_blocking(stage, extra=None):
    """阻塞跑一段 build_all（子进程 subprocess.run，捕获 stdout 不污染 HTTP 响应）。返回 returncode。"""
    env = _child_env()
    cmd = [sys.executable, str(C.APP / "build_all.py"), "--stage", stage] + (extra or [])
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env)
    try:
        out = (p.stdout or b"").decode("utf-8", errors="replace")
        BUILD["log"].append(f"[deep_agent:{stage}] rc={p.returncode}")
        for ln in out.splitlines()[-20:]:
            BUILD["log"].append(ln)
        BUILD["log"] = BUILD["log"][-300:]
    except Exception:
        pass
    return p.returncode

def _extracted_excerpt(stem, limit=1800):
    """取该篇提取正文前 ~limit 字，供 Agent 写摘要。无提取文本（未切块/扫描件）→ ""。"""
    f = C.EXTRACTED / f"{stem}.json"
    if not f.exists():
        return ""
    try:
        rec = json.loads(f.read_text(encoding="utf-8"))
    except Exception:
        return ""
    txt = "\n".join((pg.get("text") or "") for pg in rec.get("pages", []))
    return txt[:limit]

def _deep_agent_run(q: DeepAgentQ):
    import sac as SAC
    wrote = 0
    # ① 带 summaries：写进 summaries.json，再阻塞跑 deep_embed 把上一批带摘要嵌入（标记已深索）
    if q.summaries:
        wrote = SAC.write_summaries([{"key": s.key, "summary": s.summary} for s in q.summaries])
        # BF16：接住子阶段退出码——此前 rc≠0（余额不足/网络断）也照样返回 ok:true，
        # agent 会向用户误报「已嵌入入库」，实际一篇都没进。
        rc = _run_stage_blocking("deep_embed")
        if rc != 0:
            return {"ok": False, "stage": "deep_embed",
                    "error": f"深索子阶段失败(rc={rc})：可能是 API 余额不足或网络问题，请稍后重试"}
        try:                              # 嵌入后重载索引，让新深索的篇立刻可检索
            R.STATE["ready"] = False
            R.load_all()
        except Exception as e:
            log_error("deep_agent reload", repr(e))
    # ② 选下一批 pending-deep（有PDF、未深索、非扫描件），阻塞跑 deep_prepare(extract+chunk)
    papers = _load_papers(); deepk = _deep_keys(); notext = _deep_no_text_keys()
    cand = [p for p in papers.values()
            if p.get("has_pdf") and not is_deep(p["key"], deepk)
            and T.safe_name(p["key"]) not in notext]
    batch_n = max(1, int(q.batch or 15))
    batch = cand[:batch_n]
    to_summarize = []
    if batch:
        keys = [p["key"] for p in batch]
        # BF16：deep_prepare 失败（抽取/切块子进程崩）时 excerpt 都是空的，agent 拿去写摘要毫无意义
        rc = _run_stage_blocking("deep_prepare", ["--scope", "keys:" + ",".join(keys)])
        if rc != 0:
            return {"ok": False, "stage": "deep_prepare",
                    "error": f"深索子阶段失败(rc={rc})：可能是 API 余额不足或网络问题，请稍后重试"}
        for p in batch:
            to_summarize.append({"key": p["key"], "title": p.get("title", ""),
                                 "excerpt": _extracted_excerpt(p["stem"])})
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    # ③ 汇总返回。finished=没有更多待处理（末批：summaries 已提交且无新篇也能收尾）
    return {"ok": True, "wrote": wrote, "to_summarize": to_summarize,
            "done": len(_deep_keys()), "with_pdf": manifest.get("with_pdf", 0),
            "remaining": max(0, len(cand) - len(batch)),
            "finished": (len(batch) == 0)}

@app.post("/index/deep_agent")
def index_deep_agent(q: DeepAgentQ):
    """#7：Agent 驱动深索的单批处理（阻塞，可跑 30-60s）。尊重构建锁：有其它构建在跑→busy。
       用法见 MCP 工具 deep_index：首次不带 summaries→返回 to_summarize；带 summaries 再调→
       嵌入上一批+返回下一批；循环至 finished=true。"""
    with _BUILD_LOCK:
        if BUILD["running"]:
            return {"ok": False, "busy": True}
        BUILD["running"] = True; BUILD["stage"] = "deep_agent"
        BUILD["started"] = time.time(); BUILD["rc"] = None
        BUILD["cancelled"] = False       # BF14：每次构建开始都复位取消标记，避免上次取消残留误报
    try:
        return _deep_agent_run(q)
    except Exception as e:
        log_error("index/deep_agent", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=200)
    finally:
        with _BUILD_LOCK:
            BUILD["running"] = False
            BUILD["stage"] = ""          # F4：deep_agent 结束同样复位 stage

# ── 检索 ──────────────────────────────────────────────────
class SearchQ(BaseModel):
    query: str
    topk: int = C.RERANK_TOPK
    sort: Optional[str] = None
    min_weight: float = 0.0                       # 引用权重下限过滤（0=不过滤）
    category: Optional[str] = None                # F11：限定检索范围到某个分类（kbc_/topic:/zotero:）

@app.post("/search")
def search(q: SearchQ):
    if not R.STATE.get("ready"):
        # UX8：两态文案——建过库只是重建/重载窗口（稍等即可），从未建库才需要走向导；
        # 此前一律"先建立即时索引"，害老用户在重载的几秒里被误导去重建。
        msg = ("索引正在重建或加载，请稍候几秒再搜" if C.INDEX_MANIFEST.exists()
               else "还没建立索引——请到 设置 → 重新查看引导 完成首次建库")
        return JSONResponse({"error": msg, "ready": False}, status_code=503)
    t0 = time.time()
    keys = _resolve_category_keys(q.category)
    try:
        res = R.search(q.query, q.topk, q.sort, q.min_weight, keys=keys)
    except Exception as e:
        # BF36：API 模式下嵌入/重排后端挂了（余额0/断网）此前抛成裸 500「服务器内部错误」，
        # 用户不知道该去哪修——detail 给人话，前端 jpost 已会读 detail 展示。
        log_error("search", repr(e), traceback.format_exc())
        raise HTTPException(status_code=500,
                            detail=f"{e}（→ 请到 设置 检查检索引擎 API Key 或余额）")
    return {"query": q.query, "mode": R.STATE.get("mode"), "category": q.category,
            "took_ms": round((time.time() - t0) * 1000), "results": res}

# ── RAG 对话 ──────────────────────────────────────────────
class ChatQ(BaseModel):
    query: str
    history: List[dict] = []
    provider: str = "deepseek"
    base_url: str = ""
    api_key: str = ""
    model: str = ""
    topk: int = 6
    sort: Optional[str] = None
    category: Optional[str] = None                # F11：限定对话检索范围到某个分类

SYS_TMPL = (
    "你是严谨的法学文献研究助手。请**只依据下面提供的文献片段**回答用户问题，"
    "每个论点后用 [编号] 标注来源；若不足以回答请如实说明不要编造。回答用中文。\n\n=== 文献片段 ===\n{ctx}"
)

@app.post("/chat")
def chat(q: ChatQ):
    keys = _resolve_category_keys(q.category)
    # C6/A5：检索后端(嵌入/重排 API)挂了不能被误报成「模型没返回内容，请检查对话模型/Key」。
    # 把同步的 R.search 包 try：失败则 hits=[]、记下真因，稍后在 SSE 里先 yield error 再照常作答。
    search_err = None
    hits = []
    if R.STATE.get("ready"):
        try:
            hits = R.search(q.query, q.topk, q.sort, keys=keys)
        except Exception as e:
            log_error("chat search", repr(e))
            search_err = "检索后端暂时不可用（可能余额/网络/超时），本次未附引用"
    ctx = "\n\n".join(
        f"[{i+1}] {h.get('citation','')}\n{(h.get('context') or h.get('text') or '')[:1200]}"
        for i, h in enumerate(hits)
    ) or "（暂无检索结果）"
    messages = ([{"role": "system", "content": SYS_TMPL.format(ctx=ctx)}]
                + list(q.history) + [{"role": "user", "content": q.query}])
    base, model = L.resolve(q.provider, q.base_url, q.model)
    api_key = q.api_key
    # 对话选 SiliconFlow 但没单独填 key 时，自动复用检索引擎已配的 SiliconFlow key（一个 key 通吃，免费模型开箱即用）
    if not api_key and "siliconflow" in (base or "").lower():
        try:
            import settings as S
            api_key = (S.api_conf() or {}).get("key", "") or api_key
        except Exception:
            pass

    def gen():
        yield "data: " + json.dumps({"sources": hits}, ensure_ascii=False) + "\n\n"
        if search_err:      # C6：先把检索失败的真因抛给前端 j.error 分支，再照常让模型作答
            yield "data: " + json.dumps({"error": search_err}, ensure_ascii=False) + "\n\n"
        try:
            for delta in L.chat_stream(messages, base, api_key, model):
                yield "data: " + json.dumps({"delta": delta}, ensure_ascii=False) + "\n\n"
        except Exception as e:
            log_error("chat", repr(e))
            yield "data: " + json.dumps({"error": str(e)}, ensure_ascii=False) + "\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")

# ── 静态 UI ───────────────────────────────────────────────
WEB = C.APP / "web"

@app.get("/")
def index():
    # 任务六：显式 charset=utf-8，避免 WebView2 按系统 GBK 解码整页
    return FileResponse(str(WEB / "index.html"), media_type="text/html; charset=utf-8")

app.mount("/static", StaticFiles(directory=str(WEB)), name="static")

if __name__ == "__main__":
    print(f"本地知识库服务启动：http://{C.DAEMON_HOST}:{C.DAEMON_PORT}", flush=True)
    uvicorn.run(app, host=C.DAEMON_HOST, port=C.DAEMON_PORT, log_level="warning")
