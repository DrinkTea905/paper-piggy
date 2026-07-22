# -*- coding: utf-8 -*-
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

APP = Path(__file__).resolve().parents[1]
if str(APP) not in sys.path:
    sys.path.insert(0, str(APP))

import document_formats as DF
import extract as E
import folder_source as FS
import cite_format as CF
import chunk as CH


class DocumentFormatTests(unittest.TestCase):
    def test_primary_attachment_priority_and_pdf_compatibility(self):
        rec = DF.apply_attachment_fields({}, [
            {"path": "book.txt"},
            {"path": "book.docx"},
            {"path": "book.epub"},
            {"path": "book.pdf"},
            {"path": "book.md"},
        ])
        self.assertEqual(rec["fulltext_format"], "pdf")
        self.assertEqual(rec["fulltext_path"], "book.pdf")
        self.assertEqual([x["format"] for x in rec["fulltext_attachments"]],
                         ["pdf", "epub", "docx", "markdown", "txt"])
        self.assertTrue(rec["has_pdf"])
        self.assertEqual(rec["pdf_path"], "book.pdf")

    def test_without_pdf_epub_wins_and_html_is_rejected(self):
        rec = DF.apply_attachment_fields({}, [
            {"path": "a.txt"}, {"path": "a.docx"}, {"path": "a.epub"},
        ])
        self.assertEqual(rec["fulltext_format"], "epub")
        self.assertFalse(rec["has_pdf"])
        self.assertFalse(DF.supported_file("snapshot.html"))

    def test_non_pdf_citation_uses_locator_not_fake_page(self):
        text = CF.compact({
            "key": "K", "author": "作者", "title": "电子书", "journal": "",
            "fulltext_format": "epub", "locator": "第二章 制度沿革", "page": 2,
        })
        self.assertIn("第二章 制度沿革", text)
        self.assertNotIn("第2页", text)

    def test_old_pdf_record_normalizes_without_rebuild(self):
        rec = DF.normalize_record({"has_pdf": True, "pdf_path": "old.pdf"})
        self.assertTrue(rec["has_fulltext"])
        self.assertEqual(rec["fulltext_format"], "pdf")
        self.assertEqual(rec["fulltext_path"], "old.pdf")

    def test_markdown_and_txt_locators(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            md = root / "note.md"
            md.write_text("# 总论\n" + "第一段正文用于验证通用全文切块。" * 8
                          + "\n## 分论\n" + "第二段正文用于验证标题定位。" * 8, encoding="utf-8")
            got = E._extract_source_document(str(md), source_format="markdown")
            self.assertEqual(got["document_format"], "markdown")
            self.assertEqual([x["heading"] for x in got["pages"]], ["总论", "分论"])
            chunks = CH.chunk_doc({"key": "M", "meta": {"title": "笔记", "fulltext_format": "markdown"}, **got})
            self.assertTrue(chunks)
            self.assertFalse(chunks[0]["has_pdf"])
            self.assertEqual(chunks[0]["heading"], "总论")

            txt = root / "note.txt"
            txt.write_text("\n".join(f"第{i}行" for i in range(1, 206)), encoding="utf-8")
            got = E._extract_source_document(str(txt), source_format="txt")
            self.assertEqual(len(got["pages"]), 2)
            self.assertEqual(got["pages"][0]["locator_label"], "第 1–200 行")
            self.assertEqual(got["pages"][1]["locator_label"], "第 201–205 行")

    def test_epub_spine_order_and_chapter_locator(self):
        with tempfile.TemporaryDirectory() as td:
            epub = Path(td) / "book.epub"
            with zipfile.ZipFile(epub, "w") as zf:
                zf.writestr("mimetype", "application/epub+zip")
                zf.writestr("META-INF/container.xml", '''<?xml version="1.0"?>
<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
 <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>''')
                zf.writestr("OEBPS/content.opf", '''<?xml version="1.0"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0">
 <manifest><item id="c1" href="one.xhtml" media-type="application/xhtml+xml"/>
 <item id="c2" href="two.xhtml" media-type="application/xhtml+xml"/></manifest>
 <spine><itemref idref="c2"/><itemref idref="c1"/></spine>
</package>''')
                zf.writestr("OEBPS/one.xhtml", "<html><body><h1>第一章</h1><p>甲内容</p></body></html>")
                zf.writestr("OEBPS/two.xhtml", "<html><body><h1>第二章</h1><p>乙内容</p></body></html>")
            got = E._extract_source_document(str(epub), source_format="epub")
            self.assertEqual(got["document_format"], "epub")
            self.assertIn("第二章", got["pages"][0]["text"])
            self.assertEqual(got["pages"][0]["heading"], "第二章")
            self.assertIn("第一章", got["pages"][1]["text"])

    def test_docx_headings_and_folder_scan(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx 未安装")
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            docx = root / "paper.docx"
            doc = Document()
            doc.add_heading("绪论", level=1)
            doc.add_paragraph("这是正文。")
            doc.add_heading("结论", level=1)
            doc.add_paragraph("这是结论。")
            doc.save(docx)
            got = E._extract_source_document(str(docx), source_format="docx")
            self.assertEqual([x["heading"] for x in got["pages"]], ["绪论", "结论"])

            (root / "a.pdf").write_bytes(b"not needed")
            (root / "b.epub").write_bytes(b"not needed")
            (root / "c.md").write_text("# c", encoding="utf-8")
            (root / "d.txt").write_text("d", encoding="utf-8")
            (root / "ignored.html").write_text("ignored", encoding="utf-8")
            agent_dir = root / "0_Agent交付物"
            agent_dir.mkdir()
            (agent_dir / "hidden.pdf").write_bytes(b"ignored")
            names = {Path(x).name for x in FS.scan(str(root))}
            self.assertEqual(names, {"paper.docx", "a.pdf", "b.epub", "c.md", "d.txt"})


if __name__ == "__main__":
    unittest.main()
