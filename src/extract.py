# -*- coding: utf-8 -*-
"""
F 档 · 提取：从 papers.jsonl 中【有可读全文】的篇提取定位单元 → data/extracted/<stem>.json。
数据源 = index_light 生成的 papers.jsonl（含 fulltext_path / fulltext_format / collections 等）。
--scope 决定深索范围（支持"选择性深索"）：
    all                 全部有可读全文的
    collection:<path>   某收藏夹（papers 的 collections 含该 path）
    keys:K1,K2,...       指定若干篇（前端勾选/按推荐深索）
断点续跑：已提取的跳过。ThreadPool（避开本机多进程 spawn 坑）。
用法: python extract.py [--scope all] [--workers 4] [--limit N]
"""
import sys, os, json, argparse, time, threading, re, zipfile, posixpath
from html.parser import HTMLParser
from urllib.parse import unquote
import xml.etree.ElementTree as ET
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
sys.path.insert(0, str(Path(__file__).parent))
import config as C
import deep_extract_status as DES
import document_formats as DF

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

META_FIELDS = ("title", "author", "year", "journal", "doi", "langid",
               "official_pages", "itemtype", "journal_tier", "has_pdf",
               "has_fulltext", "fulltext_format")

# PDF 逐页取文本：用 pypdfium2（Google PDFium 的绑定，BSD-3/Apache-2.0）。
#
# 【为什么不是 PyMuPDF/pymupdf4llm】——2026-07 换掉，三个理由，缺一不可：
#   ① 许可证：pymupdf4llm 强制依赖 pymupdf_layout（Polyform Noncommercial，禁止商业使用、
#      非 OSI 开源），PyMuPDF 本身是 AGPL。随开源包分发它们与本项目的 Apache-2.0 冲突。
#   ② 质量：实测 6 篇真实法学 PDF，pymupdf4llm 的 markdown 转换会把**中文标点错序**
#      （句号/引号跑到句首），而三者提取的中文字符数完全相同——markdown 那点结构收益
#      根本不抵这个损失。pypdfium2 的断行最少、标点最准。
#   ③ 速度与稳定性：pymupdf4llm 单篇 9s 量级（pypdfium2 0.1~0.4s），且它 import 时会无条件
#      `import onnxruntime`（其 OCR 模块）——打包版若缺 VC++ 运行库，会让整个 import 抛错，
#      表现为**每篇 PDF 提取都失败、深索静默产出空正文**。换掉后这条隐患一并消失。
#
import pypdfium2 as _pdfium

# ★★ PDFium 不是线程安全的 ★★（2026-07-15 用户实机事故，Windows 事件日志确诊）：
# extract 用 4 线程 ThreadPool 并发调 pypdfium2，多线程同时进 PDFium → 损坏其内部堆，
# 表现为 ① 崩进程（0xc0000374 heap corruption / 0x80000003 断言，故障模块 pdfium.dll）——
#        分发版里 server 拉起的 launcher 进程随之关窗；
#      ② 大量随机「PdfiumError: Failed to load」（同一 PDF 单独打开却好好的），
#        被上层误标成「扫描件·需 OCR」，害得能深索的篇数虚低。
# 修法：所有 pypdfium2 调用**全局串行**。OCR 只接收锁内复制出的单页像素，在锁外运行；
# 这样既不让 PDFium 并发，又不让较慢的 OCR 长时间占住 PDFium。folder_ingest 也走同一入口。
_PDFIUM_LOCK = threading.Lock()
_OCR_LOCK = threading.Lock()       # 初期保守单并发；实测内存/线程安全后再考虑放宽
_OCR_ENGINE = None
_OCR_ENGINE_ERROR = None           # 初始化失败也缓存，避免一本扫描件每页都重复加载失败
OCR_DPI = 300


class OCRUnavailable(RuntimeError):
    pass


def _get_ocr_engine():
    """懒加载本地 RapidOCR；原生有文字的 PDF 完全不会 import 它。"""
    global _OCR_ENGINE, _OCR_ENGINE_ERROR
    if _OCR_ENGINE_ERROR is not None:
        raise OCRUnavailable(_OCR_ENGINE_ERROR)
    if _OCR_ENGINE is None:
        try:
            from rapidocr import RapidOCR
            # rapidocr>=3.9 的默认模型即 ONNX Runtime + PP-OCRv6 small，且模型随
            # wheel 分发；不传自定义模型路径，运行时不会为默认模型联网下载。
            _OCR_ENGINE = RapidOCR(params={"Global.log_level": "critical"})
        except Exception as e:
            _OCR_ENGINE_ERROR = (
                "本地 OCR 组件 RapidOCR 不可用；请重新安装或更新 PaperPiggy"
                "（源码运行请按 requirements.lock 安装依赖）：" + str(e))
            raise OCRUnavailable(_OCR_ENGINE_ERROR) from e
    return _OCR_ENGINE


def _ocr_result_text(result):
    """兼容 RapidOCR v3 对象和旧版 list/tuple 输出，返回 (正文, 平均置信度)。"""
    if result is None:
        return "", None
    texts = getattr(result, "txts", None)
    scores = getattr(result, "scores", None)
    if texts is not None:
        clean = [str(x).strip() for x in texts if str(x).strip()]
        vals = []
        for score in (scores if scores is not None else []):
            try:
                vals.append(float(score))
            except (TypeError, ValueError):
                pass
        return "\n".join(clean), (sum(vals) / len(vals) if vals else None)
    # RapidOCR 1.x/2.x 常见返回：(result_list, elapsed)
    if isinstance(result, tuple) and len(result) == 2 and isinstance(result[0], list):
        result = result[0]
    lines, vals = [], []
    for item in result if isinstance(result, (list, tuple)) else []:
        if isinstance(item, dict):
            text = item.get("text") or item.get("txt") or ""
            score = item.get("score")
        elif isinstance(item, (list, tuple)) and len(item) >= 3:
            text, score = item[1], item[2]
        else:
            continue
        text = str(text).strip()
        if text:
            lines.append(text)
        try:
            if score is not None:
                vals.append(float(score))
        except (TypeError, ValueError):
            pass
    return "\n".join(lines), (sum(vals) / len(vals) if vals else None)


def _ocr_image(image, engine=None):
    """识别一页内存图像；不落临时图片、不联网。"""
    with _OCR_LOCK:
        result = (engine or _get_ocr_engine())(image)
    return _ocr_result_text(result)


def _read_page(doc, index, render_empty=False):
    """所有 PDFium 调用均在全局锁内；返回后图像已复制，可在锁外 OCR。"""
    with _PDFIUM_LOCK:
        page = doc[index]
        try:
            tp = page.get_textpage()
            try:
                text = (tp.get_text_range() or "").strip()
            finally:
                tp.close()
            image = None
            if not text and render_empty:
                bitmap = page.render(scale=OCR_DPI / 72.0)
                try:
                    pixels = bitmap.to_numpy()
                    # PDFium 在 Windows 常给 BGRA；RapidOCR 的 numpy 输入按 OpenCV
                    # 约定接收 BGR。只在当前页内存中去掉 alpha，不落中间图片。
                    if getattr(pixels, "ndim", 0) == 3 and pixels.shape[2] == 4:
                        pixels = pixels[:, :, :3]
                    image = pixels.copy()
                finally:
                    bitmap.close()
            return text, image
        finally:
            page.close()


def _extract_document(pdf, _tries=3, max_pages=None, ocr_mode="off",
                      ocr_engine=None, on_ocr_pending=None):
    """提取整篇并返回 pages + 统计。

    ``ocr_mode='empty_pages'`` 时仅对原生文字为空的页做本地 OCR；``max_pages``
    在打开 PDF 后立刻截断，供 folder 题录只读前两页，避免先 OCR 整本再切片。
    """
    if ocr_mode not in ("off", "empty_pages"):
        raise ValueError(f"未知 ocr_mode：{ocr_mode}")
    last = None
    for attempt in range(_tries):
        doc = None
        try:
            with _PDFIUM_LOCK:
                doc = _pdfium.PdfDocument(pdf)
                total = len(doc)
            limit = total if not max_pages else min(total, max(0, int(max_pages)))
            pages, native_n, ocr_n = [], 0, 0
            confidences, errors = [], []
            pending_sent = False
            for i in range(limit):
                text, image = _read_page(doc, i, render_empty=(ocr_mode == "empty_pages"))
                if text:
                    native_n += 1
                    pages.append({"page": i + 1, "text": text,
                                  "source": "native", "confidence": 1.0})
                    continue
                if ocr_mode != "empty_pages":
                    continue
                if not pending_sent and on_ocr_pending:
                    try:
                        on_ocr_pending({"total_pages": limit, "native_pages": native_n})
                    except Exception:
                        # 状态 sidecar 写失败不能把一个可读 PDF 误判为 invalid_pdf；最终产物
                        # 仍照常提取，调用方稍后会再写终态。
                        pass
                    pending_sent = True
                try:
                    text, confidence = _ocr_image(image, engine=ocr_engine)
                except Exception as e:
                    errors.append(f"p{i + 1}: {type(e).__name__}: {e}")
                    continue
                if text.strip():
                    ocr_n += 1
                    if confidence is not None:
                        confidences.append(confidence)
                    pages.append({"page": i + 1, "text": text.strip(),
                                  "source": "ocr", "confidence": confidence})
            pages.sort(key=lambda x: x["page"])
            return {
                "pages": pages,
                "total_pages": limit,
                "native_pages": native_n,
                "ocr_pages": ocr_n,
                "empty_pages": max(0, limit - native_n - ocr_n),
                "ocr_confidence": (sum(confidences) / len(confidences)
                                   if confidences else None),
                "ocr_errors": errors,
            }
        except Exception as e:
            last = e
            if attempt < _tries - 1:
                time.sleep(0.3 * (attempt + 1))
        finally:
            if doc is not None:
                try:
                    with _PDFIUM_LOCK:
                        doc.close()
                except Exception:
                    pass
    raise last


def _extract_pages(pdf, _tries=3, max_pages=None, ocr_mode="off", ocr_engine=None):
    """兼容旧调用的 pages-only 包装；默认仍只做原生提取。"""
    return _extract_document(pdf, _tries=_tries, max_pages=max_pages,
                             ocr_mode=ocr_mode, ocr_engine=ocr_engine)["pages"]


_ZIP_TOTAL_LIMIT = 250 * 1024 * 1024
_ZIP_MEMBER_LIMIT = 40 * 1024 * 1024


def _check_zip_limits(zf):
    """拒绝明显的 zip bomb；EPUB/DOCX 均是不可信的 ZIP 容器。"""
    total = 0
    for info in zf.infolist():
        total += int(info.file_size or 0)
        if info.file_size > _ZIP_MEMBER_LIMIT or total > _ZIP_TOTAL_LIMIT:
            raise ValueError("压缩文档展开后过大，已拒绝读取")


class _HTMLText(HTMLParser):
    """仅用于 EPUB 容器内部的 XHTML；独立 HTML 文件仍不属于支持格式。"""
    _BREAK = {"p", "div", "section", "article", "li", "br", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}
    _IGNORE = {"script", "style", "svg", "math", "noscript"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts, self.ignore_depth, self.heading_depth = [], 0, 0
        self.heading_parts = []

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in self._IGNORE:
            self.ignore_depth += 1
        if tag in self._BREAK and self.parts and self.parts[-1] != "\n":
            self.parts.append("\n")
        if tag in {"h1", "h2", "h3"}:
            self.heading_depth += 1

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self._IGNORE and self.ignore_depth:
            self.ignore_depth -= 1
        if tag in self._BREAK and self.parts and self.parts[-1] != "\n":
            self.parts.append("\n")
        if tag in {"h1", "h2", "h3"} and self.heading_depth:
            self.heading_depth -= 1

    def handle_data(self, data):
        if self.ignore_depth:
            return
        s = re.sub(r"[\t\r\f\v ]+", " ", data or "")
        if not s.strip():
            return
        self.parts.append(s)
        if self.heading_depth and len("".join(self.heading_parts)) < 160:
            self.heading_parts.append(s)

    def result(self):
        text = "".join(self.parts)
        text = re.sub(r" *\n *", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        heading = re.sub(r"\s+", " ", "".join(self.heading_parts)).strip()[:120]
        return text, heading


def _local_name(tag):
    return str(tag or "").split("}")[-1]


def _extract_epub(epub, max_units=None):
    pages = []
    with zipfile.ZipFile(epub) as zf:
        _check_zip_limits(zf)
        container = ET.fromstring(zf.read("META-INF/container.xml"))
        rootfile = next((e.attrib.get("full-path") for e in container.iter()
                         if _local_name(e.tag) == "rootfile" and e.attrib.get("full-path")), "")
        if not rootfile:
            raise ValueError("EPUB 缺少 OPF 根文件")
        opf = ET.fromstring(zf.read(rootfile))
        manifest = {}
        for e in opf.iter():
            if _local_name(e.tag) == "item" and e.attrib.get("id") and e.attrib.get("href"):
                manifest[e.attrib["id"]] = (e.attrib["href"], e.attrib.get("media-type", ""))
        spine = [e.attrib.get("idref") for e in opf.iter()
                 if _local_name(e.tag) == "itemref" and e.attrib.get("idref")]
        base = posixpath.dirname(rootfile)
        for idref in spine:
            if max_units and len(pages) >= max_units:
                break
            href, media = manifest.get(idref, ("", ""))
            if not href or (media and "html" not in media and "xml" not in media):
                continue
            member = posixpath.normpath(posixpath.join(base, unquote(href.split("#", 1)[0])))
            if member.startswith("../"):
                continue
            raw = zf.read(member)
            parser = _HTMLText()
            parser.feed(raw.decode("utf-8", errors="replace"))
            text, heading = parser.result()
            if not text:
                continue
            n = len(pages) + 1
            label = heading or f"EPUB 第 {n} 章节"
            pages.append({"page": n, "text": text, "source": "epub", "confidence": 1.0,
                          "heading": heading, "locator_type": "chapter", "locator_label": label})
    if not pages:
        raise ValueError("EPUB 未提取到可读正文")
    return pages


def _extract_docx(docx_path, max_units=None):
    # 先用 zipfile 做体积预检，再交给项目已安装的 python-docx。
    with zipfile.ZipFile(docx_path) as zf:
        _check_zip_limits(zf)
    from docx import Document
    doc = Document(str(docx_path))
    blocks = list(doc.iter_inner_content()) if hasattr(doc, "iter_inner_content") else list(doc.paragraphs)
    pages, heading, buf = [], "", []

    def flush():
        nonlocal buf
        text = "\n".join(x for x in buf if x.strip()).strip()
        if text:
            n = len(pages) + 1
            label = heading or f"DOCX 第 {n} 节"
            pages.append({"page": n, "text": text, "source": "docx", "confidence": 1.0,
                          "heading": heading, "locator_type": "section", "locator_label": label})
        buf = []

    for block in blocks:
        if max_units and len(pages) >= max_units:
            break
        if hasattr(block, "rows"):
            rows = []
            for row in block.rows:
                vals = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells]
                if any(vals):
                    rows.append("\t".join(vals))
            if rows:
                buf.append("\n".join(rows))
            continue
        text = (getattr(block, "text", "") or "").strip()
        style = str(getattr(getattr(block, "style", None), "name", "") or "")
        if text and style.lower().startswith("heading"):
            flush()
            heading = text[:120]
            buf.append(text)
        elif text:
            buf.append(text)
    flush()
    if max_units:
        pages = pages[:max_units]
    if not pages:
        raise ValueError("DOCX 未提取到可读正文")
    return pages


def _decode_text_file(file_path):
    raw = Path(file_path).read_bytes()
    if b"\x00" in raw[:4096] and not (raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff")):
        raise ValueError("文本文件疑似二进制内容")
    for enc in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            return raw.decode(enc)
        except UnicodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_markdown(md_path, max_units=None):
    text = _decode_text_file(md_path)
    pages, heading, buf = [], "", []
    for line in text.splitlines():
        m = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$", line)
        if m:
            if buf:
                body = "\n".join(buf).strip()
                if body:
                    n = len(pages) + 1
                    pages.append({"page": n, "text": body, "source": "markdown", "confidence": 1.0,
                                  "heading": heading, "locator_type": "section",
                                  "locator_label": heading or f"Markdown 第 {n} 节"})
                    if max_units and len(pages) >= max_units:
                        break
            heading, buf = m.group(2).strip()[:120], [line]
        else:
            buf.append(line)
    if (not max_units or len(pages) < max_units) and buf:
        body = "\n".join(buf).strip()
        if body:
            n = len(pages) + 1
            pages.append({"page": n, "text": body, "source": "markdown", "confidence": 1.0,
                          "heading": heading, "locator_type": "section",
                          "locator_label": heading or f"Markdown 第 {n} 节"})
    if not pages:
        raise ValueError("Markdown 未提取到可读正文")
    return pages


def _extract_txt(txt_path, max_units=None):
    lines = _decode_text_file(txt_path).splitlines()
    pages, start = [], 0
    while start < len(lines):
        if max_units and len(pages) >= max_units:
            break
        end = min(len(lines), start + 200)
        # 尽量在空行边界切分，避免一段话被截断。
        if end < len(lines):
            for i in range(end, max(start + 1, end - 30), -1):
                if not lines[i - 1].strip():
                    end = i
                    break
        body = "\n".join(lines[start:end]).strip()
        if body:
            n = len(pages) + 1
            pages.append({"page": n, "text": body, "source": "txt", "confidence": 1.0,
                          "heading": "", "locator_type": "lines",
                          "locator_label": f"第 {start + 1}–{end} 行",
                          "line_start": start + 1, "line_end": end})
        start = max(end, start + 1)
    if not pages:
        raise ValueError("TXT 未提取到可读正文")
    return pages


def _extract_source_document(source_path, source_format="", max_units=None,
                             ocr_mode="off", ocr_engine=None, on_ocr_pending=None):
    """统一提取入口；返回结构继续使用 pages 兼容旧切块/Agent API。"""
    fmt = source_format or DF.detect_format(source_path)
    if fmt == "pdf":
        result = _extract_document(source_path, max_pages=max_units, ocr_mode=ocr_mode,
                                   ocr_engine=ocr_engine, on_ocr_pending=on_ocr_pending)
        for pg in result["pages"]:
            pg.setdefault("locator_type", "page")
            pg.setdefault("locator_label", f"PDF 第 {pg.get('page')} 页")
            pg.setdefault("heading", "")
        result["document_format"] = "pdf"
        return result
    extractors = {
        "epub": _extract_epub,
        "docx": _extract_docx,
        "markdown": _extract_markdown,
        "txt": _extract_txt,
    }
    if fmt not in extractors:
        raise ValueError(f"不支持的全文格式：{fmt or Path(source_path).suffix}")
    pages = extractors[fmt](source_path, max_units=max_units)
    return {"pages": pages, "total_pages": len(pages), "native_pages": len(pages),
            "ocr_pages": 0, "empty_pages": 0, "ocr_confidence": None,
            "ocr_errors": [], "document_format": fmt}

def load_papers():
    if not C.PAPERS_JSONL.exists():
        return []
    return [DF.normalize_record(json.loads(l))
            for l in open(C.PAPERS_JSONL, encoding="utf-8") if l.strip()]

def filter_scope(papers, scope):
    todo = [p for p in papers if p.get("has_fulltext") and p.get("fulltext_path")]
    if not scope or scope == "all":
        return todo
    if scope.startswith("collection:"):
        col = scope.split(":", 1)[1]
        return [p for p in todo if col in (p.get("collections") or [])]
    if scope.startswith("keys:"):
        ks = set(k.strip() for k in scope.split(":", 1)[1].split(",") if k.strip())
        return [p for p in todo if p.get("key") in ks]
    return todo

def _prev_ok(out):
    """读既有提取产物的 ok 标志：读不到/损坏 → None。"""
    try:
        return json.loads(out.read_text(encoding="utf-8")).get("ok")
    except Exception:
        return None

def extract_one(p):
    # 兼容直接调用本函数的旧客户端/单测：旧记录只有 pdf_path/has_pdf。
    p = DF.normalize_record(dict(p))
    out = C.EXTRACTED / f"{p['stem']}.json"
    if out.exists() and out.stat().st_size > 0:
        prev = _prev_ok(out)
        # 成功产物（有正文）→ 跳过；损坏/读不出 → 也跳过（重抽也是同样输入）。
        # 但失败产物(ok=False：曾 PDF 不在盘/被占用/坏 PDF) 若现在 PDF 可读 → 重抽，
        # 别再因『产物文件已存在』把修好的 PDF 永久跳过（BF：Zotero 链接附件/OneDrive 占用等瞬时错误）。
        if prev is None or prev:
            return "skip"
        _source = p.get("fulltext_path")
        if not (_source and os.path.exists(_source)):
            return "skip"     # 附件仍不可读，重抽无意义（避免每轮空转）
        # 落到下面重抽
    source_path = p.get("fulltext_path")
    source_format = p.get("fulltext_format") or DF.detect_format(source_path)
    meta = {k: p.get(k) for k in META_FIELDS}
    meta["collections"] = p.get("collections", [])
    rec = {"key": p["key"], "stem": p["stem"], "meta": meta, "pages": [],
           "ok": False, "status": "ocr_pending", "document_format": source_format}
    if not source_path or not os.path.exists(source_path):
        rec["error"] = "no_source_on_disk"
        rec["status"] = "missing_pdf" if source_format == "pdf" else "missing_file"
        out.write_text(json.dumps(rec, ensure_ascii=False), encoding="utf-8")
        try:
            DES.set_status(p["stem"], rec["status"], error=rec["error"], total_pages=0,
                           native_pages=0, ocr_pages=0, empty_pages=0, ocr_confidence=None)
        except Exception:
            pass
        return "nofile"
    try:
        def _pending(stats):
            DES.set_status(p["stem"], "ocr_pending", error="", **stats,
                           ocr_pages=0,
                           empty_pages=max(0, stats.get("total_pages", 0)
                                           - stats.get("native_pages", 0)),
                           ocr_confidence=None)

        result = _extract_source_document(source_path, source_format=source_format,
                                          ocr_mode="empty_pages", on_ocr_pending=_pending)
        rec["pages"] = result["pages"]
        for fld in ("total_pages", "native_pages", "ocr_pages", "empty_pages",
                    "ocr_confidence"):
            rec[fld] = result[fld]
        if result["ocr_errors"]:
            rec["ocr_errors"] = result["ocr_errors"]
            rec["error"] = (f"OCR 有 {len(result['ocr_errors'])} 页失败："
                            + "；".join(result["ocr_errors"][:3]))
        elif result["empty_pages"]:
            rec["error"] = f"OCR 后仍有 {result['empty_pages']} 页没有识别出有效文字"
        rec["ok"] = len(rec["pages"]) > 0
        if source_format != "pdf":
            rec["status"] = "ok_text" if rec["ok"] else "invalid_file"
        elif result["ocr_pages"]:
            rec["status"] = "ok_ocr"
        elif rec["ok"]:
            rec["status"] = "ok_native"
        else:
            rec["status"] = "ocr_failed"
            rec["error"] = rec.get("error") or "OCR 未识别出有效文字"
        out.write_text(json.dumps(rec, ensure_ascii=False), encoding="utf-8")
        try:
            DES.set_status(p["stem"], rec["status"], error=rec.get("error", ""),
                           total_pages=result["total_pages"], native_pages=result["native_pages"],
                           ocr_pages=result["ocr_pages"], empty_pages=result["empty_pages"],
                           ocr_confidence=result["ocr_confidence"])
        except Exception:
            pass
        return "ocr" if result["ocr_pages"] else ("ok" if rec["ok"] else "empty")
    except Exception as e:
        rec["error"] = f"{type(e).__name__}: {e}"
        rec["status"] = "invalid_pdf" if source_format == "pdf" else "invalid_file"
        try:
            out.write_text(json.dumps(rec, ensure_ascii=False), encoding="utf-8")
            DES.set_status(p["stem"], rec["status"], error=rec["error"], total_pages=0,
                           native_pages=0, ocr_pages=0, empty_pages=0,
                           ocr_confidence=None)
        except Exception:
            pass
        return "error"

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--scope", default="all")
    ap.add_argument("--workers", type=int, default=4)
    ap.add_argument("--limit", type=int, default=0)
    args = ap.parse_args()

    # 升级迁移：把旧 deep_no_text 的「扫描件/附件丢失/坏 PDF」拆开，并让
    # ocr_pending 从旧排除集合重新进入本轮候选。幂等，可每次 extract 开头调用。
    DES.reconcile_legacy()
    papers = load_papers()
    todo = filter_scope(papers, args.scope)
    if args.limit:
        todo = todo[:args.limit]
    def _needs(p):
        f = C.EXTRACTED / f"{p['stem']}.json"
        if not (f.exists() and f.stat().st_size > 0):
            return True
        try:
            old_rec = json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            old_rec = {}
        ok = old_rec.get("ok")
        if ok is None or ok:
            return False              # 成功/损坏产物：跳过
        # OCR 已真实跑过仍失败：保持粘性，只有「重试失败篇」清产物后才重跑，避免每轮
        # 对同一本扫描件反复烧 CPU。ocr_pending（旧版本迁来的）则必须自动重试。
        if old_rec.get("status") == "ocr_failed":
            return False
        source_path = p.get("fulltext_path")       # 失败产物：仅当附件现在可读才重抽
        return bool(source_path and os.path.exists(source_path))
    todo = [p for p in todo if _needs(p)]
    print(f"[extract] scope={args.scope}  待提取 {len(todo)} 篇（有可读全文、未提取或失败可重抽）", flush=True)

    t0 = time.time(); done = {}
    with ThreadPoolExecutor(max_workers=args.workers) as ex:
        futs = {ex.submit(extract_one, p): p["key"] for p in todo}
        for i, f in enumerate(as_completed(futs), 1):
            try:
                s = f.result()
            except Exception:
                s = "error"
            done[s] = done.get(s, 0) + 1
            if i % 10 == 0 or i == len(todo):
                print(f"  {i}/{len(todo)}  {time.time()-t0:.0f}s  {done}", flush=True)
    print(f"[extract] 完成 {done}  用时 {time.time()-t0:.0f}s", flush=True)

if __name__ == "__main__":
    main()
