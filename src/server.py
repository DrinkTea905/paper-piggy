# -*- coding: utf-8 -*-
"""
本地知识库 web 服务（检索 + RAG对话 + 三档索引，一个进程全包）。
数据源：直接读 zotero.sqlite（绕过 Better BibTeX）。
两种用法：① 内置 chat（用户自备 LLM key）② 检索 API（任何 agent 调 /search，不需 key）。
错误日志：后端异常 + 前端上报都写 logs/errors.log，GET /errors 可导出（方便反馈问题）。
"""
import sys, os, json, time, threading, subprocess, traceback, re
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
import config as C
import retriever as R
import llm as L
import wiki_store as W
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional, List
import uvicorn

app = FastAPI(title="本地知识库")
#  proc/cancelled：整库深索(scope=all)跑在 subprocess 里，此前不留句柄 → 一旦开始无法停，
#  可能空跑数小时并烧掉 API 额度。存下 Popen 供 POST /build/cancel 终止。
BUILD = {"running": False, "stage": None, "log": [], "started": None, "rc": None,
         "proc": None, "cancelled": False,
         "bulk": False}   # BF35：当前 deep 构建是否整库深索(scope=all)，队列批次为 False
# B1：build 守卫锁——把「判 running + 置 True」做成原子并在调用线程内同步置位，
# 杜绝多路触发并发起两个 build_all 子进程写坏同一 LanceDB/bm25/papers.jsonl。
_BUILD_LOCK = threading.Lock()
# 补生成摘要（SAC backfill）后台任务状态：给「已深索但缺检索摘要」的篇补摘要 + 重嵌入。
# running 期间禁止重复触发；phase=生成中/重嵌入中；前端经 /index/status 的 sac_backfill 读它。
BACKFILL = {"running": False, "phase": "", "done": 0, "total": 0, "fail": 0, "msg": "", "at": 0.0}
_BACKFILL_LOCK = threading.Lock()

@app.middleware("http")
async def _no_cache(request, call_next):
    """前端频繁迭代阶段禁用浏览器缓存，避免 HTML/JS 版本不匹配。"""
    resp = await call_next(request)
    p = request.url.path
    if p == "/" or p.startswith("/static"):
        resp.headers["Cache-Control"] = "no-store, must-revalidate"
        # 任务六：给文本类响应补 charset=utf-8，避免 WebView2 按系统 GBK 解码整页出现乱码
        # （StaticFiles 在 Windows 上常返回不带 charset 的 text/html、text/css）。
        ct = resp.headers.get("content-type", "")
        if ct and "charset" not in ct.lower() and (
                ct.startswith("text/") or "javascript" in ct or "json" in ct):
            resp.headers["content-type"] = ct + "; charset=utf-8"
    return resp
ERRLOG = C.LOGS / "errors.log"

def _last4(s):
    """K3：返回 key 的末4位（用于掩码显示「已填」状态）；不足4位或空则返回 ""。绝不回完整明文。"""
    s = (s or "").strip()
    return s[-4:] if len(s) >= 4 else ""

# ── 错误日志 ──────────────────────────────────────────────
def log_error(where, err, tb=""):
    try:
        ERRLOG.parent.mkdir(parents=True, exist_ok=True)
        with open(ERRLOG, "a", encoding="utf-8") as f:
            f.write(f"\n===== {time.strftime('%Y-%m-%d %H:%M:%S')} · {where} =====\n{err}\n")
            if tb:
                f.write(tb + "\n")
    except Exception:
        pass

@app.exception_handler(Exception)
async def _all_exc(request: Request, exc: Exception):
    log_error(f"{request.method} {request.url.path}", repr(exc), traceback.format_exc())
    return JSONResponse({"error": "服务器内部错误（已记入错误日志，可在设置里导出）", "detail": str(exc)}, status_code=500)

class LogQ(BaseModel):
    level: str = "error"
    msg: str = ""
    ctx: str = ""

@app.post("/log")
def client_log(q: LogQ):
    log_error(f"前端[{q.level}] {q.ctx}", q.msg)
    return {"ok": True}

@app.get("/errors")
def get_errors(n: int = 200):
    if not ERRLOG.exists():
        return {"errors": "（暂无错误记录）", "lines": 0}
    lines = ERRLOG.read_text(encoding="utf-8", errors="replace").splitlines()
    return {"errors": "\n".join(lines[-n:]), "lines": len(lines)}

@app.post("/errors/clear")
def clear_errors():
    try:
        if ERRLOG.exists():
            ERRLOG.unlink()
    except Exception:
        pass
    return {"ok": True}

# ── 自动更新：源变化(Zotero 新条目/文件夹新 PDF)时后台【按天+指定时刻】增量更新 ──
AUTO = {"sig": None, "last": 0, "last_check": 0.0, "last_build": 0.0, "last_result": ""}
_AUTO_STATE_FILE = C.STATE / "auto_update_state.json"

def _auto_days(conf):
    """按天间隔，clamp 到 1..30。"""
    try:
        return min(30, max(1, int(conf.get("interval_days", 1))))
    except Exception:
        return 1

def _parse_hhmm(s):
    try:
        h, m = str(s).split(":")
        h = min(23, max(0, int(h))); m = min(59, max(0, int(m)))
        return h, m
    except Exception:
        return 7, 0

def _auto_next_after(last_ts, interval_days, at_time):
    """下次应跑的 epoch 秒：从 last_ts 那天起隔 interval_days 天、当天 at_time 时刻；
       从没跑过(last_ts<=0)则取今天 at_time。"""
    import datetime as _dt
    h, m = _parse_hhmm(at_time)
    try:
        if last_ts and last_ts > 0:
            base = _dt.datetime.fromtimestamp(last_ts).date() + _dt.timedelta(days=max(1, int(interval_days)))
        else:
            base = _dt.date.today()
        return _dt.datetime.combine(base, _dt.time(h, m)).timestamp()
    except Exception:
        return 0.0

def _auto_state_load():
    try:
        if _AUTO_STATE_FILE.exists():
            d = json.loads(_AUTO_STATE_FILE.read_text(encoding="utf-8"))
            return float(d.get("last_check", 0) or 0), float(d.get("last_build", 0) or 0)
    except Exception:
        pass
    return 0.0, 0.0

def _auto_state_save():
    try:
        _atomic_json_write(_AUTO_STATE_FILE, {
            "last_check": AUTO.get("last_check", 0), "last_build": AUTO.get("last_build", 0),
            "saved_at": time.strftime("%Y-%m-%d %H:%M:%S")})
    except Exception as e:
        log_error("auto state save", repr(e))
def _source_signature():
    """返回当前数据源的"内容指纹"，变化即代表有新增/改动。"""
    try:
        import settings as S
        if S.source() == "folder":
            import folder_source as FS
            pdfs = FS.scan(S.folder_dir())
            mt = max((Path(p).stat().st_mtime for p in pdfs), default=0)
            return f"folder:{len(pdfs)}:{int(mt)}"
        import zotero_source as Z
        dd = Z.detect_data_dir()
        sq = (Path(dd) / "zotero.sqlite") if dd else None
        if not (sq and sq.exists()):
            return "zotero:none"
        # R7：WAL 模式下新增条目常只落 -wal、主库 mtime 暂不动 → 自动更新检测不到。
        # 把 zotero.sqlite-wal 的 mtime/size 一并纳入指纹，新增即可被捕获。
        sig = f"zotero:{int(sq.stat().st_mtime)}"
        wal = sq.with_name("zotero.sqlite-wal")
        if wal.exists():
            try:
                ws = wal.stat()
                sig += f":wal{int(ws.st_mtime)}:{ws.st_size}"
            except Exception:
                pass
        return sig
    except Exception:
        return None

# ── EN-W1：wiki lint 挂上自动更新定时器——gist 点名 drift（wiki 与文献库渐行渐远）是头号
#    失败模式，体检必须自动跑，不能指望用户想起来点按钮。lint 纯读、零 LLM、毫秒级，
#    顺搭在自动更新循环里即可，不必单开线程。
_WIKI_LINT_TTL = 24 * 3600      # 距上次体检超过 24h（或从没体检过）才重跑

def _wiki_lint_refresh():
    """距 data/state/wiki_lint.json 的 checked_at 超过 24h（或文件不存在/损坏）→ 重跑 W.lint()，
       把 {"issues":总数, "by_type":{...}, "checked_at":now} 原子写回。/stats 读它（契约3）。
       wiki 尚无任何页（index.json 由首次写页产生）时优雅跳过——别为了体检去初始化空库。"""
    f = C.STATE / "wiki_lint.json"
    try:
        if f.exists():
            d = json.loads(f.read_text(encoding="utf-8"))
            ts = time.mktime(time.strptime(d.get("checked_at", ""), "%Y-%m-%d %H:%M:%S"))
            if time.time() - ts < _WIKI_LINT_TTL:
                return
    except Exception:
        pass                                # 文件损坏/时间戳不合法 → 视为过期，重算一份盖掉
    if not C.WIKI_INDEX.exists():
        return                              # wiki 目录为空/未初始化：跳过
    res = W.lint()
    _atomic_json_write(f, {
        "issues": int(res.get("n_issues", 0)),
        "by_type": {k: len(v) for k, v in (res.get("issues") or {}).items()},
        "checked_at": time.strftime("%Y-%m-%d %H:%M:%S")})

def _auto_update_loop():
    import settings as S
    time.sleep(20)                         # 等首次加载完
    AUTO["sig"] = _source_signature()      # 建基线：启动时不误触发
    lc, lb = _auto_state_load()            # 跨重启：上次检查/构建时间（catch-up 命脉）
    AUTO["last_check"] = lc; AUTO["last_build"] = lb
    # 关了「补跑」而当前已到点(关机期间错过)：把 last_check 提到现在，跳过错过的档、等下一个 at_time。
    try:
        conf0 = S.load().get("auto_update", {}) or {}
        if not conf0.get("catch_up_on_launch", True):
            if time.time() >= _auto_next_after(AUTO["last_check"], _auto_days(conf0),
                                               conf0.get("at_time", "07:00")):
                AUTO["last_check"] = time.time(); _auto_state_save()
    except Exception as e:
        log_error("auto catch-up init", repr(e))
    while True:
        try:
            conf = S.load().get("auto_update", {}) or {}
            days = _auto_days(conf)
            at = conf.get("at_time", S.DEFAULT["auto_update"]["at_time"])
            time.sleep(60)
            # EN-W1：每轮顺带保鲜 wiki 体检（刻意放在 enabled 判断之前——体检零成本，
            # 不该随「自动更新」开关一起被关掉）；异常只记日志，绝不阻断自动更新主流程。
            try:
                _wiki_lint_refresh()
            except Exception as e:
                log_error("wiki lint refresh", repr(e))
            # 自动备份：同样刻意放在「自动更新 enabled」判断之前 —— 备份有自己的开关
            # （settings.backup.auto），不该被「自动更新知识库」这个无关的开关一起关掉。
            try:
                _auto_backup_tick()
            except Exception as e:
                log_error("auto backup", repr(e))
            if not conf.get("enabled", True):
                continue
            # 按天+时刻调度：未到下次计划时刻就跳过（catch-up 靠持久化的 last_check 自然生效）。
            if time.time() < _auto_next_after(AUTO["last_check"], days, at):
                continue
            AUTO["last_check"] = time.time(); _auto_state_save()
            sig = _source_signature()
            if sig and sig != AUTO["sig"] and not BUILD["running"]:
                stage = "folder" if S.source() == "folder" else "all"   # 只轻量层+语义，深索永远手动
                # A4：只有 build 真正成功(rc==0)才推进 sig，失败留待下轮重试（与 A3 returncode 贯通）。
                def _auto_done(rc, _sig=sig):
                    # append 到尾部（与其余日志一致；insert(0) 会插在窗口外、且下次构建即被重置，用户永远看不到）；
                    # 另存 AUTO["last_result"] 供 /index/status 直接返回，前端可展示上次自动更新结果。
                    if rc == 0:
                        AUTO["sig"] = _sig
                        AUTO["last_build"] = time.time(); _auto_state_save()
                        AUTO["last_result"] = "检测到新增文献，已自动增量更新。"
                        BUILD["log"].append("[auto] 检测到新增文献，已自动增量更新。")
                    else:
                        AUTO["last_result"] = f"自动增量更新未成功(rc={rc})，下轮将重试。"
                        BUILD["log"].append(f"[auto] 自动增量更新未成功(rc={rc})，下轮将重试。")
                    try:
                        _drain_deep_queue()   # 沿用旧默认行为：build 后推进自动深索队列
                    except Exception as e:
                        log_error("auto build drain", repr(e))
                _run_build(stage, on_done=_auto_done)
        except Exception as e:
            log_error("auto_update loop", repr(e))
            time.sleep(60)


def _retrieval_idle_loop():
    """按用户设置释放检索组件；R 内部的活动计数保证不会打断正在进行的 Agent/界面检索。"""
    import settings as S
    while True:
        time.sleep(5)
        try:
            mins = int(S.retrieval_conf().get("idle_unload_min", 10))
            if mins > 0:
                R.release_retrieval_if_idle(mins * 60)
        except Exception as e:
            log_error("retrieval idle unload", repr(e))
            time.sleep(10)


class AutoUpdateQ(BaseModel):
    enabled: Optional[bool] = None
    interval_days: Optional[int] = None
    at_time: Optional[str] = None
    catch_up_on_launch: Optional[bool] = None
    delete_sync: Optional[bool] = None       # C3：folder 模式是否同步删除（默认关）
    interval_min: Optional[int] = None       # 兼容旧前端字段（不再用于调度）

def _auto_update_view():
    """自动更新配置的对外视图（兜底统一走 settings.DEFAULT）。"""
    import settings as S
    c = S.load().get("auto_update", {}) or {}
    D = S.DEFAULT["auto_update"]
    return {"enabled": bool(c.get("enabled", True)),
            "interval_days": min(30, max(1, int(c.get("interval_days", D["interval_days"])))),
            "at_time": str(c.get("at_time", D["at_time"])),
            "catch_up_on_launch": bool(c.get("catch_up_on_launch", D["catch_up_on_launch"])),
            "delete_sync": bool(c.get("delete_sync", D.get("delete_sync", False))),
            "interval_min": int(c.get("interval_min", D["interval_min"])),
            "source": S.source(),
            "last_build": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(AUTO["last_build"])) if AUTO.get("last_build") else ""}

@app.get("/setup/auto_update")
def get_auto_update():
    return _auto_update_view()

@app.post("/setup/auto_update")
def set_auto_update(q: AutoUpdateQ):
    import settings as S
    patch = {}
    if q.enabled is not None: patch["enabled"] = bool(q.enabled)
    if q.interval_days is not None: patch["interval_days"] = min(30, max(1, int(q.interval_days)))
    if q.at_time is not None:
        h, m = _parse_hhmm(q.at_time)
        patch["at_time"] = f"{h:02d}:{m:02d}"
    if q.catch_up_on_launch is not None: patch["catch_up_on_launch"] = bool(q.catch_up_on_launch)
    if q.delete_sync is not None: patch["delete_sync"] = bool(q.delete_sync)
    if q.interval_min is not None: patch["interval_min"] = max(5, int(q.interval_min))
    S.save({"auto_update": patch})
    return {"ok": True, **_auto_update_view()}


class RetrievalMemoryQ(BaseModel):
    idle_unload_min: int


def _retrieval_memory_view():
    import settings as S
    mins = int(S.retrieval_conf().get("idle_unload_min", 10))
    st = R.retrieval_status()
    remaining = None
    if mins > 0 and st["loaded"] and not st["loading"]:
        remaining = max(0, mins * 60 - int(st["idle_s"]))
    return {"idle_unload_min": mins, "remaining_s": remaining, **st}


@app.get("/setup/retrieval_memory")
def get_retrieval_memory():
    return _retrieval_memory_view()


@app.post("/setup/retrieval_memory")
def set_retrieval_memory(q: RetrievalMemoryQ):
    import settings as S
    mins = int(q.idle_unload_min)
    if mins != 0 and not 1 <= mins <= 1440:
        raise HTTPException(status_code=400, detail="空闲释放时间请设为 1–1440 分钟；0 表示不自动释放")
    S.save({"retrieval": {"idle_unload_min": mins}})
    if mins > 0:
        R.release_retrieval_if_idle(mins * 60)
    return {"ok": True, **_retrieval_memory_view()}

@app.post("/setup/purge_deleted")
def purge_deleted():
    """删除同步：把「源里已删、库里还留着」的文献清出索引。
       folder 模式——删除同步在 stage=folder 增量构建里已内建，这里直接触发一次增量。
       zotero 模式——算差集手动清理（带安全阈值：读不到活库/一次要删过半则中止，防误抹整库）。"""
    import settings as S
    src = S.source()
    if src == "folder":
        ok = _run_build("folder")
        return {"ok": ok, "mode": "folder",
                "msg": "已触发文件夹增量更新——删除的 PDF 会在本次更新中自动清出。" if ok else "有任务在跑，稍后再试。"}
    # zotero 模式
    try:
        import zotero_source as Z
        live = {p.get("key") for p in Z.load_papers() if p.get("key")}
    except Exception as e:
        return JSONResponse({"ok": False, "msg": f"读取 Zotero 失败，已中止：{e}"}, status_code=400)
    if not live:
        return JSONResponse({"ok": False, "msg": "未从 Zotero 读到任何文献，已中止（避免误删整库）。"}, status_code=400)
    indexed = set(_load_papers().keys())
    gone = [k for k in indexed if k not in live]
    if not gone:
        return {"ok": True, "mode": "zotero", "removed": 0, "msg": "没有需要清理的已删除文献。"}
    if len(gone) > max(20, len(indexed) // 2):
        return JSONResponse({"ok": False, "removed": 0,
            "msg": f"检测到 {len(gone)} 篇需清理（超过库的一半），已中止以防误删。请确认 Zotero 库完整后重试。"},
            status_code=400)
    if BUILD["running"]:
        return JSONResponse({"ok": False, "msg": "有构建任务在跑，稍后再清理。"}, status_code=400)
    try:
        import folder_ingest as FI
        FI._purge_db_rows(gone)
        FI._purge_key_artifacts(gone)
    except Exception as e:
        log_error("purge_deleted zotero", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "msg": f"清理失败：{e}"}, status_code=400)
    try:
        R.STATE["ready"] = False; R.load_all()
    except Exception as e:
        log_error("purge_deleted reload", repr(e))
    return {"ok": True, "mode": "zotero", "removed": len(gone),
            "msg": f"已清理 {len(gone)} 篇 Zotero 中已删除的文献。建议随后点一次「手动更新知识库」刷新题录层。"}

# ── 启动加载 ──────────────────────────────────────────────
_LOADING = True   # 冷启动读取库目录与轻量句柄期间为 True；/health 暴露给前端显示启动遮罩，完成后置 False。
@app.on_event("startup")
def _startup():
    try:
        # 建 wiki 目录 + 把 WIKI.md 升到当前 schema 版本。它会被 MCP 整篇下发给 agent，
        # 过期的话 agent 就照着旧规约干活（不知道 entity 页、不知道该 mark_stale 而非覆盖）。
        W.ensure_scaffold()
    except Exception as e:
        log_error("startup ensure_scaffold", repr(e))
    try:
        # BF34：深索标记文件（追加写）在重试/并发下会积累重复行，计数与判断随之虚高——启动自愈去重
        _dedup_deep_marks()
    except Exception as e:
        log_error("startup dedup deep marks", repr(e))
    try:
        # 旧版把附件缺失、坏 PDF、真扫描件混在 deep_no_text.txt。启动即迁移，
        # 用户不必先手动跑一次深索，浏览页也能看到真实原因。
        import deep_extract_status as DES
        DES.reconcile_legacy()
    except Exception as e:
        log_error("startup reconcile deep extract status", repr(e))
    try:
        # Agent 专属文件夹（0_Agent交付物 / 0_Agent资料库）幂等脚手架，人类可读、换 agent 可续。
        import agent_ws as AW
        AW.ensure_scaffold()
    except Exception as e:
        log_error("startup agent workspace scaffold", repr(e))
    threading.Thread(target=_safe_load, daemon=True).start()
    threading.Thread(target=_auto_update_loop, daemon=True).start()
    threading.Thread(target=_retrieval_idle_loop, daemon=True).start()

def _safe_load():
    global _LOADING
    try:
        R.load_all()
    except Exception as e:
        log_error("startup load_all", repr(e), traceback.format_exc())
        print("[server] 索引加载失败：", e, flush=True)
    _LOADING = False   # 准备阶段结束（成败都撤）——失败时露出空/错误态，而不是永久遮罩。
    # F38-B：启动后台预热当前学科的期刊分级缓存（首次冷算 20+s，异步不阻塞；之后落盘永久快）
    try:
        import grading_svc as GS
        GS.warm_async(_load_papers())
    except Exception as e:
        log_error("startup grading warm", repr(e))
    # F10：回灌上次崩溃残留的自动深索队列，续跑
    try:
        _q_boot()
    except Exception as e:
        log_error("startup deep queue boot", repr(e))

# ── 状态 ──────────────────────────────────────────────────
def _lib_rev():
    """便宜的「库修订号」（分域），供前端 4s 轮询感知、外部/后台改动后自动刷新当前可见页而无需手点或重开。
       只 stat 少数几个「一改就重写 / 一增就变」的规范落点，不读内容、不递归 —— 每次 /health 调用即算，<1ms。
       ★ 有意不含 index_manifest.json：深索时它每批都被重写，含进来会让前端每 4s 过刷。深索进度另有专门扳机。
       分域（lib/wiki/agent）是为了让前端只刷「真变了的那一域对应的可见页」，避免跨域误刷（如浏览页看列表时
       agent 写了综述，不该把浏览列表也刷得跳一下）。"""
    def _mt(p):
        try:
            return int(p.stat().st_mtime)
        except Exception:
            return 0
    lib = _mt(C.PAPERS_JSONL)                       # 题录（入库/去重后重写）
    wiki = _mt(C.DATA / "wiki" / "index.json")      # 综述层权威事实源（save/update/mark_stale 都重写）
    agent = 0
    try:
        import agent_ws as AW
        agent = _mt(AW.output_dir()) + _mt(AW.tasks_dir())   # 交付物主题夹 / 定时任务夹（新增即变）
    except Exception:
        pass
    return {"lib": lib, "wiki": wiki, "agent": agent}

@app.get("/health")
def health():
    mode = R.STATE.get("mode")
    papers = len(R.M.get("papers", {}))          # 去重篇数（题录数，L/F 档都在内存）
    blocks = int(R.M.get("row_count", 0)) if mode == "full" else 0  # LanceDB 总行数；正文/向量不常驻 Python
    n = blocks if mode == "full" else papers     # 兼容旧字段
    return {"ready": R.STATE.get("ready", False), "mode": mode,
            "n": n, "papers": papers, "blocks": blocks, "building": BUILD["running"],
            "deep": len(_deep_keys()),          # F10：「全部文献」显示 已深索/总数
            "rev": _lib_rev(),                  # 库修订号（分域）：前端据此在库/综述/交付物变动后自动刷可见页
            "loading": _LOADING,                # 冷启动读取库目录期间为 True → 前端显示遮罩，完成后淡出
            "retrieval": R.retrieval_status(),  # 检索组件冷/热态；设置页据此解释内存策略
            "pid": os.getpid()}                 # 本 server 进程 pid → launcher 关窗时按 pid taskkill /T 杀整棵树，根治 orphan 堆积

# ── Agent / MCP 接入信息（给应用内 Agent 页，吐出本机真实可用的接入命令）──
@app.get("/agent/mcp-config")
def agent_mcp_config():
    py = sys.executable                        # 正在跑本服务的解释器＝mcp_server 拉起 server 用的同一个
    # 兜底：若 launcher 用非 python.exe 宿主启动（sys.executable 指向 launcher），
    # 改拼分发版内置的 LocalKB/python/python.exe（已确认该目录存在）。
    try:
        cand = C.APP.parent / "python" / "python.exe"
        if (not py or not Path(py).exists() or Path(py).name.lower() not in ("python.exe", "pythonw.exe", "python", "python3")) and cand.exists():
            py = str(cand)
    except Exception:
        pass
    mcp = str(C.APP / "mcp_server.py")
    def q(s): return '"' + s + '"'             # 路径带空格/中文，命令行统一加引号
    add_core = f'claude mcp add localkb -- {q(py)} {q(mcp)}'
    mcp_json = json.dumps(
        {"mcpServers": {"localkb": {"command": py, "args": [mcp]}}},
        ensure_ascii=False, indent=2)
    codex_toml = (f'[mcp_servers.localkb]\n'
                  f'command = {json.dumps(py)}\n'
                  f'args = [{json.dumps(mcp)}]')
    try:                                       # 真实 MCP 工具数（前端"看到 N 个工具"用它，别写死）
        import mcp_server as MCP
        tool_count = len(MCP.TOOLS)
    except Exception:
        tool_count = 0
    # Agent 专属文件夹路径（Agent 页展示「打开交付物/资料库/技能」用）
    ag = {}
    try:
        import agent_ws as AW
        AW.ensure_scaffold()
        ag = AW.paths_info()
    except Exception as e:
        log_error("agent mcp-config paths", repr(e))
    try:
        import upgrade_health as UH
        upgrade = UH.health()
    except Exception as e:
        log_error("agent upgrade health", repr(e))
        upgrade = {"pending_count": 0, "template_items": [], "error": str(e)}
    return {
        "python": py, "mcp_server": mcp,
        "daemon_url": C.DAEMON_URL, "server_running": True,
        "wiki_schema_md": str(C.WIKI_SCHEMA_MD), "tool_count": tool_count,
        "claude_cmd": add_core,
        "claude_cmd_user": add_core.replace("claude mcp add localkb ",
                                            "claude mcp add localkb --scope user "),
        "mcp_json": mcp_json, "codex_toml": codex_toml,
        # 0_Agent交付物 / 0_Agent资料库 落点（前端「打开文件夹」用）
        # 注：曾有个 skill_src_dir 字段指向 app/skills/localkb-paper，是历史遗留——
        # 技能/工作流现在由 agent_ws 统一写到「0_Agent资料库/技能」，不再有独立技能包，前端也从不读它。
        "agent_output_dir": ag.get("output_dir", ""),
        "agent_rely_dir": ag.get("rely_dir", ""),
        "agent_memory_file": ag.get("memory_file", ""),
        "upgrade_health": upgrade,
    }


@app.get("/upgrade/health")
def upgrade_health_status(include_ignored: bool = False):
    import upgrade_health as UH
    return UH.health(include_ignored=include_ignored)


@app.get("/upgrade/diff")
def upgrade_diff(kind: str, key: str):
    import upgrade_health as UH
    try:
        return {"ok": True, "diff": UH.diff(kind, key)}
    except (KeyError, ValueError) as e:
        raise HTTPException(status_code=400, detail=str(e))


class UpgradeActionQ(BaseModel):
    kind: str
    key: str
    current_hash: str
    confirm: str = ""


@app.post("/upgrade/ack")
def upgrade_ack(q: UpgradeActionQ):
    import upgrade_health as UH
    try:
        UH.acknowledge(q.kind, q.key, q.current_hash)
        return {"ok": True}
    except (KeyError, ValueError) as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/upgrade/replace")
def upgrade_replace(q: UpgradeActionQ):
    if q.confirm != "replace_with_factory":
        raise HTTPException(status_code=400, detail="需要明确确认采用新版")
    import upgrade_health as UH
    try:
        backup = UH.replace(q.kind, q.key, q.current_hash)
        return {"ok": True, "backup": backup}
    except (KeyError, ValueError) as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/setup/detect")
def setup_detect():
    zdir = None
    try:
        import zotero_source as Z
        d = Z.detect_data_dir()
        zdir = str(d) if d else None
    except Exception as e:
        log_error("setup/detect zotero", repr(e))
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    import settings as S
    st = S.load()
    backend = st.get("backend", "local")
    api_key_set = bool((st.get("api") or {}).get("key"))
    models_local = (C.MODELS / "bge-m3-onnx" / "model_quantized.onnx").exists()
    reranker_local = (C.MODELS / "bge-reranker-v2-m3-onnx" / "model_quantized.onnx").exists()
    # 数据源：优先 settings.source（用户/向导已选），否则据 manifest / 探到 zotero 推断
    src = st.get("source")
    if src not in ("zotero", "folder"):
        man_src = str(manifest.get("source") or "")
        src = "folder" if man_src.startswith("folder") else ("zotero" if zdir else None)
    try:
        import folder_meta as FM
        meta_ready = FM.available()
    except Exception:
        meta_ready = api_key_set or bool((st.get("sac") or {}).get("key"))
    return {
        "zotero_dir": zdir,
        "source": src,                             # zotero | folder | None
        "folder_dir": st.get("folder_dir", ""),
        "import_only_pdf": bool(st.get("import_only_pdf")),
        "zotero_detected": bool(zdir),             # 与 source 解耦：探到 zotero≠一定用它
        "meta_ready": bool(meta_ready),            # 抽题录 LLM key 是否就绪（folder 模式用）
        "folder_meta_ready": bool(meta_ready),
        "backend": backend,                        # local | api
        "api_key_set": api_key_set,
        # K3：各 key 末4位掩码（只回末4位、绝不回明文），供前端展示「已填 ••••1234」
        "api_key_last4": _last4((st.get("api") or {}).get("key")),
        "sac_key_last4": _last4((st.get("sac") or {}).get("key")),
        # 引擎就绪：本地模式看模型文件；API 模式看 key 是否已填
        "models_ready": models_local if backend == "local" else api_key_set,
        "reranker_ready": reranker_local if backend == "local" else api_key_set,
        "models_local": models_local, "reranker_local": reranker_local,
        "indexed": C.INDEX_MANIFEST.exists(),
        "mode": R.STATE.get("mode"),
    }

class ConnectQ(BaseModel):
    zotero_dir: Optional[str] = None
    source: Optional[str] = None                  # "folder" 时走文件夹分支
    folder_dir: Optional[str] = None

@app.post("/setup/connect")
def setup_connect(q: ConnectQ):
    """校验数据源并返回条目数。source=folder → 选/建文件夹并计 PDF 数；否则读 zotero.sqlite。"""
    import settings as S
    # BF17b：建库子进程正在读源（zotero 临时副本/受管文件夹），此时切换数据源会与之互踩
    if BUILD["running"]:
        return JSONResponse({"ok": False, "msg": "正在建库，稍后再试"}, status_code=400)
    if q.source == "folder":
        try:
            import folder_source as FS
            p = Path(q.folder_dir or "")
            if not str(p).strip():
                return JSONResponse({"ok": False, "msg": "请指定一个文件夹路径"}, status_code=400)
            p.mkdir(parents=True, exist_ok=True)      # 支持"新建空文件夹"
            n = len(FS.scan(str(p)))
            S.save({"source": "folder", "folder_dir": str(p.resolve())})
            try:
                import folder_meta as FM
                mr = FM.available()
            except Exception:
                mr = False
            return {"ok": True, "source": "folder", "entries": n,
                    "dir": str(p.resolve()), "folder_meta_ready": mr}
        except Exception as e:
            log_error("setup/connect folder", repr(e))
            return JSONResponse({"ok": False, "msg": f"无法创建/访问该文件夹：{e}"}, status_code=400)
    try:
        import zotero_source as Z
        if Z.available(q.zotero_dir):
            papers = Z.load_papers(q.zotero_dir)
            n = len(papers)
            with_pdf = sum(1 for p in papers if p.get("has_pdf"))   # F1：向导计数区分「全库 / 将入库」
            # BF2：用户手选的 zotero 数据目录此前校验完即丢，重启后 zotero_source 又退回自动探测；
            # 空串=没手选（沿用自动探测），键名 zotero_dir 与 settings.DEFAULT / zotero_source 约定一致。
            S.save({"source": "zotero", "zotero_dir": q.zotero_dir or ""})
            return {"ok": True, "source": "zotero.sqlite", "entries": n,
                    "with_pdf": with_pdf, "no_pdf": n - with_pdf,
                    "dir": q.zotero_dir or str(Z.detect_data_dir())}
    except Exception as e:
        log_error("setup/connect zotero", repr(e))
        return JSONResponse({"ok": False, "msg": f"读取 zotero.sqlite 失败：{e}"}, status_code=400)
    return JSONResponse({"ok": False, "msg": "未探测到 zotero.sqlite（请确认已安装 Zotero 且库中有文献）"}, status_code=400)

class FolderQ(BaseModel):
    folder_dir: str

@app.post("/setup/folder")
def setup_folder(q: FolderQ):
    """保存文件夹模式选择（选/建文件夹）。"""
    import settings as S
    try:
        import folder_source as FS
        p = Path(q.folder_dir)
        p.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return JSONResponse({"ok": False, "msg": f"无法创建/访问该文件夹：{e}"}, status_code=400)
    n = len(FS.scan(str(p)))
    S.save({"source": "folder", "folder_dir": str(p.resolve())})
    try:
        # 选定受管文件夹后，在其内建 Agent 专属文件夹脚手架（0_Agent* 已排除出文献扫描）。
        import agent_ws as AW
        AW.ensure_scaffold()
    except Exception as e:
        log_error("setup/folder agent scaffold", repr(e))
    try:
        import folder_meta as FM
        mr = FM.available()
    except Exception:
        mr = False
    return {"ok": True, "folder_dir": str(p.resolve()), "pdf_count": n, "folder_meta_ready": mr}

def _default_papers_dir():
    """文件夹模式的建议默认目录：应用自己的数据目录旁 <HOME>/papers（DATA=HOME/data）。"""
    try:
        return str((C.DATA.parent / "papers"))
    except Exception:
        return ""

@app.get("/setup/folder_default")
def setup_folder_default():
    return {"default_dir": _default_papers_dir()}

class ImportOptQ(BaseModel):
    only_pdf: bool = False

@app.post("/setup/import_only_pdf")
def setup_import_only_pdf(q: ImportOptQ):
    """Zotero 模式：是否只导入有 PDF 的条目（切换后需重建即时索引才生效）。"""
    import settings as S
    S.save({"import_only_pdf": bool(q.only_pdf)})
    return {"ok": True, "only_pdf": bool(q.only_pdf)}

@app.post("/setup/open_folder")
def setup_open_folder():
    """在系统文件管理器里打开受管文件夹（副本#7：跳转让用户直接放 PDF）。"""
    import settings as S
    d = S.folder_dir() or _default_papers_dir()
    try:
        Path(d).mkdir(parents=True, exist_ok=True)
        if sys.platform == "win32":
            os.startfile(d)  # noqa
        elif sys.platform == "darwin":
            subprocess.Popen(["open", d])
        else:
            subprocess.Popen(["xdg-open", d])
        return {"ok": True, "dir": d}
    except Exception as e:
        return JSONResponse({"ok": False, "msg": str(e), "dir": d}, status_code=400)

class AgentOpenQ(BaseModel):
    which: str = "rely"                          # output | rely | skills


def _open_system_dir(d: Path):
    """用系统文件管理器打开既有目录；调用方负责限定目录范围。"""
    target = str(d)
    if sys.platform == "win32":
        os.startfile(target)  # noqa
    elif sys.platform == "darwin":
        subprocess.Popen(["open", target])
    else:
        subprocess.Popen(["xdg-open", target])


@app.post("/agent/open_folder")
def agent_open_folder(q: AgentOpenQ):
    """在系统文件管理器里打开 Agent 专属文件夹（交付物 / 资料库 / 技能 / 定时任务）。which 由前端固定传，
       不接受任意路径——避免把打开任意目录的能力暴露给前端。"""
    import agent_ws as AW
    try:
        AW.ensure_scaffold()
        roots = {"output": AW.output_dir, "rely": AW.rely_dir,
                 "skills": AW.skills_dir, "tasks": AW.tasks_dir}
        if q.which not in roots:
            raise HTTPException(status_code=400, detail="不支持的 Agent 文件夹")
        # 技能统一落点=「0_Agent资料库/技能」（含 agent 中立的 工作流.md）；不再打开 app/skills，
        # 也不再往 .claude/skills 自动装——技能只此一处，任何助手读它即可。
        d = Path(roots[q.which]())
        d.mkdir(parents=True, exist_ok=True)
        _open_system_dir(d)
        return {"ok": True, "dir": str(d)}
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse({"ok": False, "msg": str(e)}, status_code=400)


class AgentOpenOutputQ(BaseModel):
    name: str


def _is_link_or_junction(p: Path) -> bool:
    """目录扫描/打开时不跟随软链接或 Windows junction。探测失败也按不安全处理。"""
    try:
        if p.is_symlink():
            return True
        isjunction = getattr(os.path, "isjunction", None)
        return bool(isjunction and isjunction(p))
    except OSError:
        return True


@app.post("/agent/open_output")
def agent_open_output(q: AgentOpenOutputQ):
    """打开一个既有交付物主题。只接受交付物根目录下的真实一级目录，不创建目录、不跟随链接。"""
    import agent_ws as AW
    try:
        AW.ensure_scaffold()
        base = Path(AW.output_dir()).resolve(strict=True)
        # 不直接用 base / 用户输入作为最终目标：只从真实一级目录中做精确名称匹配，
        # 因而绝对路径、..、混合分隔符都不可能越过交付物根目录。
        target = next((p for p in base.iterdir() if p.name == q.name), None)
        if target is None:
            raise HTTPException(status_code=404, detail="交付物主题不存在")
        if _is_link_or_junction(target) or not target.is_dir():
            raise HTTPException(status_code=400, detail="只能打开交付物内的真实主题目录")
        resolved = target.resolve(strict=True)
        if resolved.parent != base:
            raise HTTPException(status_code=400, detail="交付物主题不在允许范围内")
        _open_system_dir(resolved)
        return {"ok": True, "dir": str(resolved)}
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse({"ok": False, "msg": str(e)}, status_code=400)

def _parse_frontmatter(txt):
    """极简 YAML front-matter 解析：首个 --- 到下一个 --- 之间的 key: value（中英文冒号皆可）。
       返回 (meta_dict, body)。值里含冒号(如时间 08:30)不受影响（只按首个冒号切）。"""
    meta = {}
    body = txt
    t = txt.lstrip("﻿").lstrip()
    if t.startswith("---"):
        rest = t[3:]
        idx = rest.find("\n---")
        if idx != -1:
            fm = rest[:idx]; body = rest[idx + 4:]
            for line in fm.splitlines():
                s = line.replace("：", ":")
                if ":" in s:
                    k, v = s.split(":", 1)
                    meta[k.strip()] = v.strip()
    return meta, body

@app.get("/agent/tasks")
def agent_tasks():
    """扫「资料库/定时任务/*/任务.md」，解析 front-matter 展示定时任务列表。
       本端只登记/展示——定时执行由用户的 AI 助手在它自己的日程里负责（应用不联网、无大模型）。"""
    import agent_ws as AW
    out, unrecognized = [], []
    try:
        AW.ensure_scaffold()
        tdir = Path(AW.tasks_dir())
        for sub in sorted(tdir.iterdir()):
            if _is_link_or_junction(sub) or not sub.is_dir():
                continue
            f = sub / "任务.md"
            if not f.exists():
                unrecognized.append({"name": sub.name, "reason": "missing_task_file"})
                continue
            try:
                meta, bodytext = _parse_frontmatter(f.read_text(encoding="utf-8"))
            except Exception as e:
                unrecognized.append({"name": sub.name, "reason": "read_error", "detail": str(e)})
                continue
            name = meta.get("名称") or meta.get("name") or sub.name
            freq = meta.get("频率") or meta.get("freq") or meta.get("frequency") or ""
            # C1：enabled 缺省不再算「启用」——没写「启用」字段的是草稿，不该显示成绿灯正常运行。
            has_en = ("启用" in meta) or ("enabled" in meta)
            en = str(meta.get("启用", meta.get("enabled", ""))).strip().lower()
            enabled = en in ("true", "1", "yes", "on", "是", "启用")
            # C1：可观测字段——agent 每次跑完回写「上次执行」，前端据此对「启用但从未执行」的任务给出中性提示，
            # 缓解「应用显示已启用 vs 实际是否在跑毫无关系」的误导（记忆里踩过的坑）。
            last_run = meta.get("上次执行") or meta.get("last_run") or ""
            scheduler = meta.get("调度器") or meta.get("scheduler") or ""
            desc = ""
            for ln in bodytext.splitlines():
                ln = ln.strip()
                if not ln:
                    continue
                desc = (ln.replace("：", ":").split(":", 1)[-1].strip()
                        if ("搜什么" in ln or "内容" in ln) else ln)
                break
            out.append({"name": name, "freq": freq, "enabled": enabled, "has_enabled": has_en,
                        "last_run": last_run, "scheduler": scheduler,
                        "desc": desc[:160], "dir": str(sub)})
    except Exception as e:
        return {"tasks": out, "unrecognized": unrecognized,
                "unrecognized_count": len(unrecognized), "error": str(e)}
    return {"tasks": out, "unrecognized": unrecognized,
            "unrecognized_count": len(unrecognized)}


def _scan_agent_output_tree(root: Path):
    """递归统计真实目录中的文件/子目录与最新 mtime；逐目录容错且绝不跟随链接。"""
    file_count = subdir_count = scan_errors = 0
    try:
        latest_mtime = root.stat().st_mtime
    except OSError:
        latest_mtime = 0.0
        scan_errors += 1
    pending = [root]
    while pending:
        current = pending.pop()
        try:
            entries = list(current.iterdir())
        except OSError:
            scan_errors += 1
            continue
        for entry in entries:
            try:
                if _is_link_or_junction(entry):
                    continue
                stat = entry.stat(follow_symlinks=False)
                latest_mtime = max(latest_mtime, stat.st_mtime)
                if entry.is_dir():
                    subdir_count += 1
                    pending.append(entry)
                elif entry.is_file():
                    file_count += 1
            except OSError:
                scan_errors += 1
    return {"file_count": file_count, "subdir_count": subdir_count,
            "latest_mtime": latest_mtime, "scan_errors": scan_errors}

@app.get("/agent/outputs")
def agent_outputs(limit: int = 8):
    """C4：列 0_Agent交付物/ 下的主题子文件夹（最近修改在前），供 Agent 页交付物卡展示「最近做了哪些主题」。
       递归统计文件数/子目录数/最新修改时间；逐目录容错，不读取文件内容、不跟随链接。"""
    import agent_ws as AW
    out = []
    try:
        AW.ensure_scaffold()
        odir = Path(AW.output_dir())
        rows = []
        for d in odir.iterdir():
            try:
                if _is_link_or_junction(d) or not d.is_dir():
                    continue
                stats = _scan_agent_output_tree(d)
                rows.append((d, stats))
            except OSError:
                continue
        rows.sort(key=lambda row: row[1]["latest_mtime"], reverse=True)
        for d, stats in rows[:max(1, min(50, limit))]:
            latest = stats["latest_mtime"]
            mt = time.strftime("%Y-%m-%d", time.localtime(latest)) if latest else ""
            out.append({"name": d.name, "mtime": mt,
                        "latest_mtime": latest,
                        "file_count": stats["file_count"],
                        "subdir_count": stats["subdir_count"],
                        "scan_errors": stats["scan_errors"],
                        # n_files 暂留作旧前端兼容；新前端使用 file_count。
                        "n_files": stats["file_count"],
                        "has_readme": (d / "README.md").is_file(),
                        "dir": str(d)})
    except Exception as e:
        return {"outputs": out, "error": str(e)}
    return {"outputs": out}

@app.post("/setup/pick_folder")
def setup_pick_folder():
    """pywebview 原生文件夹选择对话框。无 pywebview（浏览器回退）时返回 ok=False，前端隐藏「浏览…」。"""
    try:
        import webview
        w = webview.windows[0] if getattr(webview, "windows", None) else None
        if not w:
            return {"ok": False, "msg": "无原生窗口"}
        res = w.create_file_dialog(webview.FOLDER_DIALOG)
        d = (res[0] if isinstance(res, (list, tuple)) and res else res) or ""
        return {"ok": bool(d), "dir": str(d)}
    except Exception as e:
        return {"ok": False, "msg": str(e)}

class BackendQ(BaseModel):
    backend: str = "local"                       # local | api
    base: Optional[str] = None
    key: Optional[str] = None
    embed_model: Optional[str] = None
    rerank_model: Optional[str] = None

def _reset_vectors_for_reembed():
    """换引擎：删掉旧引擎建的全部向量与嵌入进度，让下次 build 用新引擎【全量】重嵌。
       保留 extracted/chunks 产物（深索只需重嵌、不必重抽）。wiki 行会在 load_all 时由
       reindex_missing_pages 自动回灌。返回 True=已重置。"""
    try:
        for f in (C.META_EMBEDDED, C.STATE / "embedded_keys.txt"):
            try:
                if f.exists():
                    f.unlink()
            except Exception as e:
                log_error("reembed unlink", repr(e))
        import lancedb
        db = lancedb.connect(str(C.LANCEDB_DIR))
        if C.TABLE_NAME in db.table_names():
            db.drop_table(C.TABLE_NAME)   # meta+chunk+wiki 行全删，下次 index_semantic 以 overwrite 重建
        R.STATE["ready"] = False          # 表已删，检索先返回未就绪，别拿到半截结果
        try:                              # 清 manifest.backend（旧值已失效，等 index_semantic 写新值）
            if C.INDEX_MANIFEST.exists():
                man = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8"))
                man.pop("backend", None)
                C.INDEX_MANIFEST.write_text(json.dumps(man, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass
        return True
    except Exception as e:
        log_error("reset vectors for reembed", repr(e), traceback.format_exc())
        return False


def _reset_index_full():
    """『清空并重建』：把全部索引产物移到一个可恢复的 stash + drop 向量表，让下次「更新知识库」
       从头重建。重建时题录会重新**去重**（合并「已深索、在现有库上删不掉」的重复副本）、并按
       「只导 PDF」严格重新导入。
       **只清索引产物**，明确**保留**：综述 wiki / 收藏夹分类 / 期刊分级 / 检索摘要 SAC /
       页码映射 pagemap / 设置 / 0_Agent 工作区（人写的、花过 API 钱的、或重解析 PDF 很贵的）。
       stash 里放的是纯文件产物（chunks/extracted/bm25/state/manifest），重建确认无误后用户可删；
       向量表是 drop（不入 stash，本就要弃）——所以调用前 UI 会强烈建议先备份。
       返回 {ok, stash, moved, dropped_table, failed}。调用方须先确认没有在建索引（占 running 锁）。"""
    import shutil
    home = C.DATA.parent
    stamp = time.strftime("%Y%m%d-%H%M%S")
    stash = home / f"_reset_index_backup_{stamp}"
    # 先置未就绪 + 释放内存里的索引句柄，免得 Windows 下移动/删除被占用（bm25 可能 mmap、表句柄占 lancedb）
    R.STATE["ready"] = False
    for k in ("tbl", "row_count", "records", "bm25", "bm25_ids", "meta_bm25", "meta_ids"):
        R.M.pop(k, None)
    try:
        import gc; gc.collect()
    except Exception:
        pass
    # ① 把纯文件索引产物移到 stash（可恢复）。只移这些——绝不碰 wiki/categories/summaries/pagemap/meta/settings
    moved, failed = [], []
    for src, name in [(C.CHUNKS, "chunks"), (C.EXTRACTED, "extracted"),
                      (C.BM25_DIR, "bm25"), (C.BM25_META_DIR, "bm25_meta"),
                      (C.STATE, "state"), (C.INDEX_MANIFEST, "index_manifest.json")]:
        if not src.exists():
            continue
        try:
            dst = stash / "data" / name
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src), str(dst))
            moved.append(name)
        except Exception as e:
            failed.append(f"{name}: {e}")
            log_error("reset move", f"{name}: {e!r}", traceback.format_exc())
    # ② drop 向量表（LanceDB API 自己处理句柄；表在重建的 semantic 阶段以 mode=overwrite 重造，
    #    wiki 行由 load_all→reindex_missing_pages 从磁盘 markdown 回灌）
    dropped = False
    try:
        import lancedb
        db = lancedb.connect(str(C.LANCEDB_DIR))
        if C.TABLE_NAME in db.table_names():
            db.drop_table(C.TABLE_NAME)
            dropped = True
    except Exception as e:
        failed.append(f"drop_table: {e}")
        log_error("reset drop_table", repr(e), traceback.format_exc())
    # ③ 重建空目录（server 进程不会再跑 config 的 mkdir；留空目录免得后续写入撞到缺目录）
    for d in (C.EXTRACTED, C.CHUNKS, C.BM25_DIR, C.BM25_META_DIR, C.STATE, C.LANCEDB_DIR):
        try:
            d.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
    # ④ 刷新内存索引状态 → 未建库/待重建（表没了、bm25_meta 没了 → load_all 落到「未建库」）
    try:
        R.load_all()
    except Exception as e:
        log_error("reset reload", repr(e), traceback.format_exc())
    return {"ok": not failed, "stash": str(stash), "moved": moved,
            "dropped_table": dropped, "failed": failed}


class ResetIndexQ(BaseModel):
    confirm: bool = False

@app.post("/index/reset")
def index_reset(q: ResetIndexQ = None):
    """『清空并重建』索引：清索引产物（保留 wiki/分类/SAC/分级/设置/Agent 工作区），供从头重建。
       destructive——须 confirm=true。忙时拒绝。见 _reset_index_full。"""
    if not (q and q.confirm):
        return {"ok": False, "need_confirm": True}
    with _BUILD_LOCK:
        if BUILD["running"]:
            return {"ok": False, "busy": True, "msg": "正在建索引，请等它结束或先停止，再清空。"}
        BUILD["running"] = True; BUILD["stage"] = "reset"   # 占位防清空途中有 build 抢锁
    try:
        r = _reset_index_full()
    finally:
        with _BUILD_LOCK:
            BUILD["running"] = False; BUILD["stage"] = ""
    if r.get("ok"):
        r["msg"] = ("索引已清空（旧的提取/切块/向量已移到 " + r["stash"] + "，确认重建无误后可自行删除）。"
                    "下一步：① 点顶栏『⟳ 更新知识库』重建题录+语义层（会自动去重、按只导 PDF 重新导入）；"
                    "② 完成后到首页或『浏览』点『深索全部』重新深索。")
    else:
        r["msg"] = "清空未完全成功（详见 failed）。可重启应用后重试；已移走的产物在 " + r.get("stash", "") + "。"
    return r

@app.post("/setup/backend")
def setup_backend(q: BackendQ):
    """保存检索引擎后端选择（本地/API）。API 模式存 SiliconFlow 等的 key。
       引擎（后端 或 嵌入模型）变化且已有向量时：清进度+删表，让随后的重建用新引擎【全量】重嵌，
       避免新旧两套向量在同一张表里混用、dense 召回静默劣化（此前只发一句 warn、实际增量根本不重嵌）。"""
    import settings as S
    old = S.load()
    old_backend = old.get("backend")
    old_embed = (old.get("api") or {}).get("embed_model")
    patch = {"backend": "api" if q.backend == "api" else "local"}
    api = {}
    if q.base: api["base"] = q.base
    if q.key is not None: api["key"] = q.key
    if q.embed_model: api["embed_model"] = q.embed_model
    if q.rerank_model: api["rerank_model"] = q.rerank_model
    if api: patch["api"] = api
    st = S.save(patch)
    new_backend = st.get("backend")
    new_embed = (st.get("api") or {}).get("embed_model")
    # 引擎变化 = 后端切换，或 API 模式下换了嵌入模型（维度可能不同，查询向量与表维度不符会直接报错）
    engine_changed = (old_backend != new_backend) or (new_backend == "api" and old_embed and old_embed != new_embed)
    has_vectors = C.META_EMBEDDED.exists() or R.STATE.get("mode") == "full"
    reembed, had_deep, warn = False, 0, None
    if engine_changed and has_vectors:
        try:
            had_deep = len(_deep_keys())
        except Exception:
            had_deep = 0
        reembed = _reset_vectors_for_reembed()
        warn = "检索引擎已切换：旧向量已清除，请重建索引以用新引擎全量重嵌。"
    return {"ok": True, "backend": st.get("backend"), "reembed": reembed, "had_deep": had_deep,
            "api_key_set": bool((st.get("api") or {}).get("key")), "warn": warn}

@app.post("/setup/test_api")
def setup_test_api(q: BackendQ):
    """用一小段文本测试 API：嵌入 + 重排是否可用，返回向量维度/延迟。"""
    import time as _t
    from embedder import APIEmbedder
    from reranker import APIReranker
    import settings as S
    a = S.api_conf()
    base = q.base or a.get("base"); key = q.key if q.key is not None else a.get("key", "")
    em = q.embed_model or a.get("embed_model", "BAAI/bge-m3")
    rm = q.rerank_model or a.get("rerank_model", "BAAI/bge-reranker-v2-m3")
    if not key:
        # 诊断接口：测试“失败”不是 HTTP 错误，返回 200 + ok:false，让前端 jpost 能读到 msg
        # （否则前端只会显示无意义的“/setup/test_api 400”，真实原因被吞）。
        return JSONResponse({"ok": False, "msg": "未填 API key"}, status_code=200)
    try:
        t0 = _t.time()
        v = APIEmbedder(base, key, em).encode(["测试：涉罪未成年人分流转处"])
        sc = APIReranker(base, key, rm).scores("未成年人分流", ["涉罪未成年人分流转处", "无关文本"])
        return {"ok": True, "dim": int(v.shape[1]), "rerank_ok": len(sc) == 2,
                "latency_ms": int((_t.time() - t0) * 1000)}
    except Exception as e:
        return JSONResponse({"ok": False, "msg": f"API 测试失败：{e}"}, status_code=200)

class SacQ(BaseModel):
    enabled: Optional[bool] = None
    generator: Optional[str] = None       # K2：server（服务端用API Key）| agent（交给Agent）| off（不生成）
    base: Optional[str] = None
    key: Optional[str] = None
    model: Optional[str] = None

@app.get("/setup/sac")
def get_sac():
    import settings as S, sac as SAC
    sc = S.sac_conf()
    return {"enabled": bool(sc.get("enabled")), "generator": sc.get("generator"),
            "base": sc.get("base"), "model": sc.get("model"),
            "key_set": bool(sc.get("key")), "key_last4": _last4(sc.get("key")),   # K3 掩码
            "effective_ready": SAC.enabled()}

@app.post("/setup/sac")
def setup_sac(q: SacQ):
    """配置深索摘要生成方（K2 generator 三选一）。key 空时会自动复用 API 后端的 key。
       generator=server→服务端自动生成；agent→交给 Agent（服务端不生成）；off→不生成。"""
    import settings as S, sac as SAC
    patch = {}
    if q.generator in ("server", "agent", "off"):
        patch["generator"] = q.generator
        patch["enabled"] = (q.generator == "server")     # 与旧 enabled 同步（向后兼容读 enabled 的代码）
    elif q.enabled is not None:                           # 老前端只传 enabled 时的兼容路径
        patch["enabled"] = bool(q.enabled)
        patch["generator"] = "server" if q.enabled else "off"
    # BF契约3：判空一律用 is not None——空字符串=用户清空该项要落盘；
    # 旧的 if q.base 会把"清空 base/model"静默吞掉（前端 base/model 每次都发，key 仅 dirty 才发）。
    if q.base is not None: patch["base"] = q.base
    if q.key is not None: patch["key"] = q.key
    if q.model is not None: patch["model"] = q.model
    S.save({"sac": patch})
    sc = S.sac_conf()
    return {"ok": True, "generator": sc.get("generator"), "effective_ready": SAC.enabled()}

# ── 期刊分级学科（整库锁定单学科；journal_grading 期刊权重引擎用）──
class DiscQ(BaseModel):
    discipline: Optional[str] = None

@app.get("/setup/discipline")
def get_discipline():
    """当前锁定学科 + 可选学科清单（供设置面板下拉）。personal=True 的是个人档（如 law_personal）。"""
    import settings as S
    cur = S.discipline()
    items = []
    try:
        import journal_grading as JG
        for did, meta in JG.load_data().disciplines().items():
            items.append({"id": did, "name": meta.get("name", did),
                          "personal": did.endswith("_personal")})
    except Exception:
        items = [{"id": cur, "name": cur, "personal": cur.endswith("_personal")}]
    return {"current": cur, "disciplines": items}

@app.post("/setup/discipline")
def setup_discipline(q: DiscQ):
    """切换整库锁定学科。检索期动态读设置——保存后下次检索即生效，无需重建索引。
       并后台预热新学科的分级分布缓存，让库总览/浏览稍后刷新即见新口径（不阻塞本请求）。"""
    import settings as S
    if q.discipline:
        S.save({"journal_discipline": q.discipline})
        try:
            import grading_svc as GS
            GS.warm_async(_load_papers())
        except Exception as e:
            log_error("discipline warm", repr(e))
    return {"ok": True, "current": S.discipline()}

@app.post("/setup/reset")
def setup_reset():
    """设置页「恢复默认」：settings.json 覆盖为默认（清 API/SAC key、学科回标准法学、后端回本地、
       检索组件空闲释放回到 10 分钟）。
       浏览器里的对话 LLM key 存 localStorage，由前端另清。"""
    import settings as S
    st = S.reset()
    R.release_retrieval_if_idle(0, force=True)
    return {"ok": True, "backend": st.get("backend"), "discipline": st.get("journal_discipline")}

@app.get("/setup/models_status")
def setup_models_status():
    """本地模型是否齐全（供本地模式的下载步骤判断）。"""
    try:
        import models_bootstrap as MB
        miss = MB.missing_models()
        return {"present": not miss, "missing": miss}
    except Exception as e:
        log_error("setup/models_status", repr(e))
        return {"present": False, "missing": ["bge-m3-onnx", "bge-reranker-v2-m3-onnx"], "err": str(e)}

@app.post("/setup/download_models")
def setup_download_models():
    """本地模式：从云端下载缺失模型，SSE 汇报进度。"""
    import models_bootstrap as MB
    def gen():
        q = []
        def cb(name, done, total, phase):
            q.append({"name": name, "done": done, "total": total, "phase": phase})
        import threading
        state = {"done": False, "ok": False, "msg": ""}
        def work():
            try:
                ok, msg = MB.ensure_models(progress_cb=cb, log=lambda *a: None)
                state["ok"], state["msg"] = ok, msg
            except Exception as e:
                state["ok"], state["msg"] = False, str(e)
                log_error("download_models", repr(e))
            state["done"] = True
        th = threading.Thread(target=work, daemon=True); th.start()
        while not state["done"] or q:
            if q:
                yield "data: " + json.dumps(q.pop(0), ensure_ascii=False) + "\n\n"
            else:
                time.sleep(0.15)
                yield ": keepalive\n\n"
        yield "data: " + json.dumps({"final": True, "ok": state["ok"], "msg": state["msg"]}, ensure_ascii=False) + "\n\n"
        yield "data: [DONE]\n\n"
    return StreamingResponse(gen(), media_type="text/event-stream")

@app.get("/index/status")
def index_status():
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    # BF33：去重计数——标记文件出现重复行时 len(split()) 虚高，与 /health 的 _deep_keys()（set）口径不一
    ek = C.STATE / "embedded_keys.txt"
    deep_set = set(ek.read_text(encoding="utf-8").split()) if ek.exists() else set()
    deep = len(deep_set)
    # 深索摘要（SAC）覆盖：已深索里多少篇有检索摘要（同 safe_name(stem) 口径）
    summary_keys = _summary_keys()
    summary_issues = _summary_issues()
    sac_done = len(deep_set & summary_keys)
    sac_invalid = len(deep_set & set(summary_issues))
    sac_missing = max(0, deep - sac_done - sac_invalid)
    # 去重计数（与 deep/_deep_keys() 同口径）：标记文件出现重复行时 len(split()) 会虚高，
    # 前端据 meta_done<papers 判断语义层是否待完成，虚高到 ≥papers 会误隐藏「正在提升检索质量」提示。
    meta_done = len(set(C.META_EMBEDDED.read_text(encoding="utf-8").split())) if C.META_EMBEDDED.exists() else 0
    extract_counts = _deep_extract_counts()
    # 旧字段保留给旧前端：语义改为“当前无法进入深索的终态”，不再谎称全是扫描件。
    deep_no_text = sum(extract_counts.get(s, 0)
                       for s in ("missing_pdf", "invalid_pdf", "ocr_failed"))
    with _Q_LOCK:
        q_pending = len(QUEUE["pending"]); q_inflight = len(QUEUE["in_flight"])
    return {
        "mode": R.STATE.get("mode"), "ready": R.STATE.get("ready", False),
        "light_done": manifest.get("light_done", False), "source": manifest.get("source"),
        "papers": manifest.get("papers", 0), "with_pdf": manifest.get("with_pdf", 0),
        "meta_done": meta_done, "deep_done": deep, "deep_no_text": deep_no_text,
        "extract_status_counts": extract_counts,
        "ocr_pending": extract_counts.get("ocr_pending", 0),
        "ocr_failed": extract_counts.get("ocr_failed", 0),
        "missing_pdf": extract_counts.get("missing_pdf", 0),
        "invalid_pdf": extract_counts.get("invalid_pdf", 0),
        # 深索摘要（SAC）：sac_done=已深索且有摘要的篇数；sac_generator=生成方(off/agent/server)；
        # sac_backfill=补生成摘要后台任务的实时进度（前端据此显示进度/禁用重复触发）。
        "sac_done": sac_done, "sac_invalid": sac_invalid, "sac_missing": sac_missing,
        "sac_generator": _sac_generator(), "sac_backfill": dict(BACKFILL),
        "building": BUILD["running"], "stage": BUILD["stage"], "log": BUILD["log"][-40:],
        "queue_pending": q_pending, "queue_in_flight": q_inflight,
        # BF14：cancelled=本次构建是否被 /build/cancel 取消；rc=上次构建子进程退出码（未结束为 null）。
        # 前端据此把「取消后仍显示构建中/误报完成」纠正为真实终态。
        "cancelled": bool(BUILD.get("cancelled")), "rc": BUILD.get("rc"),
        "auto_last": AUTO.get("last_result"),   # 上次自动增量更新的结果（成功/失败/无），前端可展示
    }

@app.get("/stats")
def stats_ep():
    if not C.STATS_CACHE.exists():
        return JSONResponse({"error": "尚未建立索引"}, status_code=404)
    s = json.loads(C.STATS_CACHE.read_text(encoding="utf-8"))
    # 深索数以实时 embedded_keys.txt 为准：stats_cache 在建 L 档时算，深索后不重算→会偏旧
    ek = C.STATE / "embedded_keys.txt"
    if ek.exists() and isinstance(s.get("coverage"), dict):
        # BF33：去重计数——重复行会让「已深索 N 篇」虚高甚至超过总篇数
        dk = _deep_keys()
        s["coverage"]["deep_indexed"] = len(dk)
        s["coverage"]["sac_indexed"] = len(dk & _summary_keys())   # 已深索里有检索摘要的篇数
    # 最近入库补 has_pdf/deep（供前端三态深索按钮，F45/副本#13）
    try:
        pap = _load_papers(); deepk = _deep_keys()
        for r in s.get("recent", []):
            p = pap.get(r.get("key"))
            r["has_pdf"] = bool(p.get("has_pdf")) if p else False
            r["deep"] = is_deep(r.get("key"), deepk)
    except Exception as e:
        # BF18：静默 pass 会把 papers.jsonl 损坏等真故障吞成「最近入库按钮不对」，至少记一笔
        log_error("stats recent enrich", repr(e))
    # F38-B：期刊分级分布按当前锁定学科现算（命中缓存则替换 by_tier/by_journal；
    # 未命中则后台预热、本次先返回建库时的旧分布并标 grading_pending，前端稍后刷新）。
    try:
        import grading_svc as GS
        import settings as _S
        s["grading_discipline"] = _S.discipline()
        try:                                          # F3：概览显示中文学科名
            import journal_grading as JG
            s["grading_discipline_name"] = JG.load_data().disciplines().get(
                _S.discipline(), {}).get("name", _S.discipline())
        except Exception:
            s["grading_discipline_name"] = _S.discipline()
        dist = GS.weight_dist(_load_papers())
        if dist:
            s["by_tier"], s["by_journal"] = dist
            s["grading_pending"] = False
        else:
            s["grading_pending"] = True
    except Exception as e:
        log_error("stats grading", repr(e))
    # EN-W1：附 wiki 体检摘要（契约3：{"issues":int,"checked_at":str}）。
    # 读 data/state/wiki_lint.json（由自动更新循环保鲜）；文件不存在/损坏则不带该键。
    try:
        lf = C.STATE / "wiki_lint.json"
        if lf.exists():
            d = json.loads(lf.read_text(encoding="utf-8"))
            s["wiki_lint"] = {"issues": int(d.get("issues", 0)),
                              "checked_at": str(d.get("checked_at", ""))}
    except Exception as e:
        log_error("stats wiki_lint", repr(e))
    return s

# ── 浏览：收藏夹 + 推荐「值得读」──────────────────────────
_PC = {"data": None, "mtime": 0}
def _load_papers():
    if not C.PAPERS_JSONL.exists():
        return {}
    mt = C.PAPERS_JSONL.stat().st_mtime
    if _PC["data"] is None or _PC["mtime"] != mt:
        d = {}
        for line in open(C.PAPERS_JSONL, encoding="utf-8"):
            if line.strip():
                p = json.loads(line); d[p["key"]] = p
        _PC["data"] = d; _PC["mtime"] = mt
    return _PC["data"]

def _load_cats():
    f = C.CATEGORIES_DIR / "zotero_collections.json"
    return json.loads(f.read_text(encoding="utf-8")) if f.exists() else {"tree": [], "by_collection": {}, "by_key": {}}

def _deep_keys():
    ek = C.STATE / "embedded_keys.txt"
    return set(ek.read_text(encoding="utf-8").split()) if ek.exists() else set()

def _deep_extract_items():
    """逐篇 PDF 提取状态。读失败时退空，不影响检索主链。"""
    try:
        import deep_extract_status as DES
        return DES.load_items()
    except Exception:
        return {}


def _deep_extract_counts(items=None):
    try:
        import deep_extract_status as DES
        return DES.counts(items if items is not None else _deep_extract_items())
    except Exception:
        return {}


def _extract_record(stem, items=None, legacy_deep=False):
    """返回一篇的有效提取状态，并兼容 OCR 状态文件出现前的深索记录。

    旧版提取器还没有本地 OCR：能成功进入 ``embedded_keys.txt`` 的正文只能来自
    PDF 原生文字层。因此，已深索但 sidecar 无记录的旧文献应视为 ``ok_native``，
    不能在“原生文字层”筛选中凭空消失。
    """
    rec = (items if items is not None else _deep_extract_items()).get(str(stem), {})
    if isinstance(rec, dict) and rec.get("status"):
        return dict(rec)
    if legacy_deep:
        return {"status": "ok_native", "legacy_inferred": True}
    return {}


def _browse_filter_matches(name, isdeep, extract_rec, has_summary, summary_invalid):
    """浏览页状态筛选的单一口径；下拉计数和实际过滤必须共用。"""
    if not name:
        return True
    if name == "yes":
        return isdeep
    if name == "no":
        return not isdeep
    if name == "ocr":
        return extract_rec.get("status") == "ok_ocr"
    if name == "native":
        return extract_rec.get("status") == "ok_native"
    if name == "summary_yes":
        return has_summary
    if name == "summary_invalid":
        return summary_invalid
    if name == "summary_no":
        return isdeep and not has_summary and not summary_invalid
    return True


def _deep_no_text_keys():
    """兼容旧调用名：返回当前不可自动继续深索的提取终态。

    ocr_pending 不在这里——它必须能重新进入 extract，自动执行本地 OCR。
    """
    items = _deep_extract_items()
    blocked = {stem for stem, rec in items.items()
               if rec.get("status") in ("missing_pdf", "invalid_pdf", "ocr_failed")}
    # 尚未迁移的极端情况仍尊重旧标记；正常启动会先 reconcile_legacy。
    nt = C.STATE / "deep_no_text.txt"
    if nt.exists() and not items:
        blocked.update(nt.read_text(encoding="utf-8").split())
    return blocked

def _summary_keys():
    """已生成且通过质量检查的检索摘要 stem 集合，与 embedded_keys.txt 同口径。
       用来统计「已深索里多少篇有检索摘要」= deep ∩ summary。"""
    try:
        import sac as _SAC
        return _SAC.summary_keys()
    except Exception:
        return set()

def _summary_issues():
    """异常检索摘要 {stem: reason}。只读，不在后台擅自重生成。"""
    try:
        import sac as _SAC
        return _SAC.summary_issues()
    except Exception:
        return {}

def _sac_generator():
    """当前深索摘要生成方（off / agent / server），供前端措辞与补生成入口判断。"""
    try:
        import settings as S
        return S.sac_conf().get("generator")
    except Exception:
        return None

def _is_siliconflow_base(base):
    """BF8 加固：判断 base_url 是否**真的**指向 SiliconFlow（按 host 精确匹配，而非子串包含）。
       旧写法 "siliconflow" in base 会把 evilsiliconflow.cn 之类的后缀伪造 base 也放行，
       导致用户填入恶意 base 时把回填的 SiliconFlow key 外发。"""
    try:
        from urllib.parse import urlsplit
        h = (urlsplit(base or "").hostname or "").lower()
    except Exception:
        return False
    return h in ("siliconflow.cn", "siliconflow.com") or h.endswith(".siliconflow.cn") or h.endswith(".siliconflow.com")


def _dedup_deep_marks():
    """BF34：深索标记文件是追加写，批次重试/进程中断会留下重复行——启动时保序去重、
       tmp+os.replace 原子重写自愈（与 BF33 的读侧去重配套，治本在此）。"""
    for f in (C.STATE / "embedded_keys.txt", C.STATE / "deep_no_text.txt", C.META_EMBEDDED):
        try:
            if not f.exists():
                continue
            lines = f.read_text(encoding="utf-8").split()
            uniq = list(dict.fromkeys(lines))          # 保序去重
            if len(uniq) < len(lines):
                tmp = f.with_name(f.name + ".tmp")
                tmp.write_text("\n".join(uniq) + "\n", encoding="utf-8")
                os.replace(tmp, f)
                print(f"[server] 已清理 {len(lines) - len(uniq)} 条重复深索标记（{f.name}）", flush=True)
        except Exception as e:
            log_error("dedup deep marks", repr(e))

def _rec_score(p, g=None):
    """值得读打分：期刊权重为主 + 新近度 + 有 PDF（可深读）。
       学科感知权重(g)优先（0–1→0–10），缺则回退旧离散 tier_rank（峰值同为 ~12）。"""
    if g and g.get("weight") is not None:
        s = g["weight"] * 10.0
    else:
        s = (6 - p.get("tier_rank", 6)) * 2.0
    try:
        yr = int(p.get("year") or 0)
    except Exception:
        yr = 0
    if yr >= 2015:
        s += min(3.0, (yr - 2015) * 0.3)
    if p.get("has_pdf"):
        s += 1.0
    return round(s, 2)

def _load_ai_topics():
    f = C.CATEGORIES_DIR / "ai_topics.json"
    return json.loads(f.read_text(encoding="utf-8")) if f.exists() else {"topics": [], "by_key": {}}

# ══════════════════════════════════════════════════════════════════
#  F10/F11：自建「知识库分类」 + 分类→keys 解析 + 自动深索串行队列
# ══════════════════════════════════════════════════════════════════
import uuid
import textutil as T

def _atomic_json_write(path, obj, indent=None):
    """原子写 JSON，对 OneDrive/杀软临时占用导致的 os.replace WinError 5 做重试 + 兜底直写。
       生产环境 LOCALKB_DATA 指向 %LOCALAPPDATA%（非 OneDrive），此重试主要保开发/便携场景稳。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    data = json.dumps(obj, ensure_ascii=False, indent=indent)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(data, encoding="utf-8")
    last = None
    for i in range(6):
        try:
            os.replace(tmp, path)
            return
        except PermissionError as e:
            last = e
            time.sleep(0.15 * (i + 1))
    # 重试仍失败：退化为直写目标（放弃原子性，但保证不丢数据）
    try:
        path.write_text(data, encoding="utf-8")
        try:
            if tmp.exists():
                tmp.unlink()
        except Exception:
            pass
    except Exception:
        raise last

def is_deep(key, deepk=None):
    """唯一的深索判定入口：分类文件存原始 key，embedded_keys.txt 存 safe_name(stem)。"""
    if deepk is None:
        deepk = _deep_keys()
    return T.safe_name(key) in deepk

_KBC_FILE = C.CATEGORIES_DIR / "kb_categories.json"
_KBC_LOCK = threading.Lock()
_KBC = {"data": None, "mtime": 0}

def _kbc_load():
    if not _KBC_FILE.exists():
        return {"version": 1, "categories": [], "updated_at": ""}
    mt = _KBC_FILE.stat().st_mtime
    if _KBC["data"] is None or _KBC["mtime"] != mt:
        _KBC["data"] = json.loads(_KBC_FILE.read_text(encoding="utf-8"))
        _KBC["mtime"] = mt
    return _KBC["data"]

def _kbc_save(doc):
    """原子写（对 OneDrive 占用重试）。调用方须持 _KBC_LOCK。"""
    doc["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
    _atomic_json_write(_KBC_FILE, doc, indent=1)
    _KBC["data"] = doc; _KBC["mtime"] = _KBC_FILE.stat().st_mtime

def _kbc_find(doc, cid):
    for c in doc["categories"]:
        if c["id"] == cid:
            return c
    return None

# ── 自动深索串行队列 ──────────────────────────────────────
_DEEP_QUEUE_FILE = C.STATE / "deep_queue.json"
_Q_LOCK = threading.Lock()
# K1：paused=深索暂停标志（不再起新批）；spp=近批「秒/篇」用于 ETA 外推；batch_start=当前批起跑时刻。
QUEUE = {"pending": [], "in_flight": [], "timer": None, "fails": 0,
         "paused": False, "spp": None, "batch_start": None}
_DEEP_BATCH = 50
_DEEP_DEBOUNCE = 4.0
_DEEP_MAX_RETRY = 3        # A3：同一批连续软失败超过此次数则暂停自动重试（避免坏 PDF/余额0 无限重试烧钱）

def _q_persist():
    """调用方须持 _Q_LOCK。"""
    try:
        _atomic_json_write(_DEEP_QUEUE_FILE, {
            "pending": QUEUE["pending"], "in_flight": QUEUE["in_flight"],
            "paused": bool(QUEUE.get("paused")),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S")})
    except Exception as e:
        log_error("deep_queue persist", repr(e))

def _q_boot():
    """startup 调用：回灌上次崩溃残留（in_flight 可能未跑完 → 回 pending，extract 幂等跳过已完成）。"""
    if _DEEP_QUEUE_FILE.exists():
        try:
            d = json.loads(_DEEP_QUEUE_FILE.read_text(encoding="utf-8"))
            QUEUE["pending"] = list(dict.fromkeys((d.get("in_flight") or []) + (d.get("pending") or [])))
            QUEUE["in_flight"] = []
            QUEUE["paused"] = bool(d.get("paused"))    # K1：恢复暂停标志（暂停态跨重启保留）
        except Exception:
            pass
    if QUEUE["pending"] and not QUEUE.get("paused"):
        _drain_deep_queue()

def _queue_keyset():
    with _Q_LOCK:
        return set(QUEUE["pending"]) | set(QUEUE["in_flight"])

def enqueue_deep(keys):
    """入队一批 raw key：剔除已深索/已在队/在跑。返回真正入队的。触发防抖排空。"""
    deepk = _deep_keys()
    with _Q_LOCK:
        have = set(QUEUE["pending"]) | set(QUEUE["in_flight"])
        newk = [k for k in keys if k not in have and not is_deep(k, deepk)]
        if newk:
            QUEUE["pending"].extend(newk); QUEUE["fails"] = 0; _q_persist()   # 新动作复位失败计数，重启自动重试
            if QUEUE["timer"]:
                QUEUE["timer"].cancel()
            QUEUE["timer"] = threading.Timer(_DEEP_DEBOUNCE, _drain_deep_queue)
            QUEUE["timer"].daemon = True; QUEUE["timer"].start()
    return newk

def _drain_deep_queue():
    """把一批 pending 转 in_flight 并起 deep build。撞锁则回滚，等 build 结束再排。"""
    with _Q_LOCK:
        QUEUE["timer"] = None
        # K1：暂停中不再起新批（正在跑的那批已在子进程里自然跑完，队列保留）。
        if BUILD["running"] or QUEUE.get("paused") or not QUEUE["pending"]:
            return
        batch = QUEUE["pending"][:_DEEP_BATCH]
        QUEUE["pending"] = QUEUE["pending"][_DEEP_BATCH:]
        QUEUE["in_flight"] = batch
        QUEUE["batch_start"] = time.time()   # K1：记本批起跑，用于 ETA「秒/篇」外推
        _q_persist()
    started = _run_build("deep", ["--scope", "keys:" + ",".join(batch)], on_done=_on_deep_done)
    if not started:
        with _Q_LOCK:
            QUEUE["pending"] = batch + QUEUE["pending"]
            QUEUE["in_flight"] = []; _q_persist()

def _on_deep_done(rc=0):
    """A3：深索子进程结束回调（收 returncode）。
       rc==0：本批已成功入库，清 in_flight、复位失败计数、继续排空后续。
       rc!=0：软失败——把 in_flight 退回 pending **队首**而非清空（否则该批 key 永久蒸发），
              在重试预算内自动重排；超预算则保留在 pending 但停自动重排，待人工/重启再试。"""
    should_drain = False
    with _Q_LOCK:
        batch = QUEUE["in_flight"]; QUEUE["in_flight"] = []
        # K1：成功批用「本批耗时/篇数」更新 seconds-per-paper 的 EMA，供 /index/queue 外推 ETA。
        if rc == 0 and batch and QUEUE.get("batch_start"):
            el = time.time() - QUEUE["batch_start"]
            spp = el / max(1, len(batch))
            prev = QUEUE.get("spp")
            QUEUE["spp"] = spp if not prev else (prev * 0.5 + spp * 0.5)
        QUEUE["batch_start"] = None
        if rc != 0 and batch:
            QUEUE["fails"] = QUEUE.get("fails", 0) + 1
            QUEUE["pending"] = batch + QUEUE["pending"]        # 退回队首，保住这批 key
            if QUEUE["fails"] <= _DEEP_MAX_RETRY:
                should_drain = True
                BUILD["log"].append(f"[deep] 本批深索失败(rc={rc})，已退回队列重试（第 {QUEUE['fails']} 次）。")
            else:
                BUILD["log"].append(f"[deep] 本批深索连续失败 {QUEUE['fails']} 次，暂停自动重试"
                                    f"（{len(batch)} 篇仍在队列，稍后可重新触发深索）。")
        else:
            QUEUE["fails"] = 0
            should_drain = bool(QUEUE["pending"])
        _q_persist()
    # EN-W2：Ingest 扳机——本批深索成功后，后台算「这批新文献影响了哪些综述页」。
    # 挂在锁外、daemon 线程里跑，绝不阻塞队列继续排空。
    if rc == 0 and batch:
        try:
            _wiki_suggest_async(batch)
        except Exception as e:
            log_error("wiki suggest (queue batch)", repr(e))
    if should_drain:
        _drain_deep_queue()

# ── EN-W2：Ingest 扳机——深索成功后自动算「新文献影响了哪些综述页」──────
# gist 三环里 Ingest 环此前为 0：新文献入库后 wiki 毫无反应，drift 就是这么积累的。
# 这里只**建议**不动手（复用 wiki_store.propose_updates：by_source 反查 + 检索同题页），
# 结果落 data/state/wiki_suggestions.json，前端/agent 拿去逐条处理或 dismiss。
_WIKI_SUGG_FILE = C.STATE / "wiki_suggestions.json"
_WIKI_SUGG_LOCK = threading.Lock()      # 读-改-写串行：深索回调与 dismiss 端点可能并发
_WIKI_SUGG_MAX = 50                     # 只保留最近 50 条（旧建议早该被处理或已过时）
_WIKI_SUGG_BATCH_CAP = 20               # 单批最多算 20 篇：整库深索后一批可能几百篇，每篇要跑
                                        # 一次检索（propose_updates 内含 R.search），全算会把
                                        # 后台线程拖住几十分钟——挑批次前 20 篇，其余等下批

def _wiki_sugg_load():
    """读建议清单（items 列表）。文件缺失/损坏退空表——建议是提示件，不值得为它报错。"""
    try:
        if _WIKI_SUGG_FILE.exists():
            d = json.loads(_WIKI_SUGG_FILE.read_text(encoding="utf-8"))
            if isinstance(d, dict) and isinstance(d.get("items"), list):
                return d["items"]
    except Exception as e:
        log_error("wiki_suggestions load", repr(e))
    return []

def _wiki_sugg_save(items):
    """调用方须持 _WIKI_SUGG_LOCK。原子写（契约2 数据源）。"""
    _atomic_json_write(_WIKI_SUGG_FILE, {
        "items": items[:_WIKI_SUGG_MAX],
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S")}, indent=1)

def _wiki_suggest_batch(keys):
    """对一批新深索的 raw key 逐篇跑 propose_updates，pages 非空的追加进建议文件。
       每篇失败单独吞并记日志——一篇坏文献（题录缺失/检索异常）不能毁掉整批。"""
    _all = [k for k in dict.fromkeys(keys or []) if k]
    keys = _all[:_WIKI_SUGG_BATCH_CAP]
    if len(_all) > len(keys):
        # bulk 整库深索没有"下一批"来补算——如实留痕，别让注释谎称"其余等下批"
        BUILD["log"].append(f"[wiki建议] 本批仅分析前 {len(keys)} 篇，其余 {len(_all)-len(keys)} 篇未算建议"
                            f"（可对个别文献用 propose_wiki_updates 单独补算）。")
    if not keys:
        return
    papers = _load_papers()
    updates, newpages = [], []
    ts = time.strftime("%Y-%m-%d %H:%M:%S")
    for k in keys:
        try:
            r = W.propose_updates(k)
            title = (papers.get(k) or {}).get("title", "")
            pages = [{"id": a.get("id", ""), "title": a.get("title", "")}
                     for a in (r.get("affected") or []) if a.get("id")]
            if pages:
                updates.append({"key": k, "title": title, "pages": pages,
                                "kind": "update", "created_at": ts})
            else:
                # A2/A3：没命中既有页 ≠ 没事做——这是「该为它新建 concept/entity 页」（wiki 空时=建首页）
                # 的信号。propose_updates 已算出提示（note=冷启动 / hints=新主题），此前被 `if pages` 丢弃，
                # 导致 wiki 只会就地修补、长不出新主题。这里保留成 new_page 建议，供 agent 决定要不要建页。
                hint = (r.get("note") or (r.get("hints") or [""])[0]
                        or "库里还没有讲这篇主题的综述页，考虑为它新建 concept / entity 页。")
                newpages.append({"key": k, "title": title, "pages": [],
                                 "kind": "new_page", "hint": hint, "created_at": ts})
        except Exception as e:
            log_error(f"wiki suggest {k}", repr(e))     # 单篇失败不毁整批
    found = updates + newpages          # 「更新既有页」优先在前、「新建页」其次，同享 50 条上限不被挤掉
    if not found:
        return
    with _WIKI_SUGG_LOCK:
        have = {x.get("key") for x in found}
        # 新建议在前；同 key 旧建议被顶掉（按 key 去重——同一篇重复深索只留最新一条）
        items = found + [x for x in _wiki_sugg_load() if x.get("key") not in have]
        _wiki_sugg_save(items)

def _wiki_suggest_async(keys):
    """起 daemon 线程算建议（propose_updates 含检索调用，不能在深索回调线程里同步跑）。"""
    threading.Thread(target=_wiki_suggest_batch, args=(list(keys),), daemon=True).start()

@app.get("/index/queue")
def deep_queue_status():
    """K1：深索队列/详情/ETA。items=当前在深索或在队首的若干篇 key+标题；
       eta_seconds 按近批速率外推（取不到给 null）。"""
    papers = _load_papers()
    deep_set = _deep_keys()
    deep_done = len(deep_set)
    summary_keys = _summary_keys()
    summary_issues = _summary_issues()
    sac_done = len(deep_set & summary_keys)   # 已深索里有有效检索摘要的篇数
    sac_invalid = len(deep_set & set(summary_issues))
    sac_missing = max(0, deep_done - sac_done - sac_invalid)
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    with _Q_LOCK:
        pending = len(QUEUE["pending"]); in_flight = len(QUEUE["in_flight"])
        paused = bool(QUEUE.get("paused")); spp = QUEUE.get("spp")
        inflight_keys = list(QUEUE["in_flight"])
        shown = (inflight_keys + list(QUEUE["pending"]))[:8]
    items = [{"key": k, "title": (papers.get(k) or {}).get("title", "")} for k in shown]
    remaining = pending + in_flight
    eta_seconds = int(remaining * spp) if (spp and remaining) else None
    extract_counts = _deep_extract_counts()
    return {"pending": pending, "in_flight": in_flight, "paused": paused,
            "deep_done": deep_done, "with_pdf": manifest.get("with_pdf", 0),
            "sac_done": sac_done, "sac_invalid": sac_invalid, "sac_missing": sac_missing,
            "sac_backfill": dict(BACKFILL),   # 深索摘要覆盖 + 补生成进度
            "deep_no_text": len(_deep_no_text_keys()),   # 旧前端兼容：当前不可继续的提取终态数
            "extract_status_counts": extract_counts,
            "ocr_pending": extract_counts.get("ocr_pending", 0),
            "ocr_failed": extract_counts.get("ocr_failed", 0),
            "missing_pdf": extract_counts.get("missing_pdf", 0),
            "invalid_pdf": extract_counts.get("invalid_pdf", 0),
            "eta_seconds": eta_seconds, "items": items,
            # 兼容旧前端字段
            "in_flight_keys": inflight_keys,
            # BF35：bulk=当前 deep 构建是整库深索(scope=all)；队列批次为 false。
            # 前端据此区分「整库深索中」与「队列批次在跑」两种进度文案。
            "bulk": bool(BUILD.get("bulk")),
            "building": BUILD["running"], "stage": BUILD["stage"]}

class DeepPauseQ(BaseModel):
    paused: bool = True

@app.post("/index/deep/pause")
def deep_pause(q: DeepPauseQ):
    """K1：置深索暂停标志。暂停后 _drain_deep_queue 不再起新批（正在跑的自然跑完），队列保留；
       恢复(paused=False)立即尝试继续排空。"""
    with _Q_LOCK:
        QUEUE["paused"] = bool(q.paused)
        _q_persist()
    if not q.paused:
        try:
            _drain_deep_queue()
        except Exception as e:
            log_error("deep resume drain", repr(e))
    return {"ok": True, "paused": bool(q.paused)}

# ── F11：分类 → keys 白名单解析 ──────────────────────────
def _resolve_category_keys(category):
    """把 category 值解析成原始 key 白名单 set；None/"" → None(不过滤)。
       前缀：kbc_（用户分类）/ topic:<id>（AI主题）/ zotero:<path>（收藏夹镜像）。未知前缀→空集(命中0，安全)。"""
    if not category:
        return None
    if category.startswith("kbc_"):
        with _KBC_LOCK:
            doc = _kbc_load(); c = _kbc_find(doc, category)
        return set(c["keys"]) if c else set()
    if category.startswith("topic:"):
        try:
            tid = int(category.split(":", 1)[1])
        except Exception:
            return set()
        ait = _load_ai_topics()
        return {k for k, v in ait.get("by_key", {}).items() if v == tid}
    if category.startswith("zotero:"):
        path = category.split(":", 1)[1]
        return set(_load_cats().get("by_collection", {}).get(path, []))
    return set()

# ── F10：知识库分类 CRUD 端点 ────────────────────────────
@app.get("/kb/categories")
def kb_categories_list():
    papers = _load_papers(); deepk = _deep_keys(); qset = _queue_keyset()
    with _KBC_LOCK:
        doc = _kbc_load()
        cats = [dict(c) for c in doc["categories"]]
    out = []
    for c in cats:
        live = [k for k in c.get("keys", []) if k in papers]
        deep = sum(1 for k in live if is_deep(k, deepk))
        nopdf = sum(1 for k in live if not papers[k].get("has_pdf"))
        pend = sum(1 for k in live if k in qset)
        out.append({"id": c["id"], "name": c["name"], "source": c.get("source", "user"),
                    "count": len(live), "deep_count": deep, "no_pdf": nopdf,
                    "pending": pend, "updated_at": c.get("updated_at", "")})
    return {"categories": out}

class KbCatNewQ(BaseModel):
    name: str
    source: str = "user"

@app.post("/kb/categories")
def kb_cat_create(q: KbCatNewQ):
    name = (q.name or "").strip()
    if not name:
        return JSONResponse({"ok": False, "detail": "分类名不能为空"}, status_code=400)
    cid = "kbc_" + uuid.uuid4().hex[:8]
    now = time.strftime("%Y-%m-%d %H:%M:%S")
    with _KBC_LOCK:
        doc = _kbc_load()
        doc["categories"].append({"id": cid, "name": name, "source": q.source,
            "keys": [], "note": "", "created_at": now, "updated_at": now})
        _kbc_save(doc)
    return {"ok": True, "id": cid, "name": name}

class KbCatRenameQ(BaseModel):
    name: str

@app.patch("/kb/categories/{cid}")
def kb_cat_rename(cid: str, q: KbCatRenameQ):
    name = (q.name or "").strip()
    if not name:
        return JSONResponse({"ok": False, "detail": "分类名不能为空"}, status_code=400)
    with _KBC_LOCK:
        doc = _kbc_load(); c = _kbc_find(doc, cid)
        if not c:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        c["name"] = name; c["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _kbc_save(doc)
    return {"ok": True, "id": cid, "name": name}

@app.delete("/kb/categories/{cid}")
def kb_cat_delete(cid: str):
    with _KBC_LOCK:
        doc = _kbc_load()
        n = len(doc["categories"])
        doc["categories"] = [c for c in doc["categories"] if c["id"] != cid]
        if len(doc["categories"]) == n:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        _kbc_save(doc)
    return {"ok": True, "id": cid}

class KbMembersQ(BaseModel):
    keys: List[str] = []

@app.post("/kb/categories/{cid}/members")
def kb_cat_add_members(cid: str, q: KbMembersQ):
    papers = _load_papers(); deepk = _deep_keys()
    req = [k for k in dict.fromkeys(q.keys or []) if k]
    added, already = [], []
    with _KBC_LOCK:
        doc = _kbc_load(); c = _kbc_find(doc, cid)
        if not c:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        have = set(c["keys"])
        for k in req:
            (already if k in have else added).append(k)
        c["keys"].extend(added)
        c["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _kbc_save(doc)
    will_deep, no_pdf, already_deep = [], [], []
    for k in added:
        p = papers.get(k)
        if not p:
            continue
        if is_deep(k, deepk):
            already_deep.append(k)
        elif p.get("has_pdf"):
            will_deep.append(k)
        else:
            no_pdf.append(k)
    queued = enqueue_deep(will_deep) if will_deep else []
    return {"ok": True, "added": added, "already": already,
            "will_deep": queued, "no_pdf": no_pdf, "already_deep": already_deep,
            "queued": len(queued)}

@app.delete("/kb/categories/{cid}/members")
def kb_cat_del_members(cid: str, q: KbMembersQ):
    rm = set(q.keys or [])
    with _KBC_LOCK:
        doc = _kbc_load(); c = _kbc_find(doc, cid)
        if not c:
            return JSONResponse({"ok": False, "detail": "无此分类"}, status_code=404)
        before = len(c["keys"])
        c["keys"] = [k for k in c["keys"] if k not in rm]
        removed = before - len(c["keys"])
        c["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
        _kbc_save(doc)
    return {"ok": True, "removed": removed}

def _enrich_cat_deep(cats):
    """给收藏夹树每个节点补 count_indexed（该夹+子孙里已深索的篇数，实时算）。
       深索是后置手动动作，建库时算会偏旧——故检索期用 embedded_keys.txt 现算，
       让浏览左树显示「已深索/总数」双数字，而非满屏 0（旧版只显示总数、无深索维度）。"""
    import textutil as T
    deepk = _deep_keys()
    by_col = cats.get("by_collection", {})
    def walk(n):
        direct = by_col.get(n.get("path", ""), [])
        di = sum(1 for k in direct if T.safe_name(k) in deepk)
        for c in n.get("children", []):
            di += walk(c)
        n["count_indexed"] = di
        return di
    for r in cats.get("tree", []):
        walk(r)
    return cats

@app.get("/categories")
def categories():
    return _enrich_cat_deep(_load_cats())

@app.get("/topics")
def topics():
    ait = _load_ai_topics()
    deepk = _deep_keys()   # F10：每主题补已深索数，前端显示 已深索/总篇数
    out = []
    for t in ait.get("topics", []):
        keys = t.get("keys") or []
        deep = sum(1 for k in keys if is_deep(k, deepk))
        out.append({"id": t["id"], "name": t["name"], "size": t["size"], "deep": deep})
    return {"topics": out}

# ── AI 主题聚类：用向量 KMeans 把已索引文献自动归类（无 LLM，簇数自适应）──
TOPICS_BUILD = {"running": False, "msg": ""}
@app.post("/topics/rebuild")
def topics_rebuild():
    """后台重跑 AI 主题聚类（build_ai_topics.py，需已建语义/深索索引）。"""
    if TOPICS_BUILD["running"]:
        return {"ok": False, "msg": "AI 主题正在归类中，请稍候"}
    if BUILD["running"]:
        return {"ok": False, "msg": "正在建库/深索中，请等结束后再归类"}
    if R.STATE.get("mode") != "full":
        return JSONResponse({"ok": False, "msg": "需先深索或建语义层（AI 主题按向量聚类，纯题录库还没有向量）", "need_index": True}, status_code=200)
    def run():
        TOPICS_BUILD["running"] = True; TOPICS_BUILD["msg"] = "归类中…"
        try:
            env = dict(os.environ); env.pop("PYTHONUTF8", None)
            p = subprocess.run([sys.executable, str(C.APP / "build_ai_topics.py")],
                               stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env, timeout=900,
                               creationflags=C.SUBPROC_NO_WINDOW)
            TOPICS_BUILD["msg"] = "完成" if p.returncode == 0 else f"失败(code={p.returncode})"
        except Exception as e:
            log_error("topics/rebuild", repr(e)); TOPICS_BUILD["msg"] = f"异常：{e}"
        finally:
            TOPICS_BUILD["running"] = False
    threading.Thread(target=run, daemon=True).start()
    return {"ok": True}

@app.get("/topics/status")
def topics_status():
    return {"running": TOPICS_BUILD["running"], "msg": TOPICS_BUILD["msg"]}

# ── AI 抽词「找相似」：从标题抽实词做检索词（默认 AI，缺 key 退本地）──
class SimilarQ(BaseModel):
    title: str = ""
    provider: str = "siliconflow"
    base_url: str = ""
    api_key: str = ""
    model: str = ""

@app.post("/similar/keywords")
def similar_keywords(q: SimilarQ):
    """把一篇标题抽成 2-4 个检索关键词。优先 LLM（准），无 key 退本地实词。"""
    title = (q.title or "").strip()
    if not title:
        return {"ok": True, "keywords": ""}
    base, model = L.resolve(q.provider, q.base_url, q.model)
    key = q.api_key
    # BF8：key 外发守卫（仿 /chat）——检索引擎的 SiliconFlow key 只准回填给 SiliconFlow 的 base，
    # 否则用户任填一个第三方 base 就会把 key 发出去。无 key 时返回空 keywords，前端退本地分词。
    if not key and _is_siliconflow_base(base):
        try:
            import settings as S
            key = (S.api_conf() or {}).get("key", "")
        except Exception:
            key = ""
    if key:
        try:
            msgs = [{"role": "system", "content": "你是学术检索助手。把用户给的论文标题提炼成 2-4 个最核心的中文检索关键词，"
                     "用空格分隔，只输出关键词本身，不要解释、不要标点。"},
                    {"role": "user", "content": title}]
            kw = L.chat_once(msgs, base, key, model or "Qwen/Qwen2.5-7B-Instruct", temperature=0.2, timeout=30)
            kw = " ".join(re.findall(r"[一-鿿A-Za-z0-9]+", kw))[:60]
            if kw:
                return {"ok": True, "keywords": kw, "by": "ai"}
        except Exception as e:
            log_error("similar/keywords", repr(e))
    return {"ok": True, "keywords": "", "by": "none"}   # 前端据空回退本地抽词

@app.get("/similar/{key}")
def similar_vector(key: str, topk: int = 8):
    """C4/F6：真正的向量「找相似」。full 模式用该 key 已存向量（或现场 encode 其标题）在
       LanceDB 做 cosine 近邻、排除自身，results 每条结构与 /search 一致（前端复用 resultCard）。
       light 模式或取不到向量 → {ok:false}，前端回退现有 /similar/keywords 抽词法。"""
    try:
        res = R.neighbors(key, topk)
    except Exception as e:
        log_error("similar vector", repr(e))
        res = None
    if res is None:
        return {"ok": False}
    return {"ok": True, "key": key, "results": res}

# ── 综合层 wiki（答案沉淀 / 按需综述）──────────────────────────
class WikiSaveQ(BaseModel):
    query: str = ""
    answer: str
    sources: List[dict] = []      # 前端传 /chat 的 sources；服务端只认 key、重解析页级引用
    model: str = ""               # generated_by（可信度审计）
    by_agent: bool = False        # True=agent 经 MCP 写回（默认采纳、标 🤖 待人复看/剔除）

@app.post("/wiki/answer")
def wiki_answer(q: WikiSaveQ):
    """Phase 0：把一次问答沉淀成 answer 综合页（存盘 + 入表可检索）。网页由"保存此答案"、agent 由 save_synthesis 触发。"""
    try:
        meta = W.save_answer(q.query, q.answer, q.sources, generated_by=q.model, by_agent=q.by_agent)
        return {"ok": True, "id": meta["id"], "title": meta["title"],
                "indexed": meta.get("indexed", False), "n_sources": len(meta.get("sources", []))}
    except W.WikiWriteDenied as e:
        # 正常的权限拒绝（agent 想覆盖人工核验页），不是故障，不进 errors.log
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except Exception as e:
        log_error("wiki/answer", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/list")
def wiki_list(offset: int = 0, limit: int = 0):
    """EN-A0：分页（契约10）。offset/limit 可选；limit<=0 = 不截断（兼容旧调用方一把全拉）。
       响应新增 total=全量页数，页内字段保持不变。"""
    pages = W.list_pages()
    total = len(pages)
    off = max(0, int(offset or 0))
    lim = int(limit or 0)
    out = pages[off:off + lim] if lim > 0 else pages[off:]
    return {"pages": out, "total": total}

@app.get("/wiki/timeline")
def wiki_timeline(limit: int = 100):
    """EN-W3：全库时间线（契约1）。git log 解析优先，无 git 退 .history 快照目录。
       events=[{time,page_id,action,by_agent}]，新的在前；source=git|history。"""
    try:
        import wiki_vcs as V
        return V.log_events(max(1, min(int(limit or 100), 500)))
    except Exception as e:
        log_error("wiki/timeline", repr(e), traceback.format_exc())
        return {"events": [], "source": "history"}     # 时间线是展示件，失败退空不报 500

# ── EN-W2：wiki 更新建议（Ingest 扳机的消费端，契约2）──────────────
@app.get("/wiki/suggestions")
def wiki_suggestions():
    """深索扳机算出的「新文献 → 受影响综述页」建议清单（新的在前，最多 50 条）。"""
    return {"items": _wiki_sugg_load()}

class WikiSuggDismissQ(BaseModel):
    key: str

@app.post("/wiki/suggestions/dismiss")
def wiki_suggestions_dismiss(q: WikiSuggDismissQ):
    """忽略某条建议（按 key 从文件剔除）。幂等：key 不在清单里也返回 ok。"""
    with _WIKI_SUGG_LOCK:
        items = [x for x in _wiki_sugg_load() if x.get("key") != q.key]
        _wiki_sugg_save(items)
    return {"ok": True}

@app.get("/wiki/page/{page_id}")
def wiki_page(page_id: str):
    p = W.get_page(page_id)
    if not p:
        return JSONResponse({"error": "无此综合页", "id": page_id}, status_code=404)
    return p

@app.delete("/wiki/page/{page_id}")
def wiki_delete(page_id: str):
    """§6.4 一键"不保存"：删 md + index 条目 + LanceDB 表行。**仅人用（UI）——不做成 MCP 工具，agent 无删权。**"""
    try:
        r = W.delete_page(page_id)
        if not r.get("deleted"):
            return JSONResponse({"ok": False, "detail": "无此综合页", "id": page_id}, status_code=404)
        return {"ok": True, "id": page_id, **r}
    except Exception as e:
        log_error("wiki/delete", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ── stale 写侧 + by_source 反查出口（此前两者都只建不用）──
class WikiStaleQ(BaseModel):
    stale: bool = True
    reason: str = ""

@app.post("/wiki/stale/{page_id}")
def wiki_stale(page_id: str, q: WikiStaleQ):
    """标记/清除某综合页「已过时」。检索降权立即生效，无需重启。"""
    try:
        return {"ok": True, **W.set_stale(page_id, q.stale, q.reason)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/stale", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

class WikiVerifyQ(BaseModel):
    page_id: str

@app.post("/wiki/verify")
def wiki_verify(q: WikiVerifyQ):
    """W3：人工核验盖章——页面 frontmatter/index/内存三处写 verified_at。
       只给 UI 用，不做成 MCP 工具（核验是人的动作，agent 不得给自己的产出盖章）。"""
    try:
        r = W.set_verified(q.page_id)
        return {"ok": True, "verified_at": r["verified_at"]}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/verify", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/backlinks")
def wiki_backlinks(key: Optional[str] = None, page_id: Optional[str] = None):
    """key=论文 → 引用它的综合页；page_id=页 → 它的来源与互链（含 orphan 判定）。"""
    try:
        return {"ok": True, **W.backlinks(key=key, page_id=page_id)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)
    except Exception as e:
        log_error("wiki/backlinks", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ── 波次2：互链写侧 / 任意页读写 / 体检 / 一源触多页 / 图 / 版本历史 ──
class WikiLinksQ(BaseModel):
    links: List[str] = []
    mode: str = "replace"           # replace | add | remove
    by_agent: bool = False          # 只用于时间线标注（MCP 侧传 true）

@app.post("/wiki/links/{page_id}")
def wiki_set_links(page_id: str, q: WikiLinksQ):
    """写 links —— 把一堆孤岛补成一张图。拒绝自链与断链，返回 skipped。"""
    try:
        return {"ok": True, **W.set_links(page_id, q.links, q.mode, by_agent=q.by_agent)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/links", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

class WikiUpdateQ(BaseModel):
    kind: Optional[str] = None
    title: Optional[str] = None
    content: str = ""
    sources: Optional[List[dict]] = None
    mode: str = "replace"           # replace | append
    links: Optional[List[str]] = None
    model: str = ""
    by_agent: bool = False

@app.post("/wiki/page/{page_id}")
def wiki_update_page(page_id: str, q: WikiUpdateQ):
    """建 / 覆盖 / 追加任意 kind 的 wiki 页（含 entity / overview）。沿用 agent 写权护栏。"""
    try:
        m = W.update_page(page_id, kind=q.kind, title=q.title, content=q.content,
                          sources=q.sources, mode=q.mode, links=q.links,
                          generated_by=q.model, by_agent=q.by_agent)
        return {"ok": True, "id": m["id"], "kind": m.get("kind"), "title": m.get("title"),
                "indexed": m.get("indexed", False), "degraded": m.get("degraded", False),
                "links": m.get("links", []), "n_sources": len(m.get("sources", []))}
    except W.WikiWriteDenied as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)
    except Exception as e:
        log_error("wiki/update", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/lint")
def wiki_lint(min_mentions: int = 2):
    """gist 的 Lint：孤儿页 / 过时页 / 断链 / 无来源 / 降级页 / 缺失概念页。纯读，零副作用。"""
    try:
        return {"ok": True, **W.lint(min_mentions)}
    except Exception as e:
        log_error("wiki/lint", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/graph")
def wiki_graph():
    """图视图数据：节点=页，边=links，孤儿页标出。给应用内的图视图用（不依赖 Obsidian）。"""
    try:
        return {"ok": True, **W.graph()}
    except Exception as e:
        log_error("wiki/graph", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/propose/{source_key}")
def wiki_propose(source_key: str, topk: int = 12):
    """gist 核心：一篇源影响了哪些 wiki 页、每页该怎么改。只建议，不动手。"""
    try:
        return {"ok": True, **W.propose_updates(source_key, topk)}
    except Exception as e:
        log_error("wiki/propose", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/history/{page_id}")
def wiki_history(page_id: str, limit: int = 30):
    try:
        return {"ok": True, **W.page_history(page_id, limit)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/history", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

class WikiRestoreQ(BaseModel):
    rev: str

@app.post("/wiki/restore/{page_id}")
def wiki_restore(page_id: str, q: WikiRestoreQ):
    """回滚某页到历史版本。回滚本身也记一版，可再滚回去。**仅人用**，不做成 MCP 工具。"""
    try:
        return {"ok": True, **W.restore_page(page_id, q.rev)}
    except ValueError as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=404)
    except Exception as e:
        log_error("wiki/restore", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.get("/wiki/vcs")
def wiki_vcs_status():
    """版本历史后端：git（开发机）或 snapshot（多数 exe 用户机器上没装 git）。"""
    try:
        import wiki_vcs as V
        return {"ok": True, **V.status()}
    except Exception as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ── Phase 1：按需生成概念/主题综述页（命中缓存 0 成本；LLM 综合 + LLM 命名）──
class WikiSynthQ(BaseModel):
    concept: str = ""
    topic_id: Optional[int] = None
    force: bool = False           # 忽略缓存、强制重生
    provider: str = "siliconflow"
    base_url: str = ""
    api_key: str = ""
    model: str = ""
    topk: int = 8

def _synth_llm(q):
    return {"provider": q.provider, "base_url": q.base_url, "api_key": q.api_key,
            "model": q.model, "topk": q.topk}

def _synth_ret(m):
    return {"ok": True, "id": m["id"], "title": m["title"], "kind": m.get("kind"),
            "cached": m.get("cached", False), "indexed": m.get("indexed", False),
            "n_sources": len(m.get("sources", []))}

@app.post("/wiki/concept")
def wiki_concept(q: WikiSynthQ):
    try:
        return _synth_ret(W.synthesize_concept(q.concept, force=q.force, **_synth_llm(q)))
    except Exception as e:
        log_error("wiki/concept", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/wiki/topic")
def wiki_topic(q: WikiSynthQ):
    try:
        return _synth_ret(W.synthesize_topic(q.topic_id, force=q.force, **_synth_llm(q)))
    except Exception as e:
        log_error("wiki/topic", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/wiki/regenerate/{page_id}")
def wiki_regen(page_id: str, q: WikiSynthQ):
    try:
        return _synth_ret(W.regenerate(page_id, **_synth_llm(q)))
    except Exception as e:
        log_error("wiki/regenerate", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

# ══════════════════════════════════════════════════════════════════
#  Phase 5：半自动研究助手（选题/框架/带页级引注的资料汇编/建议补文献）
# ══════════════════════════════════════════════════════════════════
class ResearchQ(BaseModel):
    query: str = ""
    topic: str = ""
    topk: int = 14
    force: bool = False
    provider: str = "siliconflow"
    base_url: str = ""
    api_key: str = ""
    model: str = ""
    by_agent: bool = False

def _research_llm(q):
    return {"provider": q.provider, "base_url": q.base_url, "api_key": q.api_key, "model": q.model}

@app.get("/research/pagemap/{key}")
def research_pagemap(key: str):
    """调试：查某篇 PDF页→期刊印刷页码 的映射与 quality。"""
    try:
        import page_map as PM
        doc = PM.build(key)
        if not doc:
            return JSONResponse({"ok": False, "detail": "该篇无提取文本（未深索或扫描件），无法建映射"}, status_code=404)
        return {"ok": True, **doc}
    except Exception as e:
        log_error("research/pagemap", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/digest")
def research_digest(q: ResearchQ):
    """能力二：给一个子题 query → 带印刷页引注的综述 + 覆盖评级 + 缺口。存 kind=digest wiki 页。"""
    try:
        import research_assistant as RA
        m = RA.digest(q.query or q.topic, topk=q.topk, llm=_research_llm(q), force=q.force, by_agent=q.by_agent)
        # 透传 degraded（未配/失效 key、LLM 调用失败、或库内无命中的降级页）：降级页不入检索表，
        # MCP/前端据此如实措辞，别再无条件宣称「已写回 wiki 页、可被检索」。
        return {"ok": True, "id": m["id"], "title": m["title"], "kind": m.get("kind"),
                "cached": m.get("cached", False), "indexed": m.get("indexed", False),
                "degraded": m.get("degraded", False), "degraded_reason": m.get("degraded_reason"),
                "coverage": m.get("coverage"), "n_sources": len(m.get("sources", []))}
    except W.WikiWriteDenied as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except Exception as e:
        log_error("research/digest", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/scope")
def research_scope(q: ResearchQ):
    """能力一：主题 → 范围映射 + 选题拆解 + 标题候选 + 三级大纲(★/☆)。存 kind=outline wiki 页。"""
    try:
        import research_assistant as RA
        m = RA.scope(q.topic or q.query, topk=q.topk, llm=_research_llm(q), force=q.force, by_agent=q.by_agent)
        return {"ok": True, "id": m["id"], "title": m["title"], "kind": m.get("kind"),
                "cached": m.get("cached", False), "indexed": m.get("indexed", False),
                "degraded": m.get("degraded", False), "degraded_reason": m.get("degraded_reason"),
                "n_sources": len(m.get("sources", []))}
    except W.WikiWriteDenied as e:
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=409)
    except Exception as e:
        log_error("research/scope", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/suggest_sources")
def research_suggest(q: ResearchQ):
    """能力三：覆盖评估 + 脚注引文挖掘缺失文献 + 库内错配（按期刊层级排）。只读，不写库。"""
    try:
        import research_assistant as RA
        return {"ok": True, **RA.suggest_sources(q.topic or q.query, topk=q.topk, llm=_research_llm(q))}
    except Exception as e:
        log_error("research/suggest_sources", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

def _export_docx(page_id: str):
    """把 digest/outline wiki 页导出。有 python-docx 则出 .docx，否则退化为 .md（诚实降级）。
       C3：末尾用 p['sources'] 追加一节「参考文献」（页级引用），产出可直接拿去写作。"""
    p = W.get_page(page_id)
    if not p:
        return JSONResponse({"ok": False, "detail": "无此页"}, status_code=404)
    sources = p.get("sources", []) or []
    def _cite(s):
        if isinstance(s, dict):
            return s.get("citation") or s.get("key") or ""
        return str(s)
    try:
        import docx  # python-docx
        from docx import Document
        doc = Document()
        doc.add_heading(p.get("title", "资料汇编"), level=0)
        for line in (p.get("markdown") or "").splitlines():
            s = line.strip()
            if s.startswith("### "):
                doc.add_heading(s[4:], level=2)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=1)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s and not s.startswith("---"):
                doc.add_paragraph(s)
        if sources:
            doc.add_heading("参考文献", level=1)
            for i, s in enumerate(sources, 1):
                doc.add_paragraph(f"{i}. {_cite(s)}")
        out = C.WIKI_DIGEST_DIR / f"{page_id}.docx"
        doc.save(str(out))
        return FileResponse(str(out), filename=f"{page_id}.docx",
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except ImportError:
        # 降级：返回 markdown 文件（本机未装 python-docx）
        md = p.get("markdown", "")
        if sources:
            md += "\n\n## 参考文献\n\n" + "\n".join(f"{i}. {_cite(s)}" for i, s in enumerate(sources, 1))
        out = C.WIKI_DIGEST_DIR / f"{page_id}.md"
        out.write_text(md, encoding="utf-8")
        return FileResponse(str(out), filename=f"{page_id}.md", media_type="text/markdown")
    except Exception as e:
        log_error("research/export_docx", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)

@app.post("/research/export_docx/{page_id}")
def research_export_docx(page_id: str):
    return _export_docx(page_id)

@app.get("/research/export_docx/{page_id}")
def research_export_docx_get(page_id: str):
    """C3：GET 版——前端用 <a href download> 直接触发下载。"""
    return _export_docx(page_id)

# ══════════════════════════════════════════════════════════════════
#  EN-A：Agent 接入的服务端地基（蓝图 G1/G2/G4 + 入库闭环）
#  单篇题录 / 引注引擎出口 / 引文定位 / 论断核验 / 本地路径入库 / AI 使用声明
# ══════════════════════════════════════════════════════════════════

@app.get("/paper/{key}")
def paper_detail(key: str):
    """EN-A1（契约4）：单篇完整题录——agent 写引注/核出处前先拿全字段，别再从 /papers
       整表里大海捞针。cited_by_wiki 走 wiki index.json 的 by_source 反查（W.backlinks）；
       statute_status 读 papers.jsonl 同名字段（法源改造写入；缺省空串）。"""
    p = _load_papers().get(key)
    if not p:
        return JSONResponse({"detail": f"未找到文献 {key}（可先用 /papers 枚举）"}, status_code=404)
    import grading_svc as GS
    g = GS.grade_paper(p, compute=False)     # 快路径：只用已预热 memo，与 /papers 同口径
    cited_by_wiki = []
    try:
        cited_by_wiki = [{"id": c.get("id"), "title": c.get("title", "")}
                         for c in W.backlinks(key=key).get("cited_by", [])]
    except Exception as e:
        log_error("paper detail backlinks", repr(e))   # 反查失败不拦题录主体
    stem = p.get("stem") or T.safe_name(key)
    _isdeep = is_deep(key)
    extract_status = _extract_record(stem, legacy_deep=_isdeep)
    return {
        "key": key, "title": p.get("title", ""), "author": p.get("author", ""),
        "year": p.get("year", ""), "journal": p.get("journal", ""),
        "itemtype": p.get("itemtype", ""),
        "weight_tier": (g["cn"] if g else p.get("journal_tier", "")),
        "collections": p.get("collections", []),
        "official_pages": p.get("official_pages", ""),
        "has_pdf": bool(p.get("has_pdf", False)),
        "deep": _isdeep,
        "no_text": stem in _deep_no_text_keys(),
        "extract_status": extract_status,
        "abstract": p.get("abstract", ""),
        "ingested_at": p.get("ingested_at", ""),
        "statute_status": p.get("statute_status", "") or "",
        "cited_by_wiki": cited_by_wiki,
    }

@app.get("/cite/{key}")
def cite_paper(key: str, page: Optional[int] = None, style: str = "footnote", heading: Optional[str] = None):
    """EN-A2（契约5，G2 引注引擎出口）：薄封装 cite_format——格式由规则做、绝不交 LLM。
       只透传 style 与 page；statute/report 模板由 cite_format 按 itemtype 自动分派。
       heading（可选）：法条条号（如「第201条」）——不传时对 statute 按 (key,page) 反查该页 chunk 的 heading 自动补。
       missing_fields 让 agent 知道哪些字段缺了该去补题录，而不是让它自己瞎编。"""
    if style not in ("footnote", "compact"):
        return JSONResponse({"detail": "style 仅支持 footnote | compact"}, status_code=400)
    p = _load_papers().get(key)
    if not p:
        return JSONResponse({"detail": f"未找到文献 {key}"}, status_code=404)
    import cite_format as CF
    import page_map as PM
    hit = dict(p)                      # papers.jsonl 字段名与检索结果 hit 一致，直接当 hit 用
    hit["page"] = page
    # 法条条号：优先用调用方传入的 heading；否则 statute + page 给定时，从命中 chunk 行按 (key,page) 反查该页 heading。
    head = (heading or "").strip()
    if not head and page is not None and (p.get("itemtype") or "").strip() == "statute":
        try:
            head = R.find_statute_heading(key, page)
        except Exception as e:
            log_error("cite heading lookup", repr(e))
    formatted = CF.footnote(hit, heading=head) if style == "footnote" else CF.compact(hit, heading=head)
    # 印刷页展示串 + 推算标记：只有映射真产出了推算值才算 estimated（没映射不算）
    printed_disp, page_estimated = "", False
    if page is not None:
        try:
            pm = PM.printed(key, page) or {}
            printed_disp = pm.get("display") or ""
            page_estimated = bool(pm.get("printed") is not None
                                  and (pm.get("method") in ("interp", "offset", "pdfseq")
                                       or float(pm.get("conf") or 0) < 0.7))
        except Exception as e:
            log_error("cite printed", repr(e))
    # 缺字段按 itemtype 分派（与 cite_format 的模板字段一一对应）
    it = (p.get("itemtype") or "").strip()
    miss = []
    has_pg = bool(printed_disp or (p.get("official_pages") or "").strip())
    if it == "statute":
        if not (p.get("title") or "").strip(): miss.append("title")
        # 与 cite_format._statute_cite 同口径：法规标题惯例「（2018修正）」不含「年」字，判四位数字
        if not str(p.get("year") or "").strip() and not re.search(r"\d{4}", p.get("title") or ""):
            miss.append("year")
    elif it == "report":
        if not ((p.get("author") or "").strip() or (p.get("journal") or "").strip()):
            miss.append("author")      # 机构作者：author 首位，退 journal 位（模板同款回退）
        if not (p.get("title") or "").strip(): miss.append("title")
        if not str(p.get("year") or "").strip(): miss.append("year")
    else:                              # 期刊式（含 itemtype 不认识的回退）
        for f, v in (("author", p.get("author")), ("title", p.get("title")),
                     ("journal", p.get("journal")), ("year", p.get("year"))):
            if not str(v or "").strip():
                miss.append(f)
        if not has_pg:
            miss.append("page")        # 手册脚注式没页码是硬伤，必须让 agent 知道
    return {"ok": True, "key": key, "style": style, "page": page,
            "printed_page": printed_disp, "itemtype": it,
            "formatted": formatted, "missing_fields": miss,
            "page_estimated": page_estimated}

class LocateQuoteQ(BaseModel):
    quote: str
    key: Optional[str] = None
    fuzzy: bool = True

@app.post("/research/locate_quote")
def research_locate_quote(q: LocateQuoteQ):
    """EN-A3（契约6）：在提取全文里定位一段引文 → 哪篇/PDF第几页/印刷第几页。
       key 给定只搜单篇；否则全库（cap 500 篇，截断会在结果里注明）。"""
    try:
        import textloc as TL
        return TL.locate(q.quote, key=q.key, fuzzy=bool(q.fuzzy))
    except ValueError as e:
        return JSONResponse({"detail": str(e)}, status_code=400)   # 引文太短等参数性拒绝
    except Exception as e:
        log_error("research/locate_quote", repr(e), traceback.format_exc())
        return JSONResponse({"detail": str(e)}, status_code=400)

class VerifyClaimQ(BaseModel):
    claim: str
    keys: Optional[List[str]] = None
    topk: int = 8

@app.post("/research/verify_claim")
def research_verify_claim(q: VerifyClaimQ):
    """EN-A4（契约7，G1 核验器）：论断→三态判定 supported/mismatch/not_in_lib。
       保守原则见 verify_claim.py 模块注释（铁律三：库内无 ≠ 论断为假）。"""
    try:
        import verify_claim as VC
        return VC.verify(q.claim, keys=q.keys, topk=q.topk)
    except ValueError as e:
        return JSONResponse({"detail": str(e)}, status_code=400)
    except Exception as e:
        log_error("research/verify_claim", repr(e), traceback.format_exc())
        return JSONResponse({"detail": str(e)}, status_code=400)

class LocalPathQ(BaseModel):
    path: str
    note: Optional[str] = None

@app.post("/ingest/local_path")
def ingest_local_path(q: LocalPathQ):
    """EN-A5（契约8）：按本地绝对路径把一个 PDF 收进受管文件夹并触发增量建库——
       agent 替用户"把桌面上这篇收进库里"的入库闭环。仅文件夹模式；Zotero 模式的
       入库动作属于 Zotero（改它的库文件是越权），400 提示走 Zotero。"""
    import settings as S, hashlib as _hl, shutil
    if S.source() != "folder":
        return JSONResponse({"detail": "Zotero 模式请把 PDF 附到 Zotero 条目上（Zotero 是它库的唯一主人），"
                                       "随后等自动更新或点「更新知识库」即可入库。"}, status_code=400)
    folder = S.folder_dir()
    if not folder:
        return JSONResponse({"detail": "未配置受管文件夹，请先在向导/设置里选定"}, status_code=400)
    # 路径安全：只接受绝对路径（相对路径取决于服务进程 cwd，agent 传来毫无意义且易被诱导）、拒绝目录
    src = Path((q.path or "").strip())
    if not str(src) or not src.is_absolute():
        return JSONResponse({"detail": "只接受绝对路径"}, status_code=400)
    if src.is_dir():
        return JSONResponse({"detail": "只接受单个 PDF 文件，不接受目录"}, status_code=400)
    if not src.exists():
        return JSONResponse({"detail": f"文件不存在：{src}"}, status_code=400)
    if src.suffix.lower() != ".pdf":
        return JSONResponse({"detail": "只支持 .pdf 文件"}, status_code=400)
    try:
        data = src.read_bytes()
    except Exception as e:
        return JSONResponse({"detail": f"读取文件失败：{e}"}, status_code=400)
    import folder_source as FS
    fp = Path(folder)
    # sha1 查重（同 /ingest/files 的 R2 思路：先按大小粗筛，同大小才读盘算 sha1，大库不卡）
    h = _hl.sha1(data).hexdigest()
    dup = None
    for pth in fp.rglob("*.pdf"):
        try:
            if pth.stat().st_size != len(data):
                continue
            if _hl.sha1(pth.read_bytes()).hexdigest() == h:
                dup = pth; break
        except Exception:
            continue
    if dup is not None:
        return {"ok": True, "status": "duplicate", "key": FS.stable_key(folder, str(dup)),
                "building": False, "need_review": True,
                "hint": f"内容相同的 PDF 已在库中（{dup.name}），未重复入库。"}
    dst = _dedupe_name(fp / src.name)     # 同名不同容 → 加序号，绝不覆盖既有文件
    try:
        shutil.copy2(str(src), str(dst))
    except Exception as e:
        log_error("ingest/local_path copy", repr(e), traceback.format_exc())
        return JSONResponse({"detail": f"复制进受管文件夹失败：{e}"}, status_code=500)
    key = FS.stable_key(folder, str(dst))
    if q.note:                            # 入库备注留痕（追加 jsonl，供人日后查"这篇谁让收的"）
        try:
            nf = C.FOLDER_DIR_STATE / "ingest_notes.jsonl"
            with open(nf, "a", encoding="utf-8") as f:
                f.write(json.dumps({"time": time.strftime("%Y-%m-%d %H:%M:%S"), "key": key,
                                    "file": dst.name, "note": q.note}, ensure_ascii=False) + "\n")
        except Exception as e:
            log_error("ingest/local_path note", repr(e))
    # 触发增量建库：BUILD 忙时不建（building:false），调用方稍后触发即可——文件已落
    # 受管文件夹，下一次任何 folder build/自动更新都会把它捎上，不会丢
    building = False
    if not BUILD["running"]:
        building = bool(_run_build("folder"))
    return {"ok": True, "status": "added", "key": key, "building": building, "need_review": True,
            "hint": ("已复制进受管文件夹并开始后台抽题录+建索引。" if building
                     else "已复制进受管文件夹；当前有构建在跑，稍后自动更新或手动「更新知识库」即可入索引。")
                    + "题录由 AI 抽取，请在「浏览」页核对（needs_review）。"}

class DisclosureQ(BaseModel):
    page_ids: List[str] = []

@app.post("/research/disclosure")
def research_disclosure(q: DisclosureQ):
    """EN-A6（契约9，G4）：按所选 wiki 页的元数据生成《生成式 AI 使用声明》模板文本。
       规则拼装、零 LLM——披露必须是机械的事实陈述（实现在 research_assistant.disclosure）。"""
    try:
        import research_assistant as RA
        return {"text": RA.disclosure(q.page_ids)}
    except Exception as e:
        log_error("research/disclosure", repr(e), traceback.format_exc())
        return JSONResponse({"detail": str(e)}, status_code=400)

@app.get("/papers")
def papers(collection: Optional[str] = None, topic: Optional[int] = None,
           category: Optional[str] = None, deep: Optional[str] = None,
           sort: str = "recommend", limit: int = 300, offset: int = 0,
           since: Optional[str] = None):
    # EN-A7：since=YYYY-MM-DD 按 ingested_at 过滤（配合 whats_new：「上次见面后新入了什么」）。
    # ingested_at 形如 "YYYY-MM-DD HH:MM:SS"，与 since 直接字典序比较即可；
    # 没有 ingested_at 的老条目（空串）在 since 模式下会被滤掉——它们本来就不是"新入库"。
    papers = _load_papers(); cats = _load_cats(); deepk = _deep_keys()
    notextk = _deep_no_text_keys()   # 兼容旧字段：当前不可继续深索的终态
    extract_items = _deep_extract_items()
    sumk = _summary_keys()           # 有效检索摘要（SAC）的 safe_name(stem) 集合
    sum_issues = _summary_issues()
    if category:
        # 统一走 _resolve_category_keys（与 /search 同一条路）：kbc_/topic:/zotero: 都能过滤；
        # 旧写法只认 kbc_ 前缀，topic:/zotero: 会静默落到全库，list_sources 聚焦失效。
        ks = _resolve_category_keys(category) or set()
        items = [papers[k] for k in ks if k in papers]
    elif topic is not None:
        ait = _load_ai_topics()
        tk = [k for k, tid in ait.get("by_key", {}).items() if tid == topic]
        items = [papers[k] for k in tk if k in papers]
    elif collection and collection in cats.get("by_collection", {}):
        items = [papers[k] for k in cats["by_collection"][collection] if k in papers]
    else:
        items = list(papers.values())
    import grading_svc as GS
    out = []
    filter_counts = {name: 0 for name in (
        "all", "yes", "no", "ocr", "native",
        "summary_yes", "summary_invalid", "summary_no",
    )}
    for p in items:
        _isdeep = is_deep(p["key"], deepk)
        stem = p.get("stem") or T.safe_name(p["key"])
        summary_stem = T.safe_name(stem)
        extract_rec = _extract_record(stem, extract_items, legacy_deep=_isdeep)
        has_summary = summary_stem in sumk
        summary_invalid = summary_stem in sum_issues
        if since and (p.get("ingested_at") or "") < since:   # EN-A7：早于 since（或无入库时间）的滤掉
            continue
        filter_counts["all"] += 1
        for filter_name in ("yes", "no", "ocr", "native",
                            "summary_yes", "summary_invalid", "summary_no"):
            if _browse_filter_matches(filter_name, _isdeep, extract_rec,
                                      has_summary, summary_invalid):
                filter_counts[filter_name] += 1
        if deep and not _browse_filter_matches(deep, _isdeep, extract_rec,
                                               has_summary, summary_invalid):
            continue
        # F38-B：按当前学科取分级（快路径 compute=False，只用已预热 memo；未预热则回退旧 journal_tier）
        # 手动改档/法源报告规则在 grade_paper 里优先命中（不走 memo，即改即显）。
        g = GS.grade_paper(p, compute=False)
        out.append({
            "key": p["key"], "title": p.get("title", ""), "author": p.get("author", ""),
            "year": p.get("year", ""), "journal": p.get("journal", ""),
            "journal_tier": p.get("journal_tier", ""), "tier_rank": p.get("tier_rank", 6),
            "weight_tier": (g["cn"] if g else ""),           # 学科感知中文档名（未预热时空→前端兜旧）
            "weight_rank": (g["rank"] if g else 6),
            "journal_weight": (g["weight"] if g else None),
            "weight_needs_review": (g["needs_review"] if g else False),
            "weight_src": (g.get("src") if g else None),     # manual=手动改档 / rule=法源报告规则
            "official_pages": p.get("official_pages", ""), "has_pdf": p.get("has_pdf", False),
            "collections": p.get("collections", []),
            "needs_review": bool(p.get("needs_review", False)),   # folder 模式：AI 抽的题录待核对
            "no_text": stem in notextk,
            "extract_status": extract_rec,
            "ingested_at": p.get("ingested_at", ""),               # 供「最新入库」排序
            "itemtype": p.get("itemtype", ""),
            "statute_status": p.get("statute_status", ""),         # EN-L5：法条时效徽标（浏览卡 statuteBadge 读它）
            "score": _rec_score(p, g), "deep": _isdeep,
            "has_summary": has_summary,
            "summary_invalid": summary_invalid,
            "summary_error": sum_issues.get(summary_stem, ""),
        })
    if sort == "recommend":
        out.sort(key=lambda x: -x["score"])
    elif sort == "year":
        def _y(x):
            try: return int(x["year"] or 0)
            except Exception: return 0
        out.sort(key=lambda x: -_y(x))
    elif sort == "ingested":
        out.sort(key=lambda x: x.get("ingested_at", ""), reverse=True)   # 最新入库优先
    # W1：分页出口——大库此前只能看到前 limit 篇，其余永远翻不到；total=过滤后总数供前端算页数
    off = max(0, int(offset or 0))
    return {"papers": out[off:off + limit], "total": len(out),
            "filter_counts": filter_counts,
            "collection": collection, "topic": topic, "category": category, "deep": deep,
            "sort": sort, "since": since}

# ── 单篇手动改档（法源权重改造 2026-07-12）──────────────────
class TierOverrideQ(BaseModel):
    key: str
    tier: Optional[str] = None      # "T1"/"T1b"/"T2"/"T3"/"T4"/"T5"；None/空 = 恢复自动

@app.post("/paper/tier")
def set_paper_tier(q: TierOverrideQ):
    """手动提高/降低某篇的档位（存 data/tier_overrides.json，优先级最高，即改即生效）。"""
    import source_rules as SR
    import grading_svc as GS
    tier = (q.tier or "").strip() or None
    if tier and tier not in SR.TIER_W:
        return JSONResponse({"detail": f"非法档位 {tier}（可选 T1/T1b/T2/T3/T4/T5 或留空恢复自动）"},
                            status_code=400)
    p = _load_papers().get(q.key)
    if not p:
        return JSONResponse({"detail": f"未找到文献 {q.key}"}, status_code=404)
    try:
        SR.set_override(q.key, tier)
    except Exception as e:
        log_error("paper tier override", repr(e))
        return JSONResponse({"detail": f"写入改档失败：{e}"}, status_code=500)
    g = GS.grade_paper(p)            # 改完立刻算出生效档（含回落到规则/期刊分级的情形）
    return {"ok": True, "key": q.key, "override": tier,
            "effective": g or {"tier": None, "cn": p.get("journal_tier", "未知"),
                               "weight": None, "rank": p.get("tier_rank", 6),
                               "needs_review": False}}

# ── 打开原文 PDF（C2/D4：页级引注可回到原文核对）─────────────
class OpenPdfQ(BaseModel):
    key: str

@app.post("/open_pdf")
def open_pdf(q: OpenPdfQ):
    """用系统默认阅读器打开某篇的原文 PDF（供深索结果卡「📄 打开原文」）。
       pdf_path 取自 papers.jsonl（zotero/folder 两种源建库时都已落该字段）。"""
    p = _load_papers().get(q.key)
    if not p:
        return JSONResponse({"ok": False, "msg": "无此文献"}, status_code=404)
    path = (p.get("pdf_path") or "").strip()
    if not path or not Path(path).exists():
        return JSONResponse({"ok": False, "msg": "未找到原文 PDF 文件（可能仅题录、或文件已移动/未随库）"}, status_code=200)
    try:
        if sys.platform == "win32":
            os.startfile(path)  # noqa
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
        return {"ok": True, "path": path}
    except Exception as e:
        log_error("open_pdf", repr(e))
        return JSONResponse({"ok": False, "msg": f"打开失败：{e}"}, status_code=200)

# ══════════════════════════════════════════════════════════════════
#  读取论文全文（agent 的 ingest 地基）
#  此前 agent 只能拿到 220 字检索片段；逐页正文一直躺在 config.EXTRACTED 里，
#  却只有 _extracted_excerpt 内部自用、没有任何出口。gist 的 ingest 主干
#  "the LLM reads it, extracts the key information" 因此走不通。
# ══════════════════════════════════════════════════════════════════
@app.get("/source/{key}")
def read_source(key: str, from_page: int = 1, to_page: int = 0, max_chars: int = 20000):
    """按 PDF 顺序页返回该篇提取正文（每页附印刷页码）。

    诚实降级：读不到时明确区分「无此篇 / 只有题录无 PDF / 尚未深索 / 扫描件无文字」，
    各给 reason + 可执行的 detail——绝不返回空串让 agent 以为这篇没内容。"""
    import textutil as T
    papers = _load_papers()
    p = papers.get(key)
    if not p:
        return JSONResponse({"ok": False, "reason": "not_found",
                             "detail": f"知识库中没有 key={key} 的文献。可先用 /papers 或 list_sources 枚举。"},
                            status_code=404)
    title = p.get("title", "")
    if not p.get("has_pdf"):
        return {"ok": False, "reason": "no_pdf", "key": key, "title": title,
                "detail": "该篇只有题录、没有 PDF 原文，读不到全文。"}

    stem = p.get("stem") or T.safe_name(key)
    f = C.EXTRACTED / f"{stem}.json"
    if not f.exists():
        return {"ok": False, "reason": "not_deep_indexed", "key": key, "title": title,
                "detail": "该篇尚未深索，没有可读全文。请在应用「浏览」页勾选它深索，"
                          "或调用 localkb_build(stage='deep')。"}
    try:
        rec = json.loads(f.read_text(encoding="utf-8"))
    except Exception as e:
        log_error("source/read", repr(e))
        return JSONResponse({"ok": False, "reason": "unreadable", "key": key,
                             "detail": f"提取文件损坏：{e}"}, status_code=500)

    pages = rec.get("pages") or []
    if not any((pg.get("text") or "").strip() for pg in pages):
        status = _extract_record(stem).get("status")
        details = {
            "missing_pdf": "该篇记录的 PDF 附件已不在磁盘上，请先在 Zotero 中修复附件路径。",
            "invalid_pdf": "该 PDF 无法打开，可能已损坏、未同步完整或被其他程序占用。",
            "ocr_pending": "该篇正在等待本地 OCR，完成深索后即可读取正文。",
            "ocr_failed": "本地 OCR 已运行，但没有识别出有效文字；可检查 PDF 后重试。",
        }
        return {"ok": False, "reason": status or "no_extracted_text", "key": key,
                "title": title, "detail": details.get(status, "该篇暂时没有可读正文。")}

    lo = max(1, int(from_page or 1))
    # pages 只保存有文字的页；混合 PDF 若有一页 OCR 仍失败，页码会有空洞。
    # 不能用 len(pages) 当末页，否则 [1, 3] 会把真实第 3 页误裁掉。
    last_page = max([int(pg.get("page", 0) or 0) for pg in pages]
                    + [int(rec.get("total_pages", 0) or 0)])
    hi = int(to_page) if to_page else last_page
    try:
        import page_map as PM
    except Exception:
        PM = None

    out, used, truncated, next_page = [], 0, False, None
    budget = max(500, int(max_chars or 20000))
    for pg in pages:
        n = int(pg.get("page", 0) or 0)
        if n < lo or n > hi:
            continue
        txt = (pg.get("text") or "").strip()
        if not txt:
            continue
        if used + len(txt) > budget and out:      # 至少给一页，别因预算太小返回空
            truncated, next_page = True, n
            break
        used += len(txt)
        pr = PM.printed(key, n) if PM else {}
        out.append({"pdf_page": n, "printed_page": (pr or {}).get("display", ""), "text": txt})

    return {"ok": True, "key": key, "title": title,
            "author": p.get("author", ""), "year": p.get("year", ""), "journal": p.get("journal", ""),
            "n_pages_total": last_page, "returned_pages": len(out),
            "chars": used, "truncated": truncated, "next_page": next_page,
            "pages": out}

# ── 三档索引触发 ──────────────────────────────────────────
@app.post("/index/light")
def index_light_ep():
    # BF9：即时索引此前完全绕开构建锁——与 build 子进程并发重写 papers.jsonl/bm25 会写坏索引。
    # 与 /index/deep(scope=all) 同款约定：忙时返回 {ok:false,busy:true}，前端已能处理 ok:false。
    with _BUILD_LOCK:
        if BUILD["running"]:
            return {"ok": False, "busy": True, "msg": "已有构建任务在跑，请稍后再试"}
        BUILD["running"] = True; BUILD["stage"] = "light"
        BUILD["started"] = time.time(); BUILD["rc"] = None; BUILD["cancelled"] = False
        BUILD["log"] = ["[light] 即时索引启动…"]   # 不留上次构建的日志尾巴
    try:
        import importlib, index_light as IL
        importlib.reload(IL)
        stats = IL.main()
        try:
            import build_categories as BC
            importlib.reload(BC); BC.main()          # 只读 zotero.sqlite，约 1s；失败不影响词法索引
        except Exception as e:
            log_error("index/light categories", repr(e), traceback.format_exc())
        # BF9：仿 _run_build 的重载窗口——先置未就绪，防 tbl 新旧错位时检索拿到错误的空命中
        R.STATE["ready"] = False
        R.load_all()
        return {"ok": True, **stats["coverage"]}
    except Exception as e:
        log_error("index/light", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "msg": str(e)}, status_code=500)
    finally:
        with _BUILD_LOCK:
            BUILD["running"] = False
            BUILD["stage"] = ""
        # 与 _run_build 收尾对齐：补跑构建期间被 4s 防抖 timer 丢弃的深索入队请求，否则它们滞留到下次偶然事件。
        try:
            _drain_deep_queue()
        except Exception as e:
            log_error("index/light drain", repr(e))

def _child_env():
    """任务五：子进程统一 env——强制 UTF-8 输出（PYTHONIOENCODING）+ 开 UTF-8 模式（PYTHONUTF8=1），
       让 build 子进程稳定输出 UTF-8，server 端按 utf-8 解码不再乱码。"""
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONUTF8"] = "1"
    return env

def _run_build(stage, extra=None, on_done=None):
    # B1：原子「判 running + 置 True」并在调用线程内同步置位（不再等子线程 start() 后才置），
    # 多路触发时后来者立即拿到 running=True → return False，不会并发跑两个子进程。
    with _BUILD_LOCK:
        if BUILD["running"]:
            return False
        BUILD["running"] = True; BUILD["stage"] = stage; BUILD["started"] = time.time()
        BUILD["log"] = [f"[{stage}] 启动…"]; BUILD["rc"] = None
        BUILD["proc"] = None; BUILD["cancelled"] = False
        # BF35：整库深索(scope=all)标记——extra 形如 ["--scope","all"]；队列批次是 "keys:..."，不算 bulk
        BUILD["bulk"] = bool(stage == "deep" and "all" in (extra or []))
    # EN-W2：整库深索(scope=all)记下构建前的已深索集合（safe_name(stem)），
    # 成功收尾后与新集合作差 → 本轮真正新深索的篇 → 触发 wiki 更新建议。
    # 队列批次不走这里（keys 已知，由 _on_deep_done 直接触发），避免双算。
    pre_deep = _deep_keys() if BUILD["bulk"] else None
    def run():
        rc = None
        try:
            env = _child_env()   # 任务五：稳定 UTF-8 输出，避免 build 日志乱码
            cmd = [sys.executable, str(C.APP / "build_all.py"), "--stage", stage] + (extra or [])
            p = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env,
                                 creationflags=C.SUBPROC_NO_WINDOW,   # ★ 不闪黑窗（Windows）
                                 # macOS/Linux：起在独立进程组，取消时 _kill_tree 用 killpg 连孙进程一起杀
                                 start_new_session=(sys.platform != "win32"))
            BUILD["proc"] = p            # 留句柄给 /build/cancel
            for raw in p.stdout:
                BUILD["log"].append(raw.decode("utf-8", errors="replace").rstrip())
                BUILD["log"] = BUILD["log"][-300:]
            # A3：取子进程 returncode（此前 p.wait() 取到却从不检查）。非 0 = 软失败（余额0/PDF异常）。
            rc = p.wait()
            if BUILD.get("cancelled"):
                BUILD["log"].append("[build] 已被用户取消。已完成的部分已保存，可随时继续。")
            elif rc != 0:
                BUILD["log"].append(f"[build] 子进程以非 0 退出（returncode={rc}）——本批未成功完成。")
        except Exception as e:
            log_error(f"build {stage}", repr(e), traceback.format_exc())
            BUILD["log"].append("[build] 异常：" + str(e))
            rc = -1
        finally:
            BUILD["rc"] = rc
            BUILD["proc"] = None
            # B2：重载期间表句柄、BM25 与题录状态并非原子切换。
            # 重载全程保持 running=True（锁未释放前不接受新 build）+ 短暂 ready=False（检索先返回未就绪，
            # 而非错误的空命中）；load_all 成功后自会把 ready 置回 True。
            try:
                R.STATE["ready"] = False
                R.load_all(); BUILD["log"].append("[build] 索引已重载。")
            except Exception as e:
                BUILD["log"].append("[build] 重载失败：" + str(e))
            with _BUILD_LOCK:
                BUILD["running"] = False
                BUILD["stage"] = ""          # F4：构建结束复位 stage，避免残留 "deep" 让顶栏误报「深索中」
                BUILD["bulk"] = False        # BF35：构建结束复位整库深索标记
            # EN-W2：整库深索成功收尾 → 差集出本轮新深索的篇，触发 wiki 更新建议。
            # 放在索引重载之后：propose_updates 里的检索要用重载后的新表/新题录才准。
            # embedded_keys.txt 存的是 safe_name(stem)，须映射回原始 key 才能给 propose_updates。
            if rc == 0 and not BUILD.get("cancelled") and pre_deep is not None:
                try:
                    new_stems = _deep_keys() - pre_deep
                    if new_stems:
                        newk = [k for k in _load_papers() if T.safe_name(k) in new_stems]
                        _wiki_suggest_async(newk)
                except Exception as e:
                    log_error("wiki suggest (bulk deep)", repr(e))
            # A3/A4：把 returncode 贯通给回调（失败时上层决定退回队列/不推进 sig）。
            if on_done:
                try:
                    on_done(rc)
                except Exception as e:
                    log_error("deep queue on_done", repr(e))
            elif not BUILD.get("cancelled") and not (pre_deep is not None and rc != 0):
                # 被取消时不排空队列——否则子进程一死就立刻起新批，用户会以为取消没生效。
                # ★ 整库深索失败(rc≠0)时也不立刻排空：否则紧接着起的队列批次会把全局 BUILD["rc"]
                #   从非0覆写成0/None（server.py:2535 起手重置 + 空批退0），前端在 deep 忙→闲的
                #   跳变里读到 rc=0/None 就误报「✓ 深索完成」，而其实半个库没入。让失败的整库深索
                #   自然落地(idle, rc≠0)，队列已持久化，改由下次 enqueue / 每日增量 / 重启时排空。
                #   条件 pre_deep is not None ⟺ 本次是整库深索；常规 all/semantic/light/folder 维持原样。
                try:
                    _drain_deep_queue()
                except Exception as e:
                    log_error("deep queue drain", repr(e))
    threading.Thread(target=run, daemon=True).start()
    return True

# ── 更新知识库（增量 all 档：读 Zotero 新增题录 + 补语义层，不动深索）──
# 前端「更新知识库」按钮 POST /build（此前无此路由、404 被静默吞→按钮假死）。
# _run_build("all") 走 build_all.py 的增量管线（已入库跳过），深索仍由用户手动触发。
@app.post("/build")
def build_ep():
    # 文件夹模式必须走 stage=folder（含 folder_ingest 的 LLM 抽题录），否则手动「更新知识库」
    # 只跑 all 管线、跳过题录抽取，新 PDF 永远停留在文件名占位题录（与自动更新循环同一分派）。
    import settings as S
    stage = "folder" if S.source() == "folder" else "all"
    return {"ok": _run_build(stage)}

def _kill_tree(p):
    """终止建库进程树。子进程（build_all）会派生嵌入/抽取 worker 孙进程，只杀本体它们会变孤儿、
       继续跑继续烧 API 额度。
       Windows：taskkill /T /F 杀整棵树；失败兜底 terminate。
       macOS/Linux：子进程用 start_new_session 起在独立进程组里（见 _run_build 的 Popen），
         这里 killpg 杀整组；拿不到进程组则兜底 terminate。"""
    if sys.platform == "win32":
        try:
            r = subprocess.run(["taskkill", "/PID", str(p.pid), "/T", "/F"],
                               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=15,
                               creationflags=C.SUBPROC_NO_WINDOW)
            if r.returncode == 0:
                return
            log_error("build/cancel taskkill", f"taskkill rc={r.returncode}，兜底 terminate")
        except Exception as e:
            log_error("build/cancel taskkill", repr(e))
    else:
        try:
            import os as _os, signal as _sig
            _os.killpg(_os.getpgid(p.pid), _sig.SIGTERM)   # 杀整个进程组（含嵌入/抽取孙进程）
            return
        except Exception as e:
            log_error("build/cancel killpg", repr(e))
    p.terminate()

@app.post("/build/cancel")
def build_cancel():
    """取消正在跑的建库/深索子进程。整库深索(scope=all)此前一旦开始就停不下来，
       可能空跑数小时并烧掉 API 额度。深索是增量的：已完成的篇已落盘，随时可继续。
       同时暂停深索队列——否则子进程一死，_drain_deep_queue 立刻起新批。"""
    p = BUILD.get("proc")
    if not BUILD["running"] or p is None:
        return JSONResponse({"ok": False, "detail": "当前没有正在运行的任务"}, status_code=409)
    BUILD["cancelled"] = True
    try:
        with _Q_LOCK:
            QUEUE["paused"] = True
            _q_persist()
    except Exception as e:
        log_error("build/cancel pause-queue", repr(e))
    try:
        _kill_tree(p)   # BF7：杀整棵进程树，孙进程不再漏网
    except Exception as e:
        log_error("build/cancel", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": f"终止失败：{e}"}, status_code=500)
    BUILD["log"].append("[build] 收到取消请求，正在停止…")
    return {"ok": True, "stage": BUILD.get("stage"), "queue_paused": True}

@app.get("/build/status")
def build_status():
    return {"running": BUILD["running"], "stage": BUILD["stage"],
            "log": BUILD["log"][-300:], "started": BUILD["started"],
            # cancellable：有活着的子进程句柄才能取消（供前端决定是否显示「取消」按钮）
            "cancellable": bool(BUILD["running"] and BUILD.get("proc") is not None),
            "cancelled": bool(BUILD.get("cancelled")),
            "rc": BUILD.get("rc")}   # 构建返回码：前端据此区分「更新完成」与「更新失败」（余额0/子进程崩溃）

@app.post("/index/retry_no_text")
def retry_no_text():
    """兼容旧路由名：清除 PDF 提取失败终态与旧产物，让用户修好附件后重新深索。"""
    notext = _deep_no_text_keys()   # safe_name(stem) 集合
    if not notext:
        return {"ok": True, "cleared": 0, "msg": "当前没有待重试的 PDF 提取失败篇。"}
    papers = _load_papers()
    keys = [k for k, p in papers.items()
            if (p.get("stem") or T.safe_name(k)) in notext]
    if not keys:
        return {"ok": True, "cleared": 0, "msg": "无正文标记对应的篇已不在库中。"}
    try:
        import folder_ingest as FI
        FI._purge_key_artifacts(keys)   # 清深索标记 + 删 extracted/chunks/pagemap/summaries
        import deep_extract_status as DES
        DES.remove([(papers.get(k) or {}).get("stem") or T.safe_name(k) for k in keys])
    except Exception as e:
        log_error("retry_no_text purge", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=500)
    return {"ok": True, "cleared": len(keys),
            "msg": f"已清除 {len(keys)} 篇的提取失败状态与旧产物；修好附件后重新深索，会自动尝试本地 OCR。"}

# ── 文件夹模式：建库（后台，含 N 次 LLM 抽题录）+ 拖入入库 ──
@app.post("/index/folder_build")
def index_folder_build():
    import settings as S
    if S.source() != "folder":
        return JSONResponse({"ok": False, "msg": "当前非文件夹模式"}, status_code=400)
    try:
        import folder_meta as FM
        ready = FM.available()
    except Exception:
        ready = False
    if not ready:
        return JSONResponse({"ok": False, "need_key": True,
                             "msg": "未配置 API Key（用于自动抽取题录）；也可先建粗库（文件名题名），配好 Key 后再更新"},
                            status_code=200)  # 200 + need_key：前端可选择继续无 key 建库
    return {"ok": _run_build("folder")}

@app.post("/index/folder_build_nokey")
def index_folder_build_nokey():
    """无 key 也建库：folder_ingest 在无 key 时正常退出（退回文件名题名），仍能建词法层。"""
    import settings as S
    if S.source() != "folder":
        return JSONResponse({"ok": False, "msg": "当前非文件夹模式"}, status_code=400)
    return {"ok": _run_build("folder")}

def _dedupe_name(dst):
    """同名不同容→加序号后缀，避免覆盖。"""
    if not dst.exists():
        return dst
    stem, suf = dst.stem, dst.suffix
    i = 2
    while True:
        cand = dst.with_name(f"{stem} ({i}){suf}")
        if not cand.exists():
            return cand
        i += 1

class IngestFile(BaseModel):
    name: str
    content_b64: str                              # 文件内容 base64（避免依赖 python-multipart）

class IngestQ(BaseModel):
    files: List[IngestFile] = []

@app.post("/ingest/files")
def ingest_files(q: IngestQ):
    """拖入/选择的 PDF（base64）复制进受管文件夹 → 去重 → 后台 folder build（抽题录+索引）。
       用 JSON base64 而非 multipart，避免分发版缺 python-multipart。"""
    import settings as S, hashlib as _hl, base64 as _b64
    if S.source() != "folder":
        return JSONResponse({"ok": False, "msg": "仅文件夹模式支持拖入入库"}, status_code=400)
    folder = S.folder_dir()
    if not folder:
        return JSONResponse({"ok": False, "msg": "未配置受管文件夹"}, status_code=400)
    if BUILD["running"]:
        return JSONResponse({"ok": False, "msg": "正在建库/入库中，请稍后再拖入", "building": True}, status_code=409)
    fp = Path(folder)
    # R2：先按文件大小建索引，只有大小撞车的候选才按需算 sha1（并缓存），
    # 避免每拖一篇就把受管文件夹里每个 PDF 全量读盘算 sha1（大库会卡几十秒~几分钟）。
    size_index = {}          # size -> [paths]
    for p in fp.rglob("*.pdf"):
        try:
            size_index.setdefault(p.stat().st_size, []).append(p)
        except Exception:
            pass
    _sha1_cache = {}         # path -> sha1（仅对大小相同的候选按需算一次）
    def _is_dup(data):
        cands = size_index.get(len(data))
        if not cands:
            return False
        h = _hl.sha1(data).hexdigest()
        for p in cands:
            hp = _sha1_cache.get(p)
            if hp is None:
                try:
                    hp = _hl.sha1(p.read_bytes()).hexdigest()
                except Exception:
                    hp = ""
                _sha1_cache[p] = hp
            if hp and hp == h:
                return True
        return False
    added, skipped, failed = [], [], []
    for f in (q.files or []):
        name = f.name or "untitled.pdf"
        if not name.lower().endswith(".pdf"):
            failed.append({"name": name, "reason": "非 PDF"}); continue
        try:
            data = _b64.b64decode((f.content_b64 or "").split(",")[-1])
            if _is_dup(data):
                skipped.append(name); continue
            dst = _dedupe_name(fp / Path(name).name)
            dst.write_bytes(data)
            # 把新写入的文件计入索引/缓存，供同批后续文件去重（原逻辑靠 existing.add）
            size_index.setdefault(len(data), []).append(dst)
            _sha1_cache[dst] = _hl.sha1(data).hexdigest()
            added.append(dst.name)
        except Exception as e:
            failed.append({"name": name, "reason": str(e)})
    try:
        import folder_meta as FM
        ready = FM.available()
    except Exception:
        ready = False
    building = False
    if added:
        building = _run_build("folder")   # 后台增量：只处理新 key（幂等）；无 key 时 folder_ingest 退文件名
    return {"ok": True, "added": len(added), "added_names": added,
            "skipped": len(skipped), "failed": failed,
            "building": building, "need_key": (not ready)}

@app.post("/index/semantic")
def index_semantic_ep():
    return {"ok": _run_build("semantic")}

class DeepQ(BaseModel):
    scope: str = "all"

@app.post("/index/deep")
def index_deep_ep(q: DeepQ):
    # C7/A1：手动深索不再静默丢弃。
    #  - scope="keys:k1,k2..."（勾选若干篇深索）→ 走持久队列 enqueue_deep：撞锁自动排队、
    #    崩溃可回灌续跑，绝不因忙而蒸发。返回真正入队数 queued。
    #  - scope="all"（整库深索）→ 忙时返回 {ok:false,busy:true} 让前端提示「已有任务在跑」，
    #    而非假装成功。
    scope = (q.scope or "all").strip()
    if scope.startswith("keys:"):
        raw = scope[len("keys:"):]
        keys = [k for k in (raw.split(",")) if k.strip()]
        queued = enqueue_deep(keys)
        return {"ok": True, "queued": len(queued), "scope": scope}
    started = _run_build("deep", ["--scope", scope])
    if not started:
        return {"ok": False, "busy": True, "scope": scope}
    return {"ok": True, "queued": scope, "scope": scope}

# ── #7：Agent 驱动深索（切块→Agent写摘要→带摘要嵌入，一趟完成，阻塞式逐批）──
class DeepAgentSummary(BaseModel):
    key: str
    summary: str = ""

class DeepAgentQ(BaseModel):
    summaries: Optional[List[DeepAgentSummary]] = None   # 带上一批 Agent 写的摘要则先写盘+嵌入
    batch: int = 15

def _run_stage_blocking(stage, extra=None):
    """阻塞跑一段 build_all（子进程 subprocess.run，捕获 stdout 不污染 HTTP 响应）。返回 returncode。"""
    env = _child_env()
    cmd = [sys.executable, str(C.APP / "build_all.py"), "--stage", stage] + (extra or [])
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, env=env,
                       creationflags=C.SUBPROC_NO_WINDOW)
    try:
        out = (p.stdout or b"").decode("utf-8", errors="replace")
        BUILD["log"].append(f"[deep_agent:{stage}] rc={p.returncode}")
        for ln in out.splitlines()[-20:]:
            BUILD["log"].append(ln)
        BUILD["log"] = BUILD["log"][-300:]
    except Exception:
        pass
    return p.returncode

def _extracted_excerpt(stem, limit=1800):
    """取该篇提取正文前 ~limit 字，供 Agent 写摘要。无提取文本（未切块/扫描件）→ ""。"""
    f = C.EXTRACTED / f"{stem}.json"
    if not f.exists():
        return ""
    try:
        rec = json.loads(f.read_text(encoding="utf-8"))
    except Exception:
        return ""
    txt = "\n".join((pg.get("text") or "") for pg in rec.get("pages", []))
    return txt[:limit]

def _deep_agent_run(q: DeepAgentQ):
    import sac as SAC
    wrote = 0
    # ① 带 summaries：写进 summaries.json，再阻塞跑 deep_embed 把上一批带摘要嵌入（标记已深索）
    if q.summaries:
        result = SAC.write_summaries([{"key": s.key, "summary": s.summary} for s in q.summaries])
        if result["errors"]:
            return {"ok": False, "stage": "summary_validation", "wrote": 0,
                    "summary_errors": result["errors"],
                    "error": "本批摘要未通过质量检查，整批未写入；请修正后重新提交"}
        wrote = result["written"]
        # BF16：接住子阶段退出码——此前 rc≠0（余额不足/网络断）也照样返回 ok:true，
        # agent 会向用户误报「已嵌入入库」，实际一篇都没进。
        rc = _run_stage_blocking("deep_embed")
        if rc != 0:
            return {"ok": False, "stage": "deep_embed",
                    "error": f"深索子阶段失败(rc={rc})：可能是 API 余额不足或网络问题，请稍后重试"}
        try:                              # 嵌入后重载索引，让新深索的篇立刻可检索
            R.STATE["ready"] = False
            R.load_all()
        except Exception as e:
            log_error("deep_agent reload", repr(e))
        # EN-W2 Ingest 扳机：Agent 深索循环也要触发 wiki 更新建议。本批 summaries 的 key 已成功嵌入并重载，
        # 逐一分析哪些综合页引用了它们、是否需标脏/重生（整库 pre_deep 差集路径不覆盖此 Agent 循环）。
        try:
            _wiki_suggest_async([s.key for s in q.summaries])
        except Exception as e:
            log_error("deep_agent wiki suggest", repr(e))
    # ② 选下一批 pending-deep（有PDF、未深索、非扫描件），阻塞跑 deep_prepare(extract+chunk)
    papers = _load_papers(); deepk = _deep_keys(); notext = _deep_no_text_keys()
    cand = [p for p in papers.values()
            if p.get("has_pdf") and not is_deep(p["key"], deepk)
            and T.safe_name(p["key"]) not in notext]
    batch_n = max(1, int(q.batch or 15))
    batch = cand[:batch_n]
    to_summarize = []
    if batch:
        keys = [p["key"] for p in batch]
        # BF16：deep_prepare 失败（抽取/切块子进程崩）时 excerpt 都是空的，agent 拿去写摘要毫无意义
        rc = _run_stage_blocking("deep_prepare", ["--scope", "keys:" + ",".join(keys)])
        if rc != 0:
            return {"ok": False, "stage": "deep_prepare",
                    "error": f"深索子阶段失败(rc={rc})：可能是 API 余额不足或网络问题，请稍后重试"}
        for p in batch:
            to_summarize.append({"key": p["key"], "title": p.get("title", ""),
                                 "excerpt": _extracted_excerpt(p["stem"])})
    manifest = json.loads(C.INDEX_MANIFEST.read_text(encoding="utf-8")) if C.INDEX_MANIFEST.exists() else {}
    # ③ 汇总返回。finished=没有更多待处理（末批：summaries 已提交且无新篇也能收尾）
    return {"ok": True, "wrote": wrote, "to_summarize": to_summarize,
            "done": len(_deep_keys()), "with_pdf": manifest.get("with_pdf", 0),
            "remaining": max(0, len(cand) - len(batch)),
            "finished": (len(batch) == 0)}

@app.post("/index/deep_agent")
def index_deep_agent(q: DeepAgentQ):
    """#7：Agent 驱动深索的单批处理（阻塞，可跑 30-60s）。尊重构建锁：有其它构建在跑→busy。
       用法见 MCP 工具 deep_index：首次不带 summaries→返回 to_summarize；带 summaries 再调→
       嵌入上一批+返回下一批；循环至 finished=true。"""
    with _BUILD_LOCK:
        if BUILD["running"]:
            return {"ok": False, "busy": True}
        BUILD["running"] = True; BUILD["stage"] = "deep_agent"
        BUILD["started"] = time.time(); BUILD["rc"] = None
        BUILD["cancelled"] = False       # BF14：每次构建开始都复位取消标记，避免上次取消残留误报
    try:
        return _deep_agent_run(q)
    except Exception as e:
        log_error("index/deep_agent", repr(e), traceback.format_exc())
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=200)
    finally:
        with _BUILD_LOCK:
            BUILD["running"] = False
            BUILD["stage"] = ""          # F4：deep_agent 结束同样复位 stage
        try:
            _drain_deep_queue()          # 与 _run_build/light 收尾对齐：补跑构建期被防抖丢弃的深索入队
        except Exception as e:
            log_error("deep_agent drain", repr(e))

# ── 补生成检索摘要（SAC backfill）────────────────────────────
# 给「已深索但没有检索摘要」的篇：读切块生成 ~150 字摘要 → 写 summaries.json →
# 撤销其深索标记 → 跑 deep_embed 重嵌入（带摘要前缀）。摘要只有重嵌入后才对检索生效，
# 所以这一步必然要重嵌入；切块/提取产物保留不删，embed_index add 前按 key 删旧行、无重复。
def _backfill_gap_stems(only=None):
    """可补生成的候选 stem：已深索、摘要缺失或异常、且正文可用。
       only 给定（stem 集合）时只取其中的，用于「为某一篇/某几篇生成」。"""
    deep = _deep_keys(); sumk = _summary_keys(); notext = _deep_no_text_keys()
    gap = [s for s in deep if s not in sumk and s not in notext]
    if only is not None:
        only = set(only)
        gap = [s for s in gap if s in only]
    return gap

def _chunk_title_body(stem, n_body=6):
    """读该 stem 的切块文件，取 title + 前 n_body 块正文供生成摘要。无块→(None, None)。"""
    f = C.CHUNKS / f"{stem}.json"
    if not f.exists():
        return None, None
    try:
        cs = json.loads(f.read_text(encoding="utf-8"))
    except Exception:
        return None, None
    if not cs:
        return None, None
    title = cs[0].get("title", "") or stem
    body = "\n".join(c.get("text", "") for c in cs[:n_body])
    return title, body

def _unmark_deep(stems):
    """从 embedded_keys.txt 撤销这些 stem 的「已深索」标记（保序重写，保留其它行）。
       调用时已持 BUILD 锁（无 build/embed 子进程在写此文件），单进程内安全。"""
    stems = set(s for s in stems if s)
    if not stems:
        return
    ek = C.STATE / "embedded_keys.txt"
    with _BACKFILL_LOCK:
        try:
            if not ek.exists():
                return
            lines = ek.read_text(encoding="utf-8").splitlines()
            kept = [l for l in lines if l.strip() and l.strip() not in stems]
            ek.write_text("".join(l + "\n" for l in kept), encoding="utf-8")
        except Exception as e:
            log_error("unmark_deep", repr(e))

def _sac_backfill_worker(gap):
    import sac as SAC
    try:
        def _items():
            for stem in gap:
                title, body = _chunk_title_body(stem)
                if title is None:      # 无切块（扫描件/产物缺失）→ 无法补，跳过
                    continue
                yield (stem, title, "", body)
        def _prog(n, fail):
            with _BACKFILL_LOCK:
                BACKFILL["done"] = n; BACKFILL["fail"] = fail
        n, fail = SAC.gen_missing(_items(), log=lambda m: BUILD["log"].append(m), on_progress=_prog)
        with _BACKFILL_LOCK:
            BACKFILL["done"] = n; BACKFILL["fail"] = fail
        if n <= 0:
            with _BACKFILL_LOCK:
                BACKFILL["phase"] = ""
                BACKFILL["msg"] = f"未生成摘要（失败 {fail}，请检查 API key / 额度 / 网络）。"
            return
        # 只对本轮真正拿到摘要的篇撤标 + 重嵌入
        newly = [s for s in gap if s in SAC.summary_keys()]
        _unmark_deep(newly)
        with _BACKFILL_LOCK:
            BACKFILL["phase"] = "重嵌入中"; BACKFILL["at"] = time.time()
        rc = _run_stage_blocking("deep_embed")
        try:
            R.STATE["ready"] = False; R.load_all()
        except Exception as e:
            log_error("sac_backfill reload", repr(e))
        with _BACKFILL_LOCK:
            BACKFILL["phase"] = ""
            BACKFILL["msg"] = (f"已为 {n} 篇补上检索摘要并重嵌入 ✓" if rc == 0
                               else f"已生成 {n} 篇摘要，但重嵌入失败(rc={rc})，可稍后重试补生成。")
    except Exception as e:
        log_error("sac_backfill worker", repr(e), traceback.format_exc())
        with _BACKFILL_LOCK:
            BACKFILL["phase"] = ""; BACKFILL["msg"] = "补生成异常：" + str(e)
    finally:
        with _BACKFILL_LOCK:
            BACKFILL["running"] = False; BACKFILL["at"] = time.time()
        with _BUILD_LOCK:
            BUILD["running"] = False; BUILD["stage"] = ""
        try:
            _drain_deep_queue()      # 收尾补跑构建期被防抖丢弃的深索入队
        except Exception as e:
            log_error("sac_backfill drain", repr(e))

class SacBackfillQ(BaseModel):
    keys: Optional[List[str]] = None   # 原始 key 列表；给定则只为这几篇生成（单篇/指定范围），空=全部缺摘要的篇

@app.post("/index/sac_backfill")
def index_sac_backfill(q: SacBackfillQ = None):
    """生成检索摘要（知识库建设第②步）：为「已深索但缺摘要」的篇生成摘要并重嵌入。
       后台跑，进度见 /index/status.sac_backfill。keys 给定则只为这几篇生成（浏览页单篇⚪点生成）。
       无论当前 generator=off/agent/server，用户点这里都强制用可用 API key 生成一次。"""
    import sac as SAC
    if not SAC.key_available():
        return {"ok": False, "need_key": True,
                "msg": "生成检索摘要需要 API key：请到 设置 → 检索 → 检索引擎 填一个 SiliconFlow 免费 Key（或在 设置 → 建库 → 检索摘要 → 高级 单独填），再来生成。"}
    only = None
    if q and q.keys:
        only = {T.safe_name(k) for k in q.keys if k}
    gap = _backfill_gap_stems(only=only)
    if not gap:
        return {"ok": True, "started": False, "total": 0,
                "msg": ("这篇已有检索摘要，无需再生成。" if only else "所有已深索文献都已有检索摘要，无需生成。")}
    with _BUILD_LOCK:
        if BUILD["running"]:
            return {"ok": False, "busy": True, "msg": "已有任务在跑（深索/更新/补生成），请稍后再试。"}
        BUILD["running"] = True; BUILD["stage"] = "sac_backfill"
        BUILD["started"] = time.time(); BUILD["rc"] = None; BUILD["cancelled"] = False
    with _BACKFILL_LOCK:
        BACKFILL.update({"running": True, "phase": "生成摘要中", "done": 0,
                         "total": len(gap), "fail": 0, "msg": "", "at": time.time()})
    threading.Thread(target=_sac_backfill_worker, args=(gap,), daemon=True).start()
    return {"ok": True, "started": True, "total": len(gap)}

@app.get("/summary")
def get_summary(key: str):
    """只读查看某篇的检索摘要（~150 字，AI 生成、给语义检索用的嵌入前缀，不是原文摘要）。
       返回 has_summary / summary / deep（该篇是否已深索——未深索无法生成摘要）。"""
    stem = T.safe_name(key)
    import sac as SAC
    info = SAC.inspect(stem)
    return {"key": key, "deep": stem in _deep_keys(),
            "has_summary": info["valid"], "summary": info["summary"],
            "summary_invalid": bool(info["reason"]), "summary_error": info["reason"]}

# ── 检索 ──────────────────────────────────────────────────
class SearchQ(BaseModel):
    query: str
    topk: int = C.RERANK_TOPK
    sort: Optional[str] = None
    min_weight: float = 0.0                       # 引用权重下限过滤（0=不过滤）
    category: Optional[str] = None                # F11：限定检索范围到某个分类（kbc_/topic:/zotero:）

@app.post("/search")
def search(q: SearchQ):
    if not R.STATE.get("ready"):
        # UX8：两态文案——建过库只是重建/重载窗口（稍等即可），从未建库才需要走向导；
        # 此前一律"先建立即时索引"，害老用户在重载的几秒里被误导去重建。
        msg = ("索引正在重建或加载，请稍候几秒再搜" if C.INDEX_MANIFEST.exists()
               else "还没建立索引——请到 设置 → 应用 → 重新查看引导 完成首次建库")
        return JSONResponse({"error": msg, "ready": False}, status_code=503)
    t0 = time.time()
    keys = _resolve_category_keys(q.category)
    try:
        res = R.search(q.query, q.topk, q.sort, q.min_weight, keys=keys)
    except Exception as e:
        # BF36：API 模式下嵌入/重排后端挂了（余额0/断网）此前抛成裸 500「服务器内部错误」，
        # 用户不知道该去哪修——detail 给人话，前端 jpost 已会读 detail 展示。
        log_error("search", repr(e), traceback.format_exc())
        raise HTTPException(status_code=500,
                            detail=f"{e}（→ 请到 设置 检查检索引擎 API Key 或余额）")
    # 给命中行补结构化提取状态；旧 no_text 字段只表示当前终态阻塞。
    try:
        notext = _deep_no_text_keys()
        extract_items = _deep_extract_items()
        papers = _load_papers()
        for x in res:
            k = x.get("key")
            if k:
                p = papers.get(k) or {}
                stem = p.get("stem") or T.safe_name(k)
                x["no_text"] = stem in notext
                x["extract_status"] = _extract_record(
                    stem, extract_items, legacy_deep=is_deep(k))
    except Exception as e:
        log_error("search no_text tag", repr(e))
    return {"query": q.query, "mode": R.STATE.get("mode"), "category": q.category,
            "took_ms": round((time.time() - t0) * 1000), "results": res}

# ── RAG 对话 ──────────────────────────────────────────────
class ChatQ(BaseModel):
    query: str
    history: List[dict] = []
    provider: str = "deepseek"
    base_url: str = ""
    api_key: str = ""
    model: str = ""
    topk: int = 6
    sort: Optional[str] = None
    category: Optional[str] = None                # F11：限定对话检索范围到某个分类

SYS_TMPL = (
    "你是严谨的法学文献研究助手。请**只依据下面提供的文献片段**回答用户问题，"
    "每个论点后用 [编号] 标注来源；若不足以回答请如实说明不要编造。回答用中文。\n\n=== 文献片段 ===\n{ctx}"
)

@app.post("/chat")
def chat(q: ChatQ):
    keys = _resolve_category_keys(q.category)
    # C6/A5：检索后端(嵌入/重排 API)挂了不能被误报成「模型没返回内容，请检查对话模型/Key」。
    # 把同步的 R.search 包 try：失败则 hits=[]、记下真因，稍后在 SSE 里先 yield error 再照常作答。
    search_err = None
    hits = []
    if R.STATE.get("ready"):
        try:
            hits = R.search(q.query, q.topk, q.sort, keys=keys)
        except Exception as e:
            log_error("chat search", repr(e))
            search_err = "检索后端暂时不可用（可能余额/网络/超时），本次未附引用"
    ctx = "\n\n".join(
        f"[{i+1}] {h.get('citation','')}\n{(h.get('context') or h.get('text') or '')[:1200]}"
        for i, h in enumerate(hits)
    ) or "（暂无检索结果）"
    messages = ([{"role": "system", "content": SYS_TMPL.format(ctx=ctx)}]
                + list(q.history) + [{"role": "user", "content": q.query}])
    base, model = L.resolve(q.provider, q.base_url, q.model)
    api_key = q.api_key
    # 对话选 SiliconFlow 但没单独填 key 时，自动复用检索引擎已配的 SiliconFlow key（一个 key 通吃，免费模型开箱即用）
    if not api_key and _is_siliconflow_base(base):
        try:
            import settings as S
            api_key = (S.api_conf() or {}).get("key", "") or api_key
        except Exception:
            pass

    def gen():
        yield "data: " + json.dumps({"sources": hits}, ensure_ascii=False) + "\n\n"
        if search_err:      # C6：先把检索失败的真因抛给前端 j.error 分支，再照常让模型作答
            yield "data: " + json.dumps({"error": search_err}, ensure_ascii=False) + "\n\n"
        try:
            for delta in L.chat_stream(messages, base, api_key, model):
                yield "data: " + json.dumps({"delta": delta}, ensure_ascii=False) + "\n\n"
        except Exception as e:
            log_error("chat", repr(e))
            yield "data: " + json.dumps({"error": str(e)}, ensure_ascii=False) + "\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")

# ── 备份与恢复 ────────────────────────────────────────────
# 设计见 backup.py 文件头。这里只负责：后台线程 + 进度 + 「正在建索引时拒绝」。
#
# ⚠️ 本节所有错误一律 **HTTP 200 + {"ok": false, "error": ...}**，不用 4xx + msg。
#    原因见 CLAUDE.md §6：前端 jpost() 在非 2xx 时只读 detail/error、**不读 msg**，
#    真实原因会被吞成「/backup/create 400」这种没用的提示。
BACKUP = {"running": False, "stage": "", "done": 0, "total": 0,
          "result": None, "error": None}


class BackupQ(BaseModel):
    include_index: Optional[bool] = False
    include_key: Optional[bool] = False


class BackupPathQ(BaseModel):
    path: str


def _backup_busy_reason():
    """能不能现在动数据？—— 建索引时打包会得到一个『看着正常、实则损坏』的副本。"""
    if BACKUP["running"]:
        return "已经有一个备份/恢复任务在跑了，等它结束。"
    if BUILD["running"]:
        return "正在建索引（或深索），此时打包会得到损坏的向量库。等它跑完再备份。"
    return None


def _auto_backup_tick():
    """自动备份（由上面的定时循环每分钟叫一次；真正动手是每 N 天一次）。

    刻意的两个决定：
      · **只打「手写资产」，不含索引** —— 自动备份要是每次几个 G，用户的云盘和硬盘都会哭。
        换机时想连索引一起带走，用手动备份勾「包含向量索引」。
      · **永不含 API key** —— 自动备份的包会一声不响地躺进云盘。要含 key 只能手动勾选。
    正在建索引就跳过，下一轮再来（打包一个写到一半的向量库 = 一个损坏的副本）。
    """
    import settings as S
    conf = S.load().get("backup") or {}
    if not conf.get("auto"):
        return
    if BUILD["running"] or BACKUP["running"]:
        return

    days = max(1, int(conf.get("every_days") or 7))
    last = conf.get("last_at") or ""
    if last:
        try:
            t = time.mktime(time.strptime(last, "%Y-%m-%d %H:%M:%S"))
            if time.time() - t < days * 86400:
                return
        except Exception:
            pass                      # 时间戳坏了 → 当作没备份过，宁可多备一次

    import backup as BK
    BACKUP.update({"running": True, "stage": "自动备份中", "done": 0, "total": 0,
                   "result": None, "error": None})
    try:
        m = BK.create(include_index=False, include_key=False)
        S.save({"backup": {"last_at": m["created"]}})
        BACKUP["result"] = m
        BACKUP["stage"] = "完成"
        BUILD["log"].append(f"[auto] 已自动备份 → {m['path']}（{m['size'] / 1e6:.1f} MB）")
    except Exception as e:
        BACKUP["error"] = str(e)
        BACKUP["stage"] = "失败"
        log_error("auto_backup", repr(e))
    finally:
        BACKUP["running"] = False


class BackupConfQ(BaseModel):
    dir: Optional[str] = None
    auto: Optional[bool] = None
    every_days: Optional[int] = None
    keep: Optional[int] = None
    include_index: Optional[bool] = None


@app.get("/backup/config")
def backup_config_get():
    import settings as S
    import backup as BK
    c = dict(S.DEFAULT["backup"])
    c.update(S.load().get("backup") or {})
    c["effective_dir"] = str(BK.backup_dir())    # dir 为空时的实际落点，前端拿来当 placeholder
    return {"ok": True, **c}


@app.post("/backup/config")
def backup_config_set(q: BackupConfQ):
    import settings as S
    patch = {k: v for k, v in q.dict().items() if v is not None}
    if "dir" in patch and patch["dir"]:
        p = Path(patch["dir"])
        try:
            p.mkdir(parents=True, exist_ok=True)
            t = p / ".write_test"
            t.write_text("x", encoding="utf-8")
            t.unlink()
        except Exception as e:
            return {"ok": False, "error": f"这个目录写不进去：{p}（{e}）"}
    if patch:
        S.save({"backup": patch})
    return backup_config_get()


@app.get("/backup/list")
def backup_list():
    import backup as BK
    try:
        return {"ok": True, "dir": str(BK.backup_dir()), "items": BK.list_backups()}
    except Exception as e:
        log_error("backup_list", repr(e))
        return {"ok": False, "error": str(e), "items": []}


@app.get("/backup/estimate")
def backup_estimate(with_index: int = 0):
    """让用户在点之前就知道「这一下要写 3.2 GB」，而不是点完等半天。"""
    import backup as BK
    try:
        return {"ok": True, "bytes": BK.estimate(bool(with_index))}
    except Exception as e:
        return {"ok": False, "error": str(e), "bytes": 0}


@app.post("/backup/create")
def backup_create(q: BackupQ = None):
    q = q or BackupQ()
    busy = _backup_busy_reason()
    if busy:
        return {"ok": False, "error": busy}

    import backup as BK
    BACKUP.update({"running": True, "stage": "打包中", "done": 0, "total": 0,
                   "result": None, "error": None})

    def work():
        try:
            def prog(d, t):
                BACKUP["done"], BACKUP["total"] = d, t
            m = BK.create(include_index=bool(q.include_index),
                          include_key=bool(q.include_key), on_progress=prog)
            try:
                import settings as S
                S.save({"backup": {"last_at": m["created"]}})
            except Exception:
                pass
            BACKUP["result"] = m
            BACKUP["stage"] = "完成"
        except Exception as e:
            log_error("backup_create", repr(e))
            BACKUP["error"] = str(e)
            BACKUP["stage"] = "失败"
        finally:
            BACKUP["running"] = False

    threading.Thread(target=work, daemon=True).start()
    return {"ok": True, "started": True}


@app.get("/backup/status")
def backup_status():
    return {"ok": True, **{k: v for k, v in BACKUP.items()}}


@app.post("/backup/inspect")
def backup_inspect(q: BackupPathQ):
    import backup as BK
    try:
        return {"ok": True, **BK.inspect(q.path)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.post("/backup/restore")
def backup_restore(q: BackupPathQ):
    busy = _backup_busy_reason()
    if busy:
        return {"ok": False, "error": busy}

    import backup as BK
    info = BK.inspect(q.path)
    if not info.get("ok"):
        return {"ok": False, "error": info.get("err") or "备份包不可用"}

    BACKUP.update({"running": True, "stage": "恢复中", "done": 0, "total": 0,
                   "result": None, "error": None})

    def work():
        try:
            r = BK.restore(q.path)
            if r.get("ok"):
                BACKUP["result"] = r
                BACKUP["stage"] = "完成（需重启）"
            else:
                BACKUP["error"] = r.get("err") or "恢复失败"
                BACKUP["stage"] = "失败"
        except Exception as e:
            log_error("backup_restore", repr(e))
            BACKUP["error"] = str(e)
            BACKUP["stage"] = "失败"
        finally:
            BACKUP["running"] = False

    threading.Thread(target=work, daemon=True).start()
    return {"ok": True, "started": True}


@app.post("/backup/open_dir")
def backup_open_dir():
    """在资源管理器里打开备份目录（用户要把 zip 拷走/确认云盘同步了没）。"""
    import backup as BK
    d = BK.backup_dir()
    try:
        d.mkdir(parents=True, exist_ok=True)
        if sys.platform == "win32":
            os.startfile(str(d))  # noqa
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(d)])
        else:
            subprocess.Popen(["xdg-open", str(d)])
        return {"ok": True, "dir": str(d)}
    except Exception as e:
        return {"ok": False, "error": str(e), "dir": str(d)}


# ── 应用内更新（版本升级）────────────────────────────────
# 只换 app\（纯代码），data\ / 0_Agent* / models\ 一律不碰 —— 数据安全的保证全在 updater.apply()。
# 分工：server 负责「查新版 + 下载增量包」；真正的替换由 launcher 的 pywebview 桥
# apply_update() 拉起独立 updater 进程执行（因为要先关掉应用、解锁 app\ 里的 .py）。
UPDATE = {"checking": False, "info": None, "checked_at": 0,
          "downloading": False, "done": 0, "total": 0, "zip": None, "error": None}


@app.get("/update/check")
def update_check(force: int = 0):
    """查 GitHub 有没有新版。结果缓存 10 分钟，避免频繁打 GitHub API。"""
    import updater as UPD
    now = time.time()
    if not force and UPDATE["info"] and (now - UPDATE["checked_at"] < 600):
        return {"ok": True, "cached": True, **UPDATE["info"]}
    try:
        info = UPD.check()
        UPDATE["info"] = info
        UPDATE["checked_at"] = now
        return {"ok": True, "cached": False, **info}
    except Exception as e:
        log_error("update_check", repr(e))
        return {"ok": False, "error": str(e)}


@app.post("/update/open_installer")
def update_open_installer():
    """在默认浏览器打开当前 Release 的完整安装器；URL 只能来自官方 GitHub Release。"""
    import updater as UPD
    try:
        info = UPDATE.get("info") or UPD.check()
        url = str(info.get("installer_url") or "")
        prefix = f"https://github.com/{UPD.GH_OWNER}/{UPD.GH_REPO}/releases/download/"
        if not url.startswith(prefix) or not url.lower().endswith(".exe"):
            raise ValueError("当前 Release 没有可用的完整安装器")
        if sys.platform == "win32":
            os.startfile(url)  # noqa: S606 —— 固定官方 HTTPS 前缀，交给默认浏览器
        else:
            import webbrowser
            if not webbrowser.open(url):
                raise OSError("系统浏览器未接受打开请求")
        return {"ok": True, "url": url}
    except Exception as e:
        log_error("update open installer", repr(e))
        return JSONResponse({"ok": False, "detail": str(e)}, status_code=400)


@app.post("/update/download")
def update_download():
    """下载 app 增量包到 update\（后台跑，进度见 /update/status）。"""
    import updater as UPD
    if UPDATE["downloading"]:
        return {"ok": False, "error": "已经在下载了"}
    info = UPDATE["info"] or UPD.check()
    if not info.get("has_update"):
        return {"ok": False, "error": "当前已是最新版，没有可下载的更新"}

    UPDATE.update({"downloading": True, "done": 0, "total": 0, "zip": None, "error": None})

    def work():
        try:
            def prog(d, t):
                UPDATE["done"], UPDATE["total"] = d, t
            r = UPD.download(info, progress=prog)
            if r.get("ok"):
                UPDATE["zip"] = r["zip"]
            else:
                UPDATE["error"] = r.get("error") or "下载失败"
        except Exception as e:
            log_error("update_download", repr(e))
            UPDATE["error"] = str(e)
        finally:
            UPDATE["downloading"] = False

    threading.Thread(target=work, daemon=True).start()
    return {"ok": True, "started": True}


@app.get("/update/status")
def update_status():
    return {"ok": True, **{k: UPDATE[k] for k in
            ("downloading", "done", "total", "zip", "error")}}


class MirrorQ(BaseModel):
    mirror_base: str = ""


@app.get("/update/mirror")
def update_mirror_get():
    import settings as S
    return {"ok": True, "mirror_base": (S.load().get("update") or {}).get("mirror_base") or ""}


@app.post("/update/mirror")
def update_mirror_set(q: MirrorQ):
    import settings as S
    v = (q.mirror_base or "").strip()
    S.save({"update": {"mirror_base": v}})
    return {"ok": True, "mirror_base": v}


# ── 静态 UI ───────────────────────────────────────────────
WEB = C.APP / "web"

@app.get("/")
def index():
    # 任务六：显式 charset=utf-8，避免 WebView2 按系统 GBK 解码整页
    return FileResponse(str(WEB / "index.html"), media_type="text/html; charset=utf-8")

app.mount("/static", StaticFiles(directory=str(WEB)), name="static")

if __name__ == "__main__":
    print(f"本地知识库服务启动：http://{C.DAEMON_HOST}:{C.DAEMON_PORT}", flush=True)
    uvicorn.run(app, host=C.DAEMON_HOST, port=C.DAEMON_PORT, log_level="warning")
